#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
MySQL巡检日志分析器
解析Shell脚本生成的巡检日志，分析问题并生成Word报告
"""

import os
import re
import logging
from datetime import datetime
from typing import Dict, List, Any, Tuple

# 尝试导入依赖包
try:
    from docx import Document
    from docx.shared import Inches, Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 全局字体设置
FONT_NAME = "微软雅黑"


class LogAnalyzer:
    """巡检日志分析器"""
    
    def __init__(self, log_file: str):
        self.log_file = log_file
        self.logger = logging.getLogger(__name__)
        self.inspection_data = {}
        
    def parse_log_file(self) -> Dict[str, Any]:
        """解析巡检日志文件"""
        self.logger.info(f"开始解析日志文件: {self.log_file}")
        
        if not os.path.exists(self.log_file):
            raise FileNotFoundError(f"日志文件不存在: {self.log_file}")
        
        with open(self.log_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # 解析日志内容
        current_section = ""
        section_data = []
        
        # 定义需要过滤掉的日志前缀
        filtered_prefixes = [
            "\033[0;34m[INFO]\033[0m",   # log_info的ANSI颜色标记
            "\033[0;32m[SUCCESS]\033[0m", # log_success的ANSI颜色标记
            "[INFO]",                   # 纯文本格式的INFO标记
            "[SUCCESS]"                 # 纯文本格式的SUCCESS标记
        ]
        
        for line in lines:
            original_line = line.rstrip('\n')
            
            # 检测是否是需要过滤掉的log_info或log_success行
            stripped_line = original_line.strip()
            should_skip = False
            for prefix in filtered_prefixes:
                if stripped_line.startswith(prefix):
                    should_skip = True
                    break
            if should_skip:
                continue
            
            # 检测章节开始
            if stripped_line.startswith("===") and stripped_line.endswith("==="):
                # 保存上一个章节的数据
                if current_section and section_data:
                    self.inspection_data[current_section] = section_data
                
                # 开始新章节
                current_section = stripped_line.replace("===", "").strip()
                section_data = []
            elif current_section:
                # 添加原始行（保留制表符），但只添加非空行
                if original_line.strip():
                    section_data.append(original_line)
        
        # 保存最后一个章节
        if current_section and section_data:
            self.inspection_data[current_section] = section_data
        
        self.logger.info(f"日志解析完成，共找到 {len(self.inspection_data)} 个章节")
        return self.inspection_data
    
    def analyze_issues(self) -> Dict[str, Any]:
        """分析巡检问题"""
        if not self.inspection_data:
            self.parse_log_file()
        
        issues_summary = {
            "critical_issues": [],
            "warning_issues": [],
            "info_issues": [],
            "summary": {
                "total_issues": 0,
                "critical_count": 0,
                "warning_count": 0,
                "info_count": 0
            }
        }
        
        # 分析性能问题
        self._analyze_performance_issues(issues_summary)
        
        # 分析安全问题
        self._analyze_security_issues(issues_summary)
        
        # 分析存储问题
        self._analyze_storage_issues(issues_summary)
        
        # 分析系统资源问题
        self._analyze_system_issues(issues_summary)
        
        # 分析备份复制问题
        self._analyze_backup_replication_issues(issues_summary)
        
        # 分析错误日志问题
        self._analyze_error_logs(issues_summary)
        
        # 分析所有巡检项，确保没有遗漏的问题
        self._analyze_all_inspection_items(issues_summary)
        
        # 更新统计信息
        issues_summary["summary"]["critical_count"] = len(issues_summary["critical_issues"])
        issues_summary["summary"]["warning_count"] = len(issues_summary["warning_issues"])
        issues_summary["summary"]["info_count"] = len(issues_summary["info_issues"])
        issues_summary["summary"]["total_issues"] = (
            issues_summary["summary"]["critical_count"] + 
            issues_summary["summary"]["warning_count"] + 
            issues_summary["summary"]["info_count"]
        )
        
        # 最终去重：检查严重问题中是否包含已被归入警告的同名问题，反之亦然
        # 策略：对每个问题保留最高严重级别，移除低级别的同名问题
        issue_severity_map = {}
        for idx, issue in enumerate(issues_summary["critical_issues"]):
            key = (issue.get("category",""), issue.get("issue",""))
            issue_severity_map[key] = ("critical", idx)
        kept_warning = []
        for idx, issue in enumerate(issues_summary["warning_issues"]):
            key = (issue.get("category",""), issue.get("issue",""))
            if key in issue_severity_map:
                continue
            issue_severity_map[key] = ("warning", idx)
            kept_warning.append(issue)
        issues_summary["warning_issues"] = kept_warning
        kept_info = []
        for idx, issue in enumerate(issues_summary["info_issues"]):
            key = (issue.get("category",""), issue.get("issue",""))
            if key in issue_severity_map:
                continue
            kept_info.append(issue)
        issues_summary["info_issues"] = kept_info
        
        return issues_summary
    
    def _analyze_all_inspection_items(self, issues_summary: Dict[str, Any]):
        """分析所有巡检项，确保没有遗漏的问题
        注意：已被 _analyze_*_issues 专项方法处理过的巡检项在此跳过，避免重复
        """
        # 被专项分析方法覆盖的巡检项，跳过避免重复归类
        items_handled_by_specialized = {
            "TABLES_WITHOUT_PRIMARY_KEY", "TABLE_FRAGMENTATION_DETAILED",
            "EMPTY_PASSWORD_USERS", "REMOTE_ROOT_USERS", "REMOTE_USERS", "SSL_STATUS",
            "SYSTEM_CPU_USAGE", "SYSTEM_MEMORY_USAGE", "SYSTEM_SWAP_USAGE", "SYSTEM_DISK_USAGE",
            "MYSQL_PROCESS", "FILE_DESCRIPTOR_LIMIT", "NETWORK_CONNECTIONS", "MYSQL_ERROR_LOG",
            "BINARY_LOGS", "GTID_MODE", "GTID_CONSISTENCY", "GTID_EXECUTED",
            "SLAVE_STATUS", "MASTER_STATUS", "SLOW_QUERIES", "TABLE_LOCKS_WAITED",
            "DATABASE_SIZES", "PARTITIONED_TABLES",
        }
        # 定义巡检项分类（与 _generate_inspection_details 中一致）
        inspection_categories = {
            "数据库信息": ["DATABASE_VERSION", "UPTIME", "MAX_CONNECTIONS", "INNODB_BUFFER_POOL_SIZE", "OS_VERSION", 
                         "INNODB_LOG_FILE_SIZE", "INNODB_LOG_FILES_IN_GROUP", "INNODB_FLUSH_LOG_AT_TRX_COMMIT",
                         "SYNC_BINLOG", "BINLOG_FORMAT", "TRANSACTION_ISOLATION", 
                         "INNODB_FLUSH_METHOD", "INNODB_FILE_PER_TABLE", "OPEN_FILES_LIMIT", 
                         "TABLE_OPEN_CACHE", "MAX_ALLOWED_PACKET", "WAIT_TIMEOUT", "INTERACTIVE_TIMEOUT"],
            "性能指标": ["CURRENT_CONNECTIONS", "QUERY_CACHE_HITS", "QUERY_CACHE_INSERTS", "INNODB_READS", 
                        "INNODB_READ_REQUESTS", "SLOW_QUERIES", "TABLE_LOCKS_WAITED", "QUESTIONS", "QUERIES",
                        "COM_SELECT", "COM_INSERT", "COM_UPDATE", "COM_DELETE", "INNODB_ROW_LOCK_WAITS",
                        "INNODB_ROW_LOCK_TIME_AVG", "INNODB_ROW_LOCK_TIME_MAX", "CREATED_TMP_TABLES",
                        "CREATED_TMP_DISK_TABLES", "ABORTED_CLIENTS", "ABORTED_CONNECTS", 
                        "COM_COMMIT", "COM_ROLLBACK", "INNODB_ENGINE_STATUS"],
            "安全配置": ["EMPTY_PASSWORD_USERS", "REMOTE_ROOT_USERS", "SSL_STATUS", "REMOTE_USERS"],
            "存储信息": ["DATABASE_SIZES", "TABLE_FRAGMENTATION_DETAILED", 
                       "TABLES_WITHOUT_PRIMARY_KEY", "PARTITIONED_TABLES"],
            "主从复制": ["MASTER_STATUS", "SLAVE_STATUS", "BINARY_LOGS", "GTID_MODE", "GTID_CONSISTENCY", "GTID_EXECUTED"],
            "系统资源": ["SYSTEM_CPU_USAGE", "SYSTEM_MEMORY_USAGE", "SYSTEM_SWAP_USAGE", "SYSTEM_DISK_USAGE", 
                        "MYSQL_PROCESS", "FILE_DESCRIPTOR_LIMIT", "NETWORK_CONNECTIONS", "MYSQL_ERROR_LOG"]
        }
        
        # 动态添加集群信息分类
        has_mgr_data = (self.inspection_data.get("MGR_GROUP_MEMBERS", []) and 
                       len(self.inspection_data.get("MGR_GROUP_MEMBERS", [])) > 1) or \
                      (self.inspection_data.get("MGR_MEMBER_STATS", []) and 
                       len(self.inspection_data.get("MGR_MEMBER_STATS", [])) > 1)
        has_pxc_data = (self.inspection_data.get("PXC_WSREP_STATUS", []) and 
                       len(self.inspection_data.get("PXC_WSREP_STATUS", [])) > 1) or \
                      (self.inspection_data.get("PXC_WSREP_VARIABLES", []) and 
                       len(self.inspection_data.get("PXC_WSREP_VARIABLES", [])) > 1)
        
        if has_mgr_data:
            inspection_categories["集群信息（MGR）"] = ["MGR_GROUP_MEMBERS"]  # 只分析一个项，避免重复
        elif has_pxc_data:
            inspection_categories["集群信息（PXC）"] = ["PXC_WSREP_STATUS"]  # 只分析一个项，避免重复
        
        # 定义问题分类映射
        category_map = {
            "数据库信息": "配置",
            "性能指标": "性能",
            "安全配置": "安全",
            "存储信息": "存储",
            "主从复制": "备份复制",
            "系统资源": "系统资源",
            "集群信息（MGR）": "集群",
            "集群信息（PXC）": "集群"
        }
        
        # 分析每个巡检项
        for category, items in inspection_categories.items():
            issue_category = category_map.get(category, "其他")
            for item in items:
                # 跳过已被专项分析方法覆盖的巡检项，避免重复归类
                if item in items_handled_by_specialized:
                    continue
                item_data = self.inspection_data.get(item, [])
                analysis_result = self._analyze_inspection_item(item, item_data)
                
                # 如果状态不是正常，则添加到问题汇总
                if analysis_result["status"] == "异常":
                    # 检查是否已经添加过该问题（避免重复）
                    already_added = False
                    for existing_issue in issues_summary["critical_issues"]:
                        if existing_issue["issue"] == self._get_item_display_name(item):
                            already_added = True
                            break
                    if not already_added:
                        issues_summary["critical_issues"].append({
                            "category": issue_category,
                            "issue": self._get_item_display_name(item),
                            "description": analysis_result["suggestion"],
                            "suggestion": analysis_result["suggestion"]
                        })
                elif analysis_result["status"] == "注意":
                    # 检查是否已经添加过该问题（避免重复）
                    already_added = False
                    for existing_issue in issues_summary["warning_issues"]:
                        if existing_issue["issue"] == self._get_item_display_name(item):
                            already_added = True
                            break
                    if not already_added:
                        issues_summary["warning_issues"].append({
                            "category": issue_category,
                            "issue": self._get_item_display_name(item),
                            "description": analysis_result["suggestion"],
                            "suggestion": analysis_result["suggestion"]
                        })
    
    def _analyze_performance_issues(self, issues_summary: Dict[str, Any]):
        """分析性能问题"""
        # 连接数使用率分析
        max_conn = self._get_value("MAX_CONNECTIONS", 1)
        current_conn = self._get_value("CURRENT_CONNECTIONS", 1)
        
        if max_conn and current_conn:
            usage_rate = current_conn / max_conn
            if usage_rate > 0.9:
                issues_summary["critical_issues"].append({
                    "category": "性能",
                    "issue": "连接数使用率过高",
                    "description": f"当前连接数: {current_conn}, 最大连接数: {max_conn}, 使用率: {usage_rate:.1%}",
                    "suggestion": "立即增加max_connections参数，检查应用连接池配置，避免连接耗尽导致服务中断"
                })
            elif usage_rate > 0.8:
                issues_summary["warning_issues"].append({
                    "category": "性能",
                    "issue": "连接数使用率较高",
                    "description": f"当前连接数: {current_conn}, 最大连接数: {max_conn}, 使用率: {usage_rate:.1%}",
                    "suggestion": "建议监控连接数趋势，适时调整max_connections参数"
                })
        
        # InnoDB缓冲池命中率分析
        innodb_reads = self._get_value("INNODB_READS", 1)
        innodb_requests = self._get_value("INNODB_READ_REQUESTS", 1)
        
        if innodb_reads and innodb_requests and (innodb_reads + innodb_requests) > 0:
            hit_ratio = innodb_requests / (innodb_reads + innodb_requests)
            if hit_ratio < 0.9:
                issues_summary["critical_issues"].append({
                    "category": "性能",
                    "issue": "InnoDB缓冲池命中率过低",
                    "description": f"缓冲池命中率: {hit_ratio:.1%}（低于90%阈值）",
                    "suggestion": "立即增加innodb_buffer_pool_size参数，建议设置为物理内存的70-80%"
                })
            elif hit_ratio < 0.95:
                issues_summary["warning_issues"].append({
                    "category": "性能",
                    "issue": "InnoDB缓冲池命中率偏低",
                    "description": f"缓冲池命中率: {hit_ratio:.1%}（低于95%阈值）",
                    "suggestion": "建议增加innodb_buffer_pool_size参数，监控命中率变化"
                })
        
        # 慢查询分析
        slow_queries = self._get_value("SLOW_QUERIES", 1)
        if slow_queries and slow_queries > 0:
            if slow_queries > 100:
                issues_summary["critical_issues"].append({
                    "category": "性能",
                    "issue": "慢查询数量过多",
                    "description": f"累计慢查询数量: {slow_queries}（超过100个）",
                    "suggestion": "立即检查慢查询日志，优化SQL语句，添加必要的索引"
                })
            elif slow_queries > 10:
                issues_summary["warning_issues"].append({
                    "category": "性能",
                    "issue": "存在较多慢查询",
                    "description": f"累计慢查询数量: {slow_queries}",
                    "suggestion": "建议定期检查慢查询日志，优化SQL性能"
                })
            else:
                issues_summary["info_issues"].append({
                    "category": "性能",
                    "issue": "存在少量慢查询",
                    "description": f"累计慢查询数量: {slow_queries}",
                    "suggestion": "建议监控慢查询趋势，适时优化"
                })
    
    def _analyze_security_issues(self, issues_summary: Dict[str, Any]):
        """分析安全问题"""
        # 空密码用户检查
        empty_password_section = self.inspection_data.get("EMPTY_PASSWORD_USERS", [])
        if len(empty_password_section) > 0:
            # 提取实际用户信息
            users = []
            for line in empty_password_section:
                if line.strip():
                    parts = line.split('\t')
                    if len(parts) >= 2:
                        users.append(f"{parts[0]}@{parts[1]}")
            
            if users:
                users_str = ", ".join(users)
                issues_summary["critical_issues"].append({
                    "category": "安全",
                    "issue": "存在空密码用户",
                    "description": f"发现 {len(users)} 个空密码用户: {users_str}",
                    "suggestion": f"立即为以下用户设置强密码: {users_str}，建议使用密码复杂度策略"
                })
        
        # 远程root用户检查
        remote_root_section = self.inspection_data.get("REMOTE_ROOT_USERS", [])
        if len(remote_root_section) > 0:
            # 提取实际用户信息
            users = []
            for line in remote_root_section:
                if line.strip():
                    parts = line.split('\t')
                    if len(parts) >= 2:
                        users.append(f"{parts[0]}@{parts[1]}")
            
            if users:
                users_str = ", ".join(users)
                issues_summary["critical_issues"].append({
                    "category": "安全",
                    "issue": "存在远程root用户",
                    "description": f"发现 {len(users)} 个远程root用户: {users_str}",
                    "suggestion": f"立即修改以下用户的host为localhost或127.0.0.1: {users_str}"
                })
        
        # SSL状态检查
        ssl_status = self._get_value("SSL_STATUS", 1)
        if ssl_status and ssl_status.upper() != "YES":
            issues_summary["warning_issues"].append({
                "category": "安全",
                "issue": "SSL未启用",
                "description": "数据传输未加密，存在安全风险",
                "suggestion": "建议启用SSL加密传输，提高数据安全性"
            })
        
        # 远程用户检查
        remote_users_section = self.inspection_data.get("REMOTE_USERS", [])
        if len(remote_users_section) > 0:
            # 提取实际用户信息
            users = []
            for line in remote_users_section:
                if line.strip():
                    parts = line.split('\t')
                    if len(parts) >= 2:
                        users.append(f"{parts[0]}@{parts[1]}")
            
            if users:
                users_str = ", ".join(users)
                issues_summary["warning_issues"].append({
                    "category": "安全",
                    "issue": "存在允许任意主机访问的用户",
                    "description": f"发现 {len(users)} 个允许任意主机访问的用户: {users_str}",
                    "suggestion": f"建议限制以下用户的访问来源为具体IP: {users_str}"
                })
    
    def _analyze_storage_issues(self, issues_summary: Dict[str, Any]):
        """分析存储问题"""
        # 表碎片检查（使用详细数据，筛选碎片大小>2GB且碎片率>20%的表）
        fragmented_tables_section = self.inspection_data.get("TABLE_FRAGMENTATION_DETAILED", [])
        if len(fragmented_tables_section) > 0:
            # 提取实际的表信息，筛选碎片大小>2GB且碎片率>20%的表
            tables = []
            for line in fragmented_tables_section:
                if line.strip():
                    parts = line.split('\t')
                    if len(parts) >= 7:
                        try:
                            frag_size_mb = float(parts[5])  # 碎片大小
                            frag_percent = float(parts[6])  # 碎片率
                            # 过滤条件：碎片大小 > 2048MB (2GB) 且 碎片率 > 20%
                            if frag_size_mb > 2048.0 and frag_percent > 20.0:
                                tables.append(f"{parts[0]}.{parts[1]}")
                        except (ValueError, IndexError):
                            continue
            
            if tables:
                if len(tables) > 10:
                    issues_summary["critical_issues"].append({
                        "category": "存储",
                        "issue": "碎片化表过多",
                        "description": f"发现 {len(tables)} 个碎片大小>2GB且碎片率>20%的表",
                        "suggestion": "立即对碎片较大的表进行OPTIMIZE TABLE操作，定期维护表结构"
                    })
                else:
                    issues_summary["warning_issues"].append({
                        "category": "存储",
                        "issue": "存在表碎片",
                        "description": f"发现 {len(tables)} 个碎片大小>2GB且碎片率>20%的表",
                        "suggestion": "建议对碎片较大的表进行OPTIMIZE TABLE操作"
                    })
        
        # 数据库大小监控
        db_sizes_section = self.inspection_data.get("DATABASE_SIZES", [])
        if db_sizes_section and len(db_sizes_section) > 0:
            total_size = 0
            large_dbs = []
            for line in db_sizes_section:
                line_stripped = line.strip()
                # 跳过空行和标题行
                if not line_stripped or line_stripped.lower().startswith('database_name'):
                    continue
                parts = line_stripped.split('\t')
                if len(parts) >= 2:
                    try:
                        size_mb = float(parts[1])
                        total_size += size_mb
                        if size_mb > 102400:  # 100GB阈值
                            large_dbs.append(f"{parts[0]}({size_mb:.1f}MB)")
                    except (ValueError, IndexError):
                        continue
            
            # 总数据库大小检查
            if total_size > 1048576:  # 1TB
                issues_summary["critical_issues"].append({
                    "category": "存储",
                    "issue": "数据库总容量过大",
                    "description": f"数据库总容量: {total_size:.1f}MB（超过1TB）",
                    "suggestion": "立即清理历史数据，考虑数据归档或分库分表策略"
                })
            elif total_size > 524288:  # 500GB
                issues_summary["warning_issues"].append({
                    "category": "存储",
                    "issue": "数据库容量较大",
                    "description": f"数据库总容量: {total_size:.1f}MB（超过500GB）",
                    "suggestion": "建议定期清理历史数据，监控存储空间使用"
                })
            
            # 大型数据库检查
            if large_dbs:
                if len(large_dbs) > 3:
                    issues_summary["critical_issues"].append({
                        "category": "存储",
                        "issue": "大型数据库过多",
                        "description": f"发现 {len(large_dbs)} 个超过100GB的数据库",
                        "suggestion": "立即优化大型数据库，考虑分表或数据归档"
                    })
                else:
                    issues_summary["info_issues"].append({
                        "category": "存储",
                        "issue": "存在大型数据库",
                        "description": f"大型数据库: {', '.join(large_dbs)}",
                        "suggestion": "建议定期清理历史数据或考虑分库分表"
                    })
        
        # 无主键表检查
        no_pk_tables_section = self.inspection_data.get("TABLES_WITHOUT_PRIMARY_KEY", [])
        if no_pk_tables_section:
            # 提取实际的表信息 - 已去重
            tables = []
            seen_pk = set()
            for line in no_pk_tables_section:
                if line.strip() and not line.startswith('database_name'):
                    if ':' in line and not line.startswith('*'):
                        parts = line.split(':', 1)
                        if len(parts) >= 2:
                            line_content = parts[1].strip()
                            if '\t' in line_content:
                                line_parts = line_content.split('\t')
                                if len(line_parts) >= 2:
                                    db_name = line_parts[0].strip()
                                    table_name = line_parts[1].strip()
                                    key = f"{db_name}.{table_name}"
                                    if key not in seen_pk:
                                        seen_pk.add(key)
                                        tables.append(key)
                    elif '\t' in line:
                        parts = line.split('\t')
                        if len(parts) >= 2:
                            db_name = parts[0].strip()
                            table_name = parts[1].strip()
                            key = f"{db_name}.{table_name}"
                            if key not in seen_pk:
                                seen_pk.add(key)
                                tables.append(key)
            
            if tables:
                if len(tables) > 10:
                    issues_summary["critical_issues"].append({
                        "category": "存储",
                        "issue": "无主键表过多",
                        "description": f"发现 {len(tables)} 个无主键表",
                        "suggestion": "立即为无主键表添加主键，建议使用自增主键或业务主键"
                    })
                else:
                    issues_summary["warning_issues"].append({
                        "category": "存储",
                        "issue": "存在无主键表",
                        "description": f"发现 {len(tables)} 个无主键表",
                        "suggestion": "建议为无主键表添加主键，提升查询性能和数据完整性"
                    })
        
        # 表锁等待分析
        table_locks_waited = self._get_value("TABLE_LOCKS_WAITED", 1)
        if table_locks_waited and table_locks_waited > 0:
            if table_locks_waited > 100:
                issues_summary["critical_issues"].append({
                    "category": "存储",
                    "issue": "表锁等待严重",
                    "description": f"表锁等待次数: {table_locks_waited}（超过100次）",
                    "suggestion": "立即优化事务处理，检查长事务，使用行级锁替代表级锁"
                })
            else:
                issues_summary["warning_issues"].append({
                    "category": "存储",
                    "issue": "存在表锁等待",
                    "description": f"表锁等待次数: {table_locks_waited}",
                    "suggestion": "建议优化事务处理，减少锁等待"
                })
        
        # 表碎片详细分析
        self._analyze_table_fragmentation_detailed(issues_summary)
        
        # 索引使用率分析
        self._analyze_index_usage(issues_summary)
    
    def _analyze_table_fragmentation_detailed(self, issues_summary: Dict[str, Any]):
        """分析表碎片详细信息"""
        frag_data = self.inspection_data.get("TABLE_FRAGMENTATION_DETAILED", [])
        if not frag_data or len(frag_data) <= 1:
            return
        
        high_frag_tables = []
        medium_frag_tables = []
        total_frag_space = 0
        
        for i in range(1, len(frag_data)):
            line = frag_data[i].strip()
            if not line:
                continue
            
            parts = line.split('\t')
            if len(parts) < 6:
                continue
            
            try:
                frag_percent = float(parts[5])
                free_size_mb = float(parts[4])
                total_frag_space += free_size_mb
                
                table_name = f"{parts[0]}.{parts[1]}"
                
                if frag_percent >= 30:
                    high_frag_tables.append(f"{table_name} ({frag_percent:.1f}%, {free_size_mb:.1f}MB)")
                elif frag_percent >= 15:
                    medium_frag_tables.append(f"{table_name} ({frag_percent:.1f}%, {free_size_mb:.1f}MB)")
            except (ValueError, IndexError):
                continue
        
        if total_frag_space > 1024:
            issues_summary["critical_issues"].append({
                "category": "存储",
                "issue": "表碎片空间过大",
                "description": f"碎片总空间: {total_frag_space:.1f}MB（超过1GB）",
                "suggestion": "立即对重要碎片化表执行OPTIMIZE TABLE操作，回收磁盘空间"
            })
        elif total_frag_space > 512:
            issues_summary["warning_issues"].append({
                "category": "存储",
                "issue": "表碎片空间较大",
                "description": f"碎片总空间: {total_frag_space:.1f}MB",
                "suggestion": "建议对碎片化表执行OPTIMIZE TABLE操作"
            })
        
        if high_frag_tables:
            if len(high_frag_tables) > 5:
                issues_summary["critical_issues"].append({
                    "category": "存储",
                    "issue": "高碎片率表过多",
                    "description": f"发现 {len(high_frag_tables)} 个碎片率≥30%的表",
                    "suggestion": "立即对高碎片率表执行OPTIMIZE TABLE操作"
                })
            else:
                issues_summary["warning_issues"].append({
                    "category": "存储",
                    "issue": "存在高碎片率表",
                    "description": f"高碎片率表: {', '.join(high_frag_tables[:5])}",
                    "suggestion": "建议对这些表执行OPTIMIZE TABLE操作"
                })
        
        if medium_frag_tables and len(medium_frag_tables) > 10:
            issues_summary["info_issues"].append({
                "category": "存储",
                "issue": "存在中等碎片率表",
                "description": f"发现 {len(medium_frag_tables)} 个碎片率15%-30%的表",
                "suggestion": "定期监控，必要时执行OPTIMIZE TABLE操作"
            })
    
    def _analyze_index_usage(self, issues_summary: Dict[str, Any]):
        """分析索引使用率"""
        index_stats = self.inspection_data.get("INDEX_STATISTICS", [])
        index_details = self.inspection_data.get("INDEX_DETAILS", [])
        
        if not index_stats or len(index_stats) <= 1:
            return
        
        tables_with_many_indexes = []
        duplicate_indexes = []
        
        # 分析索引数量过多的表
        for i in range(1, len(index_stats)):
            line = index_stats[i].strip()
            if not line:
                continue
            
            parts = line.split('\t')
            if len(parts) < 3:
                continue
            
            try:
                index_count = int(parts[2])
                table_name = f"{parts[0]}.{parts[1]}"
                
                if index_count >= 10:
                    tables_with_many_indexes.append(f"{table_name} ({index_count}个索引)")
            except (ValueError, IndexError):
                continue
        
        if tables_with_many_indexes:
            if len(tables_with_many_indexes) > 3:
                issues_summary["warning_issues"].append({
                    "category": "性能",
                    "issue": "多索引表过多",
                    "description": f"发现 {len(tables_with_many_indexes)} 个表拥有≥10个索引",
                    "suggestion": "审核这些表的索引，删除未使用或重复的索引"
                })
            else:
                issues_summary["info_issues"].append({
                    "category": "性能",
                    "issue": "存在多索引表",
                    "description": f"多索引表: {', '.join(tables_with_many_indexes)}",
                    "suggestion": "建议审核这些表的索引，避免索引过多影响写入性能"
                })
        
        # 简单检查重复索引（基于列组合）
        if index_details and len(index_details) > 1:
            table_indexes = {}
            for i in range(1, len(index_details)):
                line = index_details[i].strip()
                if not line:
                    continue
                
                parts = line.split('\t')
                if len(parts) < 4:
                    continue
                
                table_schema = parts[0]
                table_name = parts[1]
                index_name = parts[2]
                columns = parts[3]
                
                key = f"{table_schema}.{table_name}"
                if key not in table_indexes:
                    table_indexes[key] = {}
                
                if columns in table_indexes[key]:
                    duplicate_indexes.append(f"{key}: {table_indexes[key][columns]} 和 {index_name} 有相同列组合")
                else:
                    table_indexes[key][columns] = index_name
        
        if duplicate_indexes:
            if len(duplicate_indexes) > 3:
                issues_summary["warning_issues"].append({
                    "category": "性能",
                    "issue": "存在多个重复索引",
                    "description": f"发现 {len(duplicate_indexes)} 个可能的重复索引",
                    "suggestion": "立即审核并删除重复索引，减少维护开销"
                })
            else:
                issues_summary["info_issues"].append({
                    "category": "性能",
                    "issue": "存在可能的重复索引",
                    "description": f"重复索引: {'; '.join(duplicate_indexes[:3])}",
                    "suggestion": "建议审核并删除重复索引"
                })
    
    def _analyze_table_fragmentation_item(self, item_data: list) -> Dict[str, str]:
        """分析表碎片详细信息（单巡检项分析）"""
        if not item_data or len(item_data) <= 1:
            return {"status": "正常", "suggestion": "未检测到表碎片信息"}
        
        high_frag_count = 0
        medium_frag_count = 0
        total_frag_space = 0
        total_tables = 0
        
        for i in range(1, len(item_data)):
            line = item_data[i].strip()
            if not line:
                continue
            
            # 尝试多种分隔符分割
            parts = line.split('\t')
            if len(parts) < 6:
                parts = line.split()
                if len(parts) < 6:
                    continue
            
            try:
                frag_percent = float(parts[5])
                free_size_mb = float(parts[4])
                total_frag_space += free_size_mb
                total_tables += 1
                
                if frag_percent >= 30:
                    high_frag_count += 1
                elif frag_percent >= 15:
                    medium_frag_count += 1
            except (ValueError, IndexError):
                continue
        
        if total_tables == 0:
            return {"status": "正常", "suggestion": "未检测到有效的表碎片信息"}
        
        if total_frag_space > 1024 or high_frag_count > 5:
            return {
                "status": "异常",
                "suggestion": f"共检查 {total_tables} 张表，发现 {high_frag_count} 个高碎片表，碎片总空间 {total_frag_space:.1f}MB，请立即执行OPTIMIZE TABLE操作"
            }
        elif total_frag_space > 512 or high_frag_count > 0 or medium_frag_count > 10:
            return {
                "status": "注意",
                "suggestion": f"共检查 {total_tables} 张表，发现 {high_frag_count} 个高碎片表、{medium_frag_count} 个中等碎片表，建议执行OPTIMIZE TABLE操作"
            }
        else:
            return {
                "status": "正常",
                "suggestion": f"共检查 {total_tables} 张表，表碎片情况良好"
            }
    
    def _analyze_index_details_item(self, item_data: list) -> Dict[str, str]:
        """分析索引详细信息（单巡检项分析）"""
        if not item_data or len(item_data) <= 1:
            return {"status": "正常", "suggestion": "未检测到索引信息"}
        
        return {
            "status": "正常",
            "suggestion": f"共检测到 {len(item_data) - 1} 个索引，请检查索引统计了解详细情况"
        }
    
    def _analyze_index_statistics_item(self, item_data: list) -> Dict[str, str]:
        """分析索引统计（单巡检项分析）"""
        if not item_data:
            return {"status": "正常", "suggestion": "未检测到索引统计信息"}
        
        # 先尝试直接返回有数据
        if len(item_data) > 0:
            # 检查是否只有表头
            if len(item_data) == 1:
                return {"status": "正常", "suggestion": "未检测到索引统计数据"}
            # 如果有数据，直接返回正常
            return {"status": "正常", "suggestion": f"共检测到 {len(item_data) - 1} 条索引统计记录"}
        
        return {"status": "正常", "suggestion": "未检测到索引统计信息"}
    
    def _analyze_backup_replication_issues(self, issues_summary: Dict[str, Any]):
        """分析备份复制问题"""
        # MGR/PXC集群无需传统主从同步检查
        is_cluster = self._is_mgr_or_pxc_cluster()
        
        # 从库状态检查（MGR/PXC集群跳过）
        if not is_cluster:
            slave_status_section = self.inspection_data.get("SLAVE_STATUS", [])
            if slave_status_section and len(slave_status_section) > 1:
                # 分析从库状态
                slave_analysis = self._analyze_slave_status(slave_status_section)
                if slave_analysis["status"] == "异常":
                    issues_summary["critical_issues"].append({
                        "category": "备份复制",
                        "issue": "从库复制中断",
                        "description": slave_analysis["suggestion"],
                        "suggestion": "立即修复从库复制问题，确保数据一致性"
                    })
                elif slave_analysis["status"] == "注意":
                    issues_summary["warning_issues"].append({
                        "category": "备份复制",
                        "issue": "从库复制存在问题",
                        "description": slave_analysis["suggestion"],
                        "suggestion": "建议检查从库状态，优化复制配置"
                    })
        
        # GTID模式检查
        gtid_mode_section = self.inspection_data.get("GTID_MODE", [])
        if gtid_mode_section:
            gtid_analysis = self._analyze_gtid_mode(gtid_mode_section)
            if gtid_analysis["status"] != "正常":
                issues_summary["warning_issues"].append({
                    "category": "备份复制",
                    "issue": "GTID模式配置问题",
                    "description": gtid_analysis["suggestion"],
                    "suggestion": "建议检查GTID模式配置，优化复制设置"
                })
        
        # GTID一致性检查
        gtid_consistency_section = self.inspection_data.get("GTID_CONSISTENCY", [])
        if gtid_consistency_section:
            gtid_consistency_analysis = self._analyze_gtid_consistency(gtid_consistency_section)
            if gtid_consistency_analysis["status"] != "正常":
                issues_summary["warning_issues"].append({
                    "category": "备份复制",
                    "issue": "GTID一致性配置问题",
                    "description": gtid_consistency_analysis["suggestion"],
                    "suggestion": "建议检查GTID一致性配置，确保数据一致性"
                })
        
        # 二进制日志检查
        binary_logs_section = self.inspection_data.get("BINARY_LOGS", [])
        if binary_logs_section:
            binary_logs_analysis = self._analyze_binary_logs(binary_logs_section)
            if binary_logs_analysis["status"] == "异常":
                issues_summary["critical_issues"].append({
                    "category": "备份复制",
                    "issue": "二进制日志占用过大",
                    "description": binary_logs_analysis["suggestion"],
                    "suggestion": "立即清理旧的二进制日志，配置 expire_logs_days 参数自动清理"
                })
            elif binary_logs_analysis["status"] == "注意":
                issues_summary["warning_issues"].append({
                    "category": "备份复制",
                    "issue": "二进制日志占用较大",
                    "description": binary_logs_analysis["suggestion"],
                    "suggestion": "建议定期清理二进制日志，配置 expire_logs_days 参数自动清理"
                })
    
    def _analyze_error_logs(self, issues_summary: Dict[str, Any]):
        """分析错误日志问题"""
        error_log_section = self.inspection_data.get("MYSQL_ERROR_LOG", [])
        
        if not error_log_section:
            return
        
        # 收集Warning和ERROR级别的日志
        warning_logs = []
        error_logs = []
        
        for line in error_log_section:
            line_upper = line.upper()
            
            if "ERROR" in line_upper or "WARN" in line_upper:
                # 检查是否是ERROR级别
                if "ERROR" in line_upper:
                    error_logs.append(line)
                else:
                    warning_logs.append(line)
        
        # 分析常见错误模式
        suggestions = []
        
        # 连接错误
        connection_errors = []
        memory_errors = []
        disk_errors = []
        replication_errors = []
        deadlock_errors = []
        
        for log in error_logs + warning_logs:
            log_lower = log.lower()
            
            # 连接错误
            if any(keyword in log_lower for keyword in ['connection', 'connect', 'timeout', 'refused', 'denied']):
                if log not in connection_errors:
                    connection_errors.append(log)
            # 内存错误
            elif any(keyword in log_lower for keyword in ['memory', 'oom', 'out of memory', 'allocation']):
                if log not in memory_errors:
                    memory_errors.append(log)
            # 磁盘错误
            elif any(keyword in log_lower for keyword in ['disk', 'space', 'full', 'i/o', 'io error']):
                if log not in disk_errors:
                    disk_errors.append(log)
            # 复制错误
            elif any(keyword in log_lower for keyword in ['replica', 'slave', 'replication']):
                if log not in replication_errors:
                    replication_errors.append(log)
            # 死锁错误
            elif any(keyword in log_lower for keyword in ['deadlock', 'lock wait', 'lock timeout']):
                if log not in deadlock_errors:
                    deadlock_errors.append(log)
        
        # 生成建议
        if connection_errors:
            suggestions.append(f"连接错误（{len(connection_errors)}条）：检查网络连接、防火墙配置、MySQL用户权限配置")
        
        if memory_errors:
            suggestions.append(f"内存错误（{len(memory_errors)}条）：检查innodb_buffer_pool_size等内存参数，考虑增加物理内存")
        
        if disk_errors:
            suggestions.append(f"磁盘错误（{len(disk_errors)}条）：检查磁盘空间、磁盘I/O性能，考虑数据归档或升级存储")
        
        if replication_errors:
            suggestions.append(f"复制错误（{len(replication_errors)}条）：检查主从复制状态，修复复制中断问题")
        
        if deadlock_errors:
            suggestions.append(f"死锁错误（{len(deadlock_errors)}条）：优化SQL语句，检查事务设计，调整锁超时参数")
        
        # 将错误日志汇总保存到 inspection_data 中供报告生成使用
        self.inspection_data["MYSQL_ERROR_LOG_SUMMARY"] = {
            "warning_count": len(warning_logs),
            "error_count": len(error_logs),
            "warning_logs": warning_logs[:50],  # 最多保存50条
            "error_logs": error_logs[:50],
            "suggestions": suggestions
        }
        
        # 如果有错误，添加到问题摘要
        if error_logs:
            issues_summary["warning_issues"].append({
                "category": "错误日志",
                "issue": f"检测到MySQL错误日志",
                "description": f"发现 {len(error_logs)} 条ERROR级别日志，{len(warning_logs)} 条WARNING级别日志",
                "suggestion": "请查看报告中'错误日志分析'章节了解详情"
            })
        
        # 如果只有警告没有错误
        elif warning_logs:
            issues_summary["info_issues"].append({
                "category": "错误日志",
                "issue": f"检测到MySQL警告日志",
                "description": f"发现 {len(warning_logs)} 条WARNING级别日志",
                "suggestion": "建议关注警告信息，及时排查潜在问题"
            })
    
    def _analyze_system_issues(self, issues_summary: Dict[str, Any]):
        """分析系统资源问题"""
        # CPU使用率检查
        cpu_usage_section = self.inspection_data.get("SYSTEM_CPU_USAGE", [])
        if cpu_usage_section:
            cpu_usage = self._extract_percentage(cpu_usage_section[0])
            if cpu_usage > 95:
                issues_summary["critical_issues"].append({
                    "category": "系统资源",
                    "issue": "CPU使用率严重过高",
                    "description": f"CPU使用率: {cpu_usage}%（超过95%）",
                    "suggestion": "立即检查高负载SQL，优化查询性能，考虑升级CPU或负载均衡"
                })
            elif cpu_usage > 80:
                issues_summary["warning_issues"].append({
                    "category": "系统资源",
                    "issue": "CPU使用率过高",
                    "description": f"CPU使用率: {cpu_usage}%（超过80%）",
                    "suggestion": "检查是否有高负载SQL，优化查询性能，监控CPU使用趋势"
                })
        
        # 内存使用率检查
        memory_usage_section = self.inspection_data.get("SYSTEM_MEMORY_USAGE", [])
        if memory_usage_section:
            memory_usage = self._extract_percentage(memory_usage_section[0])
            if memory_usage > 95:
                issues_summary["critical_issues"].append({
                    "category": "系统资源",
                    "issue": "内存使用率严重过高",
                    "description": f"内存使用率: {memory_usage}%（超过95%）",
                    "suggestion": "立即检查内存泄漏，优化内存使用，考虑增加物理内存"
                })
            elif memory_usage > 90:
                issues_summary["warning_issues"].append({
                    "category": "系统资源",
                    "issue": "内存使用率过高",
                    "description": f"内存使用率: {memory_usage}%（超过90%）",
                    "suggestion": "检查内存使用情况，优化内存配置，监控内存使用趋势"
                })
        
        # 磁盘使用率检查
        disk_usage_section = self.inspection_data.get("SYSTEM_DISK_USAGE", [])
        if disk_usage_section:
            disk_usage = self._extract_percentage(disk_usage_section[0])
            if disk_usage > 95:
                issues_summary["critical_issues"].append({
                    "category": "系统资源",
                    "issue": "磁盘使用率严重过高",
                    "description": f"磁盘使用率: {disk_usage}%（超过95%）",
                    "suggestion": "立即清理磁盘空间，检查大文件，考虑扩容或数据迁移"
                })
            elif disk_usage > 90:
                issues_summary["warning_issues"].append({
                    "category": "系统资源",
                    "issue": "磁盘使用率过高",
                    "description": f"磁盘使用率: {disk_usage}%（超过90%）",
                    "suggestion": "建议清理磁盘空间，监控磁盘使用趋势，适时扩容"
                })
        
        # Swap使用率检查
        swap_usage_section = self.inspection_data.get("SYSTEM_SWAP_USAGE", [])
        if swap_usage_section:
            swap_usage = self._extract_percentage(swap_usage_section[0])
            if swap_usage > 50:
                issues_summary["critical_issues"].append({
                    "category": "系统资源",
                    "issue": "Swap使用率过高",
                    "description": f"Swap使用率: {swap_usage}%（超过50%）",
                    "suggestion": "立即检查内存不足问题，优化内存使用，增加物理内存"
                })
            elif swap_usage > 20:
                issues_summary["warning_issues"].append({
                    "category": "系统资源",
                    "issue": "Swap使用率较高",
                    "description": f"Swap使用率: {swap_usage}%（超过20%）",
                    "suggestion": "检查内存使用情况，优化内存配置，监控Swap使用"
                })
        
        # 磁盘使用率检查
        disk_usage_section = self.inspection_data.get("SYSTEM_DISK_USAGE", [])
        if disk_usage_section:
            disk_usage = self._extract_percentage(disk_usage_section[0])
            if disk_usage > 85:
                issues_summary["critical_issues"].append({
                    "category": "系统资源",
                    "issue": "磁盘空间不足",
                    "description": f"磁盘使用率: {disk_usage}%",
                    "suggestion": "立即清理磁盘空间或扩容"
                })
    
    def _get_value(self, section_name: str, value_index: int = 1):
        """从指定章节获取数值"""
        section = self.inspection_data.get(section_name, [])
        if len(section) > value_index:
            try:
                # 尝试提取数字
                line = section[value_index]
                numbers = re.findall(r'\d+\.?\d*', line)
                if numbers:
                    return float(numbers[0])
            except (ValueError, IndexError):
                pass
        return None
    
    def _clean_xml_text(self, text: str) -> str:
        """清理XML不兼容字符 - 极速优化版本"""
        if not text:
            return ""
        
        # 直接构建结果字符串，避免两次遍历
        cleaned_chars = []
        for c in text:
            if 32 <= ord(c) <= 65535 or c in '\t\n\r':
                cleaned_chars.append(c)
        
        return ''.join(cleaned_chars)
    
    def _extract_percentage(self, text: str) -> float:
        """从文本中提取百分比数值"""
        match = re.search(r'(\d+\.?\d*)%', text)
        if match:
            return float(match.group(1))
        return 0.0
    
    def _extract_item_value(self, line: str) -> str:
        """
        从单条巡检数据中提取值 - 通用方法
        支持格式：
            - name\tvalue (制表符分隔)
            - name: value (冒号分隔)
            - value (直接值)
        """
        if not line:
            return ""
        
        line = line.rstrip('\n')
        
        # 先尝试制表符分隔
        if "\t" in line:
            parts = line.split("\t")
            if len(parts) >= 2:
                val = parts[1].strip()
                if val:
                    return val
        
        # 再尝试冒号分隔
        if ":" in line:
            parts = line.split(":", 1)
            if len(parts) >= 2:
                val = parts[1].strip()
                if val:
                    return val
        
        # 都不行，直接返回整行
        return line.strip()
    
    def generate_report(self, output_format: str = "docx") -> str:
        """生成巡检报告"""
        # 分析日志
        inspection_data = self.parse_log_file()
        issues_summary = self.analyze_issues()
        
        if output_format.lower() == "docx" and DOCX_AVAILABLE:
            return self._generate_docx_report(inspection_data, issues_summary)
        else:
            return self._generate_text_report(inspection_data, issues_summary)
    
    def _generate_table_of_contents(self, doc: Document):
        """在封面后插入Word原生目录域，打开后在Word中更新域即可生成带页码、可点击的目录"""
        # 添加"目录"标题（使用普通段落+样式，避免作为Heading出现在目录自身）
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(self._clean_xml_text("目　　录"))
        title_run.font.name = FONT_NAME
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 51, 102)
        title_para.paragraph_format.space_after = Pt(12)

        # 插入TOC域代码：TOC \o "1-3" \h \z \u
        # \o "1-3" — 捕获大纲级别1-3的标题
        # \h       — 目录条目自动变为超链接（可点击跳转）
        # \z       — Web版式视图中隐藏页码和制表符前导符
        # \u       — 使用段落的大纲级别（Heading样式自动带大纲级别）
        toc_para = doc.add_paragraph()
        toc_run = toc_para.add_run()

        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        toc_run._element.append(fld_begin)

        toc_run2 = toc_para.add_run()
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = ' TOC \\o "1-3" \\h \\z \\u '
        toc_run2._element.append(instr)

        toc_run3 = toc_para.add_run()
        fld_sep = OxmlElement('w:fldChar')
        fld_sep.set(qn('w:fldCharType'), 'separate')
        toc_run3._element.append(fld_sep)

        toc_run4 = toc_para.add_run()
        toc_run4.text = '（请在Word中右键点击此处，选择"更新域" → "更新整个目录" 以生成完整目录）'
        toc_run4.font.color.rgb = RGBColor(160, 160, 160)
        toc_run4.font.size = Pt(11)
        toc_run4.font.name = FONT_NAME
        toc_run4._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

        toc_run5 = toc_para.add_run()
        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        toc_run5._element.append(fld_end)

        doc.add_page_break()

    def _enable_auto_update_fields(self, doc: Document):
        """设置文档打开时自动更新所有域（包括目录）"""
        settings = doc.settings._element
        update_fields = settings.find(qn('w:updateFields'))
        if update_fields is None:
            update_fields = OxmlElement('w:updateFields')
            settings.append(update_fields)
        update_fields.set(qn('w:val'), 'true')

    def _generate_docx_report(self, inspection_data: Dict[str, Any], issues_summary: Dict[str, Any]) -> str:
        """生成Word格式报告"""
        try:
            doc = Document()
            
            self._enable_auto_update_fields(doc)
            
            # 设置文档样式
            self._setup_document_styles(doc)
            
            # 生成报告内容
            # 使用脚本所在目录的LOGO.jpg
            import os
            logo_path = os.path.join(os.path.dirname(__file__), "LOGO.jpg")
            self._generate_cover_page(doc, inspection_data, logo_image_path=logo_path)
            self._generate_table_of_contents(doc)
            self._generate_executive_summary(doc, issues_summary)
            self._generate_inspection_details(doc, inspection_data, issues_summary)  # 新增：巡检项详情
            self._generate_error_log_analysis(doc, inspection_data)  # 新增：错误日志分析
            self._generate_issue_details(doc, issues_summary)
            self._generate_recommendations(doc, issues_summary)
            
            # 保存报告
            report_filename = self._save_docx_report(doc, inspection_data)
            
            self.logger.info(f"Word报告生成完成: {report_filename}")
            return report_filename
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            self.logger.error(f"生成报告失败: {e}\n详细错误信息:\n{error_details}")
            raise
    
    def _generate_text_report(self, inspection_data: Dict[str, Any], issues_summary: Dict[str, Any]) -> str:
        """生成文本格式报告"""
        report_content = []
        report_content.append("=" * 60)
        report_content.append("MySQL数据库巡检报告")
        report_content.append("=" * 60)
        
        # 基本信息
        summary_section = inspection_data.get("INSPECTION_SUMMARY", [])
        for line in summary_section:
            report_content.append(line)
        
        report_content.append("")
        
        # 问题摘要
        summary = issues_summary.get('summary', {})
        report_content.append("执行摘要:")
        report_content.append(f"总问题数: {summary.get('total_issues', 0)}")
        report_content.append(f"严重问题: {summary.get('critical_count', 0)}")
        report_content.append(f"警告问题: {summary.get('warning_count', 0)}")
        report_content.append(f"信息提示: {summary.get('info_count', 0)}")
        report_content.append("")
        
        # 问题详情
        for severity_type in ['critical_issues', 'warning_issues', 'info_issues']:
            issues = issues_summary.get(severity_type, [])
            if issues:
                severity_name = {
                    'critical_issues': '严重问题',
                    'warning_issues': '警告问题',
                    'info_issues': '信息提示'
                }[severity_type]
                
                report_content.append(f"{severity_name}:")
                for i, issue in enumerate(issues, 1):
                    report_content.append(f"  {i}. 【{issue['category']}】{issue['issue']}")
                    report_content.append(f"     描述: {issue['description']}")
                    report_content.append(f"     建议: {issue['suggestion']}")
                    report_content.append("")
        
        # 从日志文件名中提取IP和端口
        log_filename = os.path.basename(self.log_file)
        instance_ip = ""
        instance_port = ""
        
        # 尝试从日志文件名中提取IP和端口
        import re
        ip_pattern = r'(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})'
        port_pattern = r'_(\d{1,5})_'
        
        ip_match = re.search(ip_pattern, log_filename)
        if ip_match:
            instance_ip = ip_match.group(1)
        
        port_match = re.search(port_pattern, log_filename)
        if port_match:
            instance_port = port_match.group(1)
        
        # 如果从日志文件名中提取失败，再尝试从巡检数据中提取
        if not instance_ip or not instance_port:
            summary_data = inspection_data.get('INSPECTION_SUMMARY', [])
            for line in summary_data:
                if not instance_ip and "实例IP:" in line:
                    instance_ip = line.replace("实例IP:", "").strip()
                elif not instance_port and "实例端口:" in line:
                    instance_port = line.replace("实例端口:", "").strip()
        
        # 生成文件名
        current_date = datetime.now().strftime('%Y%m%d')
        
        if instance_ip and instance_port:
            filename = f"reports/MySQL_inspection_{instance_ip}_{instance_port}_{current_date}.txt"
        else:
            # 如果没有IP和端口，使用默认格式
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"reports/MySQL巡检报告_{timestamp}.txt"
        
        os.makedirs("reports", exist_ok=True)
        with open(filename, 'w', encoding='utf-8') as f:
            f.write('\n'.join(report_content))
        
        return filename
    
    def _setup_document_styles(self, doc: Document):
        """设置专业商务文档样式"""
        from datetime import datetime
        
        # 设置页面边距和页眉页脚
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)
            
            # 添加页眉
            header = section.header
            header_para = header.paragraphs[0]
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            header_run = header_para.add_run(self._clean_xml_text("MySQL数据库巡检报告"))
            header_run.font.size = Pt(10)
            header_run.font.name = FONT_NAME
            header_run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            header_run.font.color.rgb = RGBColor(102, 102, 102)
            
            # 添加页脚
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 左侧：公司信息
            footer_run1 = footer_para.add_run(self._clean_xml_text("专业数据库服务提供商"))
            footer_run1.font.size = Pt(9)
            footer_run1.font.name = FONT_NAME
            footer_run1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            footer_run1.font.color.rgb = RGBColor(102, 102, 102)
            
            # 中间：页码
            footer_para.add_run(" | ")
            footer_run2 = footer_para.add_run("第 ")
            footer_run2.font.size = Pt(9)
            footer_run2.font.color.rgb = RGBColor(102, 102, 102)
            
            footer_run3 = footer_para.add_run()
            footer_run3._element.append(parse_xml('<w:fldSimple {} w:instr="PAGE \\* MERGEFORMAT"/>'.format(nsdecls('w'))))
            footer_run3.font.size = Pt(9)
            footer_run3.font.color.rgb = RGBColor(102, 102, 102)
            
            footer_run4 = footer_para.add_run(" 页")
            footer_run4.font.size = Pt(9)
            footer_run4.font.color.rgb = RGBColor(102, 102, 102)
            
            # 右侧：生成时间
            footer_para.add_run(" | ")
            footer_run5 = footer_para.add_run(datetime.now().strftime('%Y-%m-%d'))
            footer_run5.font.size = Pt(9)
            footer_run5.font.color.rgb = RGBColor(102, 102, 102)
        
        # 设置专业字体
        doc.styles['Normal'].font.name = FONT_NAME
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        doc.styles['Normal'].font.size = Pt(11)
        doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        
        # 创建专业标题样式
        title_style = doc.styles.add_style('ProfessionalTitle', 1)
        title_style.font.name = FONT_NAME
        title_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        title_style.font.size = Pt(28)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(24)
        
        # 创建章节标题样式
        heading_style = doc.styles.add_style('SectionHeading', 1)
        heading_style.font.name = FONT_NAME
        heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        heading_style.font.size = Pt(18)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(12)
        
        # 创建子标题样式
        subheading_style = doc.styles.add_style('SubHeading', 1)
        subheading_style.font.name = FONT_NAME
        subheading_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        subheading_style.font.size = Pt(14)
        subheading_style.font.bold = True
        subheading_style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        subheading_style.paragraph_format.space_before = Pt(12)
        subheading_style.paragraph_format.space_after = Pt(6)
        
        # 设置 Word 内置标题样式 - Heading 1 (一级标题)
        heading1_style = doc.styles['Heading 1']
        heading1_style.font.name = FONT_NAME
        heading1_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        heading1_style.font.size = Pt(18)
        heading1_style.font.bold = True
        heading1_style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        heading1_style.paragraph_format.space_before = Pt(18)
        heading1_style.paragraph_format.space_after = Pt(12)
        
        # 设置 Word 内置标题样式 - Heading 2 (二级标题)
        heading2_style = doc.styles['Heading 2']
        heading2_style.font.name = FONT_NAME
        heading2_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        heading2_style.font.size = Pt(16)
        heading2_style.font.bold = True
        heading2_style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        heading2_style.paragraph_format.space_before = Pt(12)
        heading2_style.paragraph_format.space_after = Pt(8)
        
        # 设置 Word 内置标题样式 - Heading 3 (三级标题)
        heading3_style = doc.styles['Heading 3']
        heading3_style.font.name = FONT_NAME
        heading3_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        heading3_style.font.size = Pt(14)
        heading3_style.font.bold = True
        heading3_style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        heading3_style.paragraph_format.space_before = Pt(8)
        heading3_style.paragraph_format.space_after = Pt(4)
    
    def _setup_cell_style(self, cell, bold=False, font_size=10.5, color=None):
        """统一设置单元格样式 - 优化版本，减少重复操作"""
        if cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.font.name = FONT_NAME
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            run.font.size = Pt(font_size)
            if bold:
                run.font.bold = True
            if color is not None:
                run.font.color.rgb = color
    
    def _set_cell_shading(self, cell, fill_color):
        """设置单元格背景色 - 优化版本，每次重新解析"""
        # 简单直接，避免复杂的复制逻辑
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color))
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def _add_heading_with_style(self, doc: Document, text: str, level: int):
        """添加带有正确样式的标题 - 优化版本，减少不必要的操作"""
        heading = doc.add_heading(text, level=level)
        
        # 快速路径：如果已经有运行对象，只设置第一个run的样式即可
        if heading.runs:
            run = heading.runs[0]
            run.font.name = FONT_NAME
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            
            # 根据标题级别设置字体大小
            if level == 1:
                run.font.size = Pt(18)
            elif level == 2:
                run.font.size = Pt(16)
            elif level == 3:
                run.font.size = Pt(14)
            
            run.font.bold = True
        
        return heading
    
    def _generate_cover_page(self, doc: Document, inspection_data: Dict[str, Any], logo_image_path: str = None):
        """生成专业商务封面页"""
        from datetime import datetime
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        # 添加公司logo
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if logo_image_path and os.path.exists(logo_image_path):
            try:
                from docx.shared import Inches
                logo_run = logo_para.add_run()
                logo_run.add_picture(logo_image_path, width=Inches(1.5))
            except Exception as e:
                print(f"添加logo失败: {e}")
                logo_run = logo_para.add_run("[LOGO]")
                logo_run.font.size = Pt(24)
                logo_run.font.color.rgb = RGBColor(0, 51, 102)
        else:
            logo_run = logo_para.add_run("[LOGO]")
            logo_run.font.size = Pt(24)
            logo_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # 主标题
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.style = 'ProfessionalTitle'
        title_run = title_para.add_run(self._clean_xml_text("MySQL数据库巡检报告"))
        
        # 副标题
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run(self._clean_xml_text("专业数据库健康检查与分析"))
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.name = FONT_NAME
        subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
        subtitle_run.italic = True
        
        doc.add_paragraph().add_run("\n" * 6)  # 增加空行
        
        # 专业信息表格
        table = doc.add_table(rows=8, cols=2)
        table.style = 'Light Grid Accent 1'
        table.autofit = False
        
        # 设置表格样式
        table.allow_autofit = False
        for row in table.rows:
            row.cells[0].width = Inches(1.8)
            row.cells[1].width = Inches(4.2)
            
            # 设置单元格背景色
            for cell in row.cells:
                cell.paragraphs[0].style = doc.styles['Normal']
        
        # 表头行
            header_cells = table.rows[0].cells
            header_cells[0].merge(header_cells[1])
            header_cells[0].text = self._clean_xml_text("巡检基本信息")
            header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # 确保有运行对象
            if header_cells[0].paragraphs[0].runs:
                header_cells[0].paragraphs[0].runs[0].font.bold = True
                header_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        
        # 设置表头背景色
        shading_elm = parse_xml(r'<w:shd {} w:fill="003366"/>'.format(nsdecls('w')))
        header_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
        
        # 填充表格数据
        summary_data = inspection_data.get('INSPECTION_SUMMARY', [])
        inspection_time = ""
        mysql_host = ""
        inspection_user = ""
        instance_ip = ""
        instance_port = ""
        
        for line in summary_data:
            if "巡检时间:" in line:
                inspection_time = line.replace("巡检时间:", "").strip()
            elif "实例IP:" in line:
                instance_ip = line.replace("实例IP:", "").strip()
            elif "实例端口:" in line:
                instance_port = line.replace("实例端口:", "").strip()
            elif "MySQL主机:" in line:
                mysql_host = line.replace("MySQL主机:", "").strip()
            elif "巡检用户:" in line:
                inspection_user = line.replace("巡检用户:", "").strip()
        
        # 构建数据库环境显示字符串
        db_env = ""
        if instance_ip and instance_port:
            db_env = f"{instance_ip}:{instance_port}"
        elif mysql_host:
            db_env = mysql_host
        else:
            db_env = "未知"
        
        data = [
            ("报告编号", f"MYSQL-{datetime.now().strftime('%Y%m%d')}"),
            ("巡检时间", inspection_time or datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ("数据库环境", db_env),
            ("巡检范围", "全面数据库健康检查"),
            ("巡检人员", inspection_user or "自动化巡检系统"),
            ("报告版本", "V1.0"),
            ("生成时间", datetime.now().strftime('%Y年%m月%d日 %H:%M:%S'))
        ]
        
        for i, (label, value) in enumerate(data, 1):
            table.cell(i, 0).text = self._clean_xml_text(label)
            table.cell(i, 1).text = self._clean_xml_text(str(value))
            
            # 设置标签单元格样式
            if table.cell(i, 0).paragraphs[0].runs:
                table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
                table.cell(i, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 51, 102)
        
        doc.add_paragraph().add_run("\n" * 4)  # 底部空行
        
        # 添加保密声明
        confidential_para = doc.add_paragraph()
        confidential_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        confidential_run = confidential_para.add_run(self._clean_xml_text("本报告包含敏感信息，请妥善保管"))
        confidential_run.font.size = Pt(10)
        confidential_run.font.color.rgb = RGBColor(153, 0, 0)
        confidential_run.italic = True
        
            
    def _generate_executive_summary(self, doc: Document, issues_summary: Dict[str, Any]):
        """生成执行摘要 - 表格样式"""
                
        self._add_heading_with_style(doc, self._clean_xml_text("一、执行摘要"), level=1)
        
        summary = issues_summary.get('summary', {})
        critical_issues = issues_summary.get('critical_issues', [])
        warning_issues = issues_summary.get('warning_issues', [])
        info_issues = issues_summary.get('info_issues', [])
        
        # 生成健康状态评级
        health_status = "良好"
        if summary.get('critical_count', 0) > 0:
            health_status = "严重"
        elif summary.get('warning_count', 0) > 3:
            health_status = "警告"
        elif summary.get('warning_count', 0) > 0:
            health_status = "注意"
        
        total = summary.get('total_issues', 0)
        critical_count = summary.get('critical_count', 0)
        warning_count = summary.get('warning_count', 0)
        info_count = summary.get('info_count', 0)
        
        overview_text = f"本次MySQL数据库巡检共发现 {total} 个问题。"
        para = doc.add_paragraph(overview_text)
        
        # 问题统计表
        stats_table = doc.add_table(rows=2, cols=4)
        stats_table.style = 'Light Grid Accent 1'
        stats_table.autofit = False
        stats_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        stats_table.columns[0].width = Inches(1.2)
        stats_table.columns[1].width = Inches(1.2)
        stats_table.columns[2].width = Inches(1.2)
        stats_table.columns[3].width = Inches(1.2)
        
        stats_headers = ['健康状态', 'P0 严重问题', 'P1 警告问题', 'P2 信息提示']
        stats_values = [health_status, str(critical_count), str(warning_count), str(info_count)]
        status_colors = {'严重': 'CC0000', '警告': 'CC7700', '注意': 'CC7700', '良好': '006600'}
        health_color = status_colors.get(health_status, '006600')
        
        for j, h in enumerate(stats_headers):
            cell = stats_table.rows[0].cells[j]
            cell.text = self._clean_xml_text(h)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._setup_cell_style(cell, bold=True, font_size=10, color=RGBColor(255, 255, 255))
            self._set_cell_shading(cell, "003366")
        
        for j, v in enumerate(stats_values):
            cell = stats_table.rows[1].cells[j]
            cell.text = self._clean_xml_text(v)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if j == 0:
                self._setup_cell_style(cell, bold=True, color=RGBColor(int(health_color[:2], 16), int(health_color[2:4], 16), int(health_color[4:], 16)))
            else:
                self._setup_cell_style(cell, bold=True, color=RGBColor(0, 0, 0))
            self._set_cell_shading(cell, "E8ECF0")
            cell.paragraphs[0].space_before = Pt(4)
            cell.paragraphs[0].space_after = Pt(4)
        
        doc.add_paragraph()
        
        # 详细展示各优先级问题（按规则：每条问题独立一行，与问题详情表一致）
        severity_config = {
            'critical_issues': ('P0 严重问题', 'CC0000', 'FFF0F0', 'FFD7D7'),
            'warning_issues': ('P1 警告问题', 'CC7700', 'FFF8E6', 'FFEDB8'),
            'info_issues': ('P2 信息提示', '006600', 'F0FFF0', 'D7FFD7')
        }
        
        for severity_type, (level_title, font_color, header_bg, row_bg) in severity_config.items():
            issues = issues_summary.get(severity_type, [])
            if not issues:
                continue
            
            level_heading = doc.add_paragraph(self._clean_xml_text(level_title))
            level_run = level_heading.add_run()
            level_run.font.bold = True
            level_run.font.size = Pt(12)
            level_run.font.color.rgb = RGBColor(int(font_color[:2], 16), int(font_color[2:4], 16), int(font_color[4:], 16))
            
            cat_table = doc.add_table(rows=1, cols=3)
            cat_table.style = 'Light Grid Accent 1'
            cat_table.autofit = False
            cat_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            cat_table.columns[0].width = Inches(0.8)
            cat_table.columns[1].width = Inches(2.4)
            cat_table.columns[2].width = Inches(3.0)
            
            cat_header_cells = cat_table.rows[0].cells
            for j, (cell, text) in enumerate(zip(cat_header_cells, ['分类', '问题描述', '处理建议'])):
                cell.text = self._clean_xml_text(text)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._setup_cell_style(cell, bold=True, font_size=10, color=RGBColor(255, 255, 255))
                self._set_cell_shading(cell, "003366")
            
            for i, issue in enumerate(issues, 1):
                row_cells = cat_table.add_row().cells
                row_cells[0].text = self._clean_xml_text(issue.get('category', ''))
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._setup_cell_style(row_cells[0], bold=True, color=RGBColor(0, 0, 0))
                
                row_cells[1].text = self._clean_xml_text(issue['issue'])
                self._setup_cell_style(row_cells[1], color=RGBColor(51, 51, 51))
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                row_cells[2].text = self._clean_xml_text(issue.get('suggestion', ''))
                self._setup_cell_style(row_cells[2], color=RGBColor(102, 102, 102))
                
                shading_color = row_bg if i % 2 == 0 else header_bg
                for cell in row_cells:
                    self._set_cell_shading(cell, shading_color)
                    cell.paragraphs[0].space_before = Pt(2)
                    cell.paragraphs[0].space_after = Pt(2)
            
            doc.add_paragraph()
        
        # 移除总体处理优先级部分
    
    def _generate_error_log_analysis(self, doc: Document, inspection_data: Dict[str, Any]):
        """生成错误日志分析章节"""
        error_log_summary = inspection_data.get("MYSQL_ERROR_LOG_SUMMARY", {})
        
        doc.add_page_break()
        self._add_heading_with_style(doc, self._clean_xml_text("三、错误日志分析"), level=1)
        
        # 如果没有错误日志数据
        if not error_log_summary:
            para = doc.add_paragraph("未检测到错误日志数据")
            para.style = 'Normal'
            return
        
        # 1. 错误日志汇总
        self._add_heading_with_style(doc, "3.1 错误日志汇总", level=2)
        
        warning_count = error_log_summary.get("warning_count", 0)
        error_count = error_log_summary.get("error_count", 0)
        
        para = doc.add_paragraph(f"ERROR级别日志: {error_count} 条")
        para = doc.add_paragraph(f"WARNING级别日志: {warning_count} 条")
        
        # 2. 优化建议
        suggestions = error_log_summary.get("suggestions", [])
        if suggestions:
            self._add_heading_with_style(doc, "3.2 优化建议", level=2)
            
            for i, suggestion in enumerate(suggestions, 1):
                para = doc.add_paragraph(f"{i}. {suggestion}")
        
        # 3. ERROR级别日志详情（最多显示50条）
        error_logs = error_log_summary.get("error_logs", [])
        if error_logs:
            self._add_heading_with_style(doc, "3.3 ERROR级别日志详情（前50条）", level=2)
            
            for log in error_logs[:50]:
                para = doc.add_paragraph(self._clean_xml_text(log))
                para.paragraph_format.space_after = Pt(6)
        
        # 4. WARNING级别日志详情（最多显示50条）
        warning_logs = error_log_summary.get("warning_logs", [])
        if warning_logs:
            self._add_heading_with_style(doc, "3.4 WARNING级别日志详情（前50条）", level=2)
            
            for log in warning_logs[:50]:
                para = doc.add_paragraph(self._clean_xml_text(log))
                para.paragraph_format.space_after = Pt(6)
    
    def _generate_issue_details(self, doc: Document, issues_summary: Dict[str, Any]):
        """生成问题详情 - 表格样式"""
        if issues_summary.get('summary', {}).get('total_issues', 0) == 0:
            return
        
        doc.add_page_break()
        self._add_heading_with_style(doc, self._clean_xml_text("四、问题详情"), level=1)
        
        severity_config = {
            'critical_issues': ('4.1 严重问题', 'CC0000', 'FFF0F0', 'FFD7D7'),
            'warning_issues': ('4.2 警告问题', 'CC7700', 'FFF8E6', 'FFEDB8'),
            'info_issues': ('4.3 信息提示', '006600', 'F0FFF0', 'D7FFD7')
        }
        
        for severity_type, (severity_title, font_color, header_bg, row_bg) in severity_config.items():
            issues = issues_summary.get(severity_type, [])
            if not issues:
                continue
            
            self._add_heading_with_style(doc, severity_title, level=2)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            table.autofit = False
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            table.columns[0].width = Inches(0.4)
            table.columns[1].width = Inches(1.0)
            table.columns[2].width = Inches(1.6)
            table.columns[3].width = Inches(3.2)
            
            header_cells = table.rows[0].cells
            header_texts = ['序号', '分类', '问题描述', '优化建议']
            for j, (cell, text) in enumerate(zip(header_cells, header_texts)):
                cell.text = self._clean_xml_text(text)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._setup_cell_style(cell, bold=True, font_size=10, color=RGBColor(255, 255, 255))
                self._set_cell_shading(cell, "003366")
            
            for i, issue in enumerate(issues, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._setup_cell_style(row_cells[0], color=RGBColor(0, 0, 0))
                
                row_cells[1].text = self._clean_xml_text(issue.get('category', ''))
                self._setup_cell_style(row_cells[1], bold=True, color=RGBColor(0, 0, 0))
                
                desc_text = issue.get('issue', '')
                if issue.get('description'):
                    desc_text += '\n' + issue['description']
                row_cells[2].text = self._clean_xml_text(desc_text)
                self._setup_cell_style(row_cells[2], color=RGBColor(51, 51, 51))
                
                row_cells[3].text = self._clean_xml_text(issue.get('suggestion', ''))
                self._setup_cell_style(row_cells[3], color=RGBColor(102, 102, 102))
                
                shading_color = row_bg if i % 2 == 0 else header_bg
                for cell in row_cells:
                    self._set_cell_shading(cell, shading_color)
                
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for cell in row_cells:
                    cell.paragraphs[0].space_before = Pt(2)
                    cell.paragraphs[0].space_after = Pt(2)
            
            doc.add_paragraph()
    
    def _generate_inspection_details(self, doc: Document, inspection_data: Dict[str, Any], issues_summary: Dict[str, Any]):
        """生成巡检项详情
        接收 issues_summary 参数，用于判断是否需要显示MGR详细表格
        """
        doc.add_page_break()
        self._add_heading_with_style(doc, self._clean_xml_text("二、巡检项详情"), level=1)
        
        # 检查是否有MGR相关问题
        has_mgr_issues = False
        mgr_item_display_name = self._get_item_display_name("MGR_GROUP_MEMBERS")
        for issue in issues_summary.get("critical_issues", []) + issues_summary.get("warning_issues", []):
            if issue["issue"] == mgr_item_display_name:
                has_mgr_issues = True
                break
        
        # 定义巡检项分类和阈值
        inspection_categories = {
            "数据库信息": ["DATABASE_VERSION", "UPTIME", "MAX_CONNECTIONS", "INNODB_BUFFER_POOL_SIZE", "OS_VERSION", 
                         "INNODB_LOG_FILE_SIZE", "INNODB_LOG_FILES_IN_GROUP", "INNODB_FLUSH_LOG_AT_TRX_COMMIT",
                         "SYNC_BINLOG", "BINLOG_FORMAT", "TRANSACTION_ISOLATION", 
                         "INNODB_FLUSH_METHOD", "INNODB_FILE_PER_TABLE", "OPEN_FILES_LIMIT", 
                         "TABLE_OPEN_CACHE", "MAX_ALLOWED_PACKET", "WAIT_TIMEOUT", "INTERACTIVE_TIMEOUT"],
            "性能指标": ["CURRENT_CONNECTIONS", "QUERY_CACHE_HITS", "QUERY_CACHE_INSERTS", "INNODB_READS", 
                        "INNODB_READ_REQUESTS", "SLOW_QUERIES", "TABLE_LOCKS_WAITED", "QUESTIONS", "QUERIES",
                        "COM_SELECT", "COM_INSERT", "COM_UPDATE", "COM_DELETE", "INNODB_ROW_LOCK_WAITS",
                        "INNODB_ROW_LOCK_TIME_AVG", "INNODB_ROW_LOCK_TIME_MAX", "CREATED_TMP_TABLES",
                        "CREATED_TMP_DISK_TABLES", "ABORTED_CLIENTS", "ABORTED_CONNECTS", 
                        "COM_COMMIT", "COM_ROLLBACK", "INNODB_ENGINE_STATUS"],
            "安全配置": ["EMPTY_PASSWORD_USERS", "REMOTE_ROOT_USERS", "SSL_STATUS", "REMOTE_USERS"],
            "存储信息": ["DATABASE_SIZES", "TABLE_FRAGMENTATION_DETAILED", 
                       "TABLES_WITHOUT_PRIMARY_KEY", "PARTITIONED_TABLES"],
            "主从复制": ["MASTER_STATUS", "SLAVE_STATUS", "BINARY_LOGS", "GTID_MODE", "GTID_CONSISTENCY", "GTID_EXECUTED"],
            "系统资源": ["SYSTEM_CPU_USAGE", "SYSTEM_MEMORY_USAGE", "SYSTEM_SWAP_USAGE", "SYSTEM_DISK_USAGE", 
                        "MYSQL_PROCESS", "FILE_DESCRIPTOR_LIMIT", "NETWORK_CONNECTIONS", "MYSQL_ERROR_LOG"]
        }
        
        # 动态添加集群信息分类（根据获取到的数据判断）
        has_mgr_data = (inspection_data.get("MGR_GROUP_MEMBERS", []) and 
                       len(inspection_data.get("MGR_GROUP_MEMBERS", [])) > 1) or \
                      (inspection_data.get("MGR_MEMBER_STATS", []) and 
                       len(inspection_data.get("MGR_MEMBER_STATS", [])) > 1)
        has_pxc_data = (inspection_data.get("PXC_WSREP_STATUS", []) and 
                       len(inspection_data.get("PXC_WSREP_STATUS", [])) > 1) or \
                      (inspection_data.get("PXC_WSREP_VARIABLES", []) and 
                       len(inspection_data.get("PXC_WSREP_VARIABLES", [])) > 1)
        
        if has_mgr_data:
            inspection_categories["集群信息（MGR）"] = ["MGR_GROUP_MEMBERS", "MGR_MEMBER_STATS"]
        elif has_pxc_data:
            inspection_categories["集群信息（PXC）"] = ["PXC_WSREP_STATUS", "PXC_WSREP_VARIABLES"]
        
        # 分析每个巡检项
        for cat_idx, (category, items) in enumerate(inspection_categories.items()):
            category_title = f"2.{cat_idx + 1} {category}"
            self._add_heading_with_style(doc, self._clean_xml_text(category_title), level=2)
            
            # 创建专业表格
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            table.autofit = False
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 表格居中
            
            # 设置表格列宽（优化布局，留出边距）
            table.columns[0].width = Inches(1.6)
            table.columns[1].width = Inches(2.0)
            table.columns[2].width = Inches(1.2)
            table.columns[3].width = Inches(2.4)
            
            # 专业表头
            header_cells = table.rows[0].cells
            header_cells[0].text = self._clean_xml_text("巡检项目")
            header_cells[1].text = self._clean_xml_text("检测结果")
            header_cells[2].text = self._clean_xml_text("健康状态")
            header_cells[3].text = self._clean_xml_text("优化建议")
            
            # 设置专业表头样式
            for cell in header_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._setup_cell_style(cell, bold=True, font_size=11, color=RGBColor(255,255,255))
                self._set_cell_shading(cell, "003366")
            
            # 添加巡检项数据 - 确保所有巡检项都展示
            for item in items:
                # 获取巡检数据
                item_data = inspection_data.get(item, [])
                analysis_result = self._analyze_inspection_item(item, item_data)
                
                # 添加专业数据行
                row_cells = table.add_row().cells
                
                # 巡检项目列
                row_cells[0].text = self._clean_xml_text(self._get_item_display_name(item))
                self._setup_cell_style(row_cells[0], bold=True, color=RGBColor(51,51,51))
                
                # 检测结果列
                row_cells[1].text = self._format_item_value(item, item_data)
                self._setup_cell_style(row_cells[1], color=RGBColor(68,68,68))
                
                # 健康状态列
                status_text = analysis_result["status"]
                if status_text == "正常":
                    row_cells[2].text = "✓ " + status_text
                    status_color = RGBColor(0,128,0)
                elif status_text == "注意":
                    row_cells[2].text = "⚠ " + status_text
                    status_color = RGBColor(255,165,0)
                elif status_text == "异常":
                    row_cells[2].text = "✗ " + status_text
                    status_color = RGBColor(255,0,0)
                else:
                    status_color = RGBColor(0,0,0)
                self._setup_cell_style(row_cells[2], color=status_color)
                row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 优化建议列
                row_cells[3].text = self._clean_xml_text(analysis_result["suggestion"])
                self._setup_cell_style(row_cells[3], color=RGBColor(102,102,102))
                
                # 设置行交替背景色
                if len(table.rows) % 2 == 0:
                    for cell in row_cells:
                        self._set_cell_shading(cell, "F5F5F5")
            
            doc.add_paragraph()  # 添加空行分隔不同分类
            
            # 如果是存储信息分类，在下方添加详细表格
            if category == "存储信息":
                # 1. 为表碎片详细信息单独生成表格
                frag_data = inspection_data.get("TABLE_FRAGMENTATION_DETAILED", [])
                if frag_data:
                    # 添加标题
                    self._add_heading_with_style(doc, self._clean_xml_text("2.4.1 表碎片详细信息"), level=3)
                    
                    # 创建表格，完全模仿主表格样式
                    frag_table = doc.add_table(rows=1, cols=4)
                    frag_table.style = 'Light Grid Accent 1'
                    frag_table.autofit = False
                    frag_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 设置表格列宽，和主表格完全一样！
                    frag_table.columns[0].width = Inches(1.6)
                    frag_table.columns[1].width = Inches(2.0)
                    frag_table.columns[2].width = Inches(1.2)
                    frag_table.columns[3].width = Inches(2.4)
                    
                    # 设置表头
                    frag_header = frag_table.rows[0].cells
                    frag_header[0].text = self._clean_xml_text("数据库名")
                    frag_header[1].text = self._clean_xml_text("表名")
                    frag_header[2].text = self._clean_xml_text("碎片大小(MB)")
                    frag_header[3].text = self._clean_xml_text("碎片率(%)")
                    
                    # 设置表头样式，完全模仿主表格
                    for cell in frag_header:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        self._setup_cell_style(cell, bold=True, font_size=11, color=RGBColor(255,255,255))
                        self._set_cell_shading(cell, "003366")
                    
                    # 填充数据，只显示碎片大小>2GB且碎片率>20%的，并且限制最多显示100行
                    filtered_frag_data = []
                    for line in frag_data:
                        stripped_line = line.strip()
                        if not stripped_line:
                            continue
                        parts = stripped_line.split('\t')
                        if len(parts) < 7:
                            continue
                        
                        # 检查碎片大小和碎片率
                        try:
                            frag_size_mb = float(parts[5])  # 碎片大小
                            frag_percent = float(parts[6])  # 碎片率
                            # 过滤条件：碎片大小 > 2048MB (2GB) 且 碎片率 > 20%
                            if frag_size_mb > 2048.0 and frag_percent > 20.0:
                                filtered_frag_data.append(parts)
                        except (ValueError, IndexError):
                            continue
                    
                    # 按碎片率降序排序，优先显示最严重的
                    filtered_frag_data.sort(key=lambda x: float(x[6]) if len(x) > 6 else 0, reverse=True)
                    
                    # 限制最多显示100行
                    max_frag_rows = 100
                    for idx, parts in enumerate(filtered_frag_data):
                        if idx >= max_frag_rows:
                            break
                        
                        row_cells = frag_table.add_row().cells
                        
                        # 填充数据 - 避免不必要的字符串转换和清理
                        row_cells[0].text = parts[0]
                        row_cells[1].text = parts[1]
                        row_cells[2].text = parts[5]
                        row_cells[3].text = parts[6]
                        
                        # 设置单元格样式
                        for cell in row_cells:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            self._setup_cell_style(cell)
                        
                        # 交替行背景色
                        if len(frag_table.rows) % 2 == 0:
                            for cell in row_cells:
                                self._set_cell_shading(cell, "F5F5F5")
                    
                    # 如果显示的是有限行数，添加提示
                    if len(filtered_frag_data) > max_frag_rows:
                        doc.add_paragraph(f"注：共发现 {len(filtered_frag_data)} 个碎片大小>2GB且碎片率>20%的表，此处仅显示最严重的前 {max_frag_rows} 个。")
                    elif filtered_frag_data:
                        doc.add_paragraph(f"注：共发现 {len(filtered_frag_data)} 个碎片大小>2GB且碎片率>20%的表。")
                    
                    doc.add_paragraph()
                
                # 2. 为索引统计单独生成表格
                idx_data = inspection_data.get("INDEX_STATISTICS", [])
                if idx_data:
                    # 添加标题
                    self._add_heading_with_style(doc, self._clean_xml_text("2.4.2 索引统计"), level=3)
                    
                    # 创建表格，完全模仿主表格样式
                    idx_table = doc.add_table(rows=1, cols=3)
                    idx_table.style = 'Light Grid Accent 1'
                    idx_table.autofit = False
                    idx_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 设置表格列宽，和主表格总宽度一致（7.2英寸）
                    idx_table.columns[0].width = Inches(1.8)
                    idx_table.columns[1].width = Inches(3.6)
                    idx_table.columns[2].width = Inches(1.8)
                    
                    # 设置表头
                    idx_header = idx_table.rows[0].cells
                    idx_header[0].text = self._clean_xml_text("数据库名")
                    idx_header[1].text = self._clean_xml_text("表名")
                    idx_header[2].text = self._clean_xml_text("索引数量")
                    
                    # 设置表头样式，完全模仿主表格
                    for cell in idx_header:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        self._setup_cell_style(cell, bold=True, font_size=11, color=RGBColor(255,255,255))
                        self._set_cell_shading(cell, "003366")
                    
                    # 过滤并排序数据
                    filtered_idx_data = []
                    for line in idx_data:
                        stripped_line = line.strip()
                        if not stripped_line:
                            continue
                        parts = stripped_line.split('\t')
                        if len(parts) < 3:
                            continue
                        
                        # 检查索引数量
                        try:
                            index_count = int(parts[2])
                            if index_count >= 5:
                                filtered_idx_data.append(parts)
                        except (ValueError, IndexError):
                            continue
                    
                    # 按索引数量降序排序，优先显示索引最多的
                    filtered_idx_data.sort(key=lambda x: int(x[2]) if len(x) > 2 else 0, reverse=True)
                    
                    # 限制最多显示100行
                    max_idx_rows = 100
                    for idx, parts in enumerate(filtered_idx_data):
                        if idx >= max_idx_rows:
                            break
                        
                        row_cells = idx_table.add_row().cells
                        
                        # 填充数据 - 避免不必要的字符串转换和清理
                        row_cells[0].text = parts[0]
                        row_cells[1].text = parts[1]
                        row_cells[2].text = parts[2]
                        
                        # 设置单元格样式
                        for cell in row_cells:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            self._setup_cell_style(cell)
                        
                        # 交替行背景色
                        if len(idx_table.rows) % 2 == 0:
                            for cell in row_cells:
                                self._set_cell_shading(cell, "F5F5F5")
                    
                    # 如果显示的是有限行数，添加提示
                    if len(filtered_idx_data) > max_idx_rows:
                        doc.add_paragraph(f"注：共发现 {len(filtered_idx_data)} 个索引数量大于等于5的表，此处仅显示索引最多的前 {max_idx_rows} 个。")
                    
                    doc.add_paragraph()
            
            # 如果是集群信息（MGR）分类，在下方添加详细表格（只有在有问题时才显示）
            elif category == "集群信息（MGR）" and has_mgr_issues:
                # 1. 为MGR集群成员信息单独生成表格
                mgr_members_data = inspection_data.get("MGR_GROUP_MEMBERS", [])
                if mgr_members_data:
                    # 添加标题
                    self._add_heading_with_style(doc, self._clean_xml_text("2.7.1 MGR集群成员信息"), level=3)
                    
                    # 创建表格
                    mgr_members_table = doc.add_table(rows=1, cols=4)
                    mgr_members_table.style = 'Light Grid Accent 1'
                    mgr_members_table.autofit = False
                    mgr_members_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 设置表格列宽
                    mgr_members_table.columns[0].width = Inches(1.5)
                    mgr_members_table.columns[1].width = Inches(2.0)
                    mgr_members_table.columns[2].width = Inches(1.5)
                    mgr_members_table.columns[3].width = Inches(1.5)
                    
                    # 设置表头
                    mgr_members_header = mgr_members_table.rows[0].cells
                    mgr_members_header[0].text = self._clean_xml_text("成员ID")
                    mgr_members_header[1].text = self._clean_xml_text("主机")
                    mgr_members_header[2].text = self._clean_xml_text("状态")
                    mgr_members_header[3].text = self._clean_xml_text("角色")
                    
                    # 设置表头样式
                    for cell in mgr_members_header:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        self._setup_cell_style(cell, bold=True, font_size=11, color=RGBColor(255,255,255))
                        self._set_cell_shading(cell, "003366")
                    
                    # 填充数据
                    for line in mgr_members_data:
                        stripped_line = line.strip()
                        if not stripped_line:
                            continue
                        parts = stripped_line.split('\t')
                        if len(parts) < 8:
                            continue
                        
                        row_cells = mgr_members_table.add_row().cells
                        row_cells[0].text = parts[1]  # 成员ID
                        row_cells[1].text = parts[2]  # 主机
                        row_cells[2].text = parts[4]  # 状态
                        row_cells[3].text = parts[5]  # 角色
                        
                        for i in range(4):
                            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            self._setup_cell_style(row_cells[i])
                        
                        # 交替行背景色
                        if len(mgr_members_table.rows) % 2 == 0:
                            for cell in row_cells:
                                self._set_cell_shading(cell, "F5F5F5")
                    
                    doc.add_paragraph()
                
                # 2. 为MGR成员统计信息单独生成表格
                mgr_stats_data = inspection_data.get("MGR_MEMBER_STATS", [])
                if mgr_stats_data:
                    # 添加标题
                    self._add_heading_with_style(doc, self._clean_xml_text("2.7.2 MGR成员统计信息"), level=3)
                    
                    # 创建表格
                    mgr_stats_table = doc.add_table(rows=1, cols=4)
                    mgr_stats_table.style = 'Light Grid Accent 1'
                    mgr_stats_table.autofit = False
                    mgr_stats_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 设置表格列宽
                    mgr_stats_table.columns[0].width = Inches(1.5)
                    mgr_stats_table.columns[1].width = Inches(1.8)
                    mgr_stats_table.columns[2].width = Inches(1.8)
                    mgr_stats_table.columns[3].width = Inches(1.8)
                    
                    # 设置表头
                    mgr_stats_header = mgr_stats_table.rows[0].cells
                    mgr_stats_header[0].text = self._clean_xml_text("成员ID")
                    mgr_stats_header[1].text = self._clean_xml_text("计数提交")
                    mgr_stats_header[2].text = self._clean_xml_text("计数冲突")
                    mgr_stats_header[3].text = self._clean_xml_text("队列事务")
                    
                    # 设置表头样式
                    for cell in mgr_stats_header:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        self._setup_cell_style(cell, bold=True, font_size=11, color=RGBColor(255,255,255))
                        self._set_cell_shading(cell, "003366")
                    
                    # 填充数据
                    for line in mgr_stats_data:
                        stripped_line = line.strip()
                        if not stripped_line:
                            continue
                        parts = stripped_line.split('\t')
                        if len(parts) < 14:
                            continue
                        
                        row_cells = mgr_stats_table.add_row().cells
                        row_cells[0].text = parts[2]  # 成员ID
                        row_cells[1].text = parts[3]  # 计数提交
                        row_cells[2].text = parts[5]  # 计数冲突
                        row_cells[3].text = parts[10] # 队列事务
                        
                        for i in range(4):
                            row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            self._setup_cell_style(row_cells[i])
                        
                        # 交替行背景色
                        if len(mgr_stats_table.rows) % 2 == 0:
                            for cell in row_cells:
                                self._set_cell_shading(cell, "F5F5F5")
                    
                    doc.add_paragraph()
    
    def _get_item_display_name(self, item: str) -> str:
        """获取巡检项显示名称"""
        name_map = {
            "DATABASE_VERSION": "数据库版本",
            "UPTIME": "运行时间",
            "MAX_CONNECTIONS": "最大连接数",
            "INNODB_BUFFER_POOL_SIZE": "InnoDB缓冲池大小",
            "OS_VERSION": "操作系统版本",
            "CURRENT_CONNECTIONS": "当前连接数",
            "QUERY_CACHE_HITS": "查询缓存命中数",
            "QUERY_CACHE_INSERTS": "查询缓存插入数",
            "INNODB_READS": "InnoDB物理读取",
            "INNODB_READ_REQUESTS": "InnoDB读取请求",
            "SLOW_QUERIES": "慢查询数量",
            "TABLE_LOCKS_WAITED": "表锁等待次数",
            "EMPTY_PASSWORD_USERS": "空密码用户",
            "REMOTE_ROOT_USERS": "远程root用户",
            "SSL_STATUS": "SSL状态",
            "REMOTE_USERS": "远程用户",
            "DATABASE_SIZES": "数据库大小",
            "FRAGMENTED_TABLES": "碎片化表",
            "TABLE_FRAGMENTATION_DETAILED": "表碎片详细信息",
            "MASTER_STATUS": "主库状态",
            "SLAVE_STATUS": "从库状态",
            "BINARY_LOGS": "二进制日志",
            "SYSTEM_CPU_USAGE": "CPU使用率",
            "SYSTEM_MEMORY_USAGE": "内存使用率",
            "SYSTEM_DISK_USAGE": "磁盘使用率",
            "MYSQL_PROCESS": "MySQL进程",
            "SYSTEM_SWAP_USAGE": "Swap使用率",
            "TABLES_WITHOUT_PRIMARY_KEY": "无主键表",
            "PARTITIONED_TABLES": "分区表",
            "GTID_MODE": "GTID模式",
            "GTID_CONSISTENCY": "GTID一致性",
            "GTID_EXECUTED": "GTID执行状态",
            "SLAVE_STATUS_DETAILED": "从库详细状态",
            # 新增巡检项
            "INNODB_LOG_FILE_SIZE": "InnoDB日志文件大小",
            "INNODB_LOG_FILES_IN_GROUP": "InnoDB日志文件数量",
            "INNODB_FLUSH_LOG_AT_TRX_COMMIT": "事务日志刷新策略",
            "SYNC_BINLOG": "Binlog同步策略",
            "BINLOG_FORMAT": "Binlog格式",
            "TRANSACTION_ISOLATION": "事务隔离级别",
            "INNODB_FLUSH_METHOD": "InnoDB刷新方法",
            "INNODB_FILE_PER_TABLE": "独立表空间",
            "OPEN_FILES_LIMIT": "打开文件限制",
            "TABLE_OPEN_CACHE": "表打开缓存",
            "MAX_ALLOWED_PACKET": "最大允许包大小",
            "WAIT_TIMEOUT": "等待超时",
            "INTERACTIVE_TIMEOUT": "交互超时",
            "QUESTIONS": "查询总数",
            "QUERIES": "执行语句总数",
            "COM_SELECT": "SELECT语句数",
            "COM_INSERT": "INSERT语句数",
            "COM_UPDATE": "UPDATE语句数",
            "COM_DELETE": "DELETE语句数",
            "INNODB_ROW_LOCK_WAITS": "InnoDB行锁等待次数",
            "INNODB_ROW_LOCK_TIME_AVG": "InnoDB行锁平均等待时间",
            "INNODB_ROW_LOCK_TIME_MAX": "InnoDB行锁最大等待时间",
            "CREATED_TMP_TABLES": "创建临时表数量",
            "CREATED_TMP_DISK_TABLES": "创建磁盘临时表数量",
            "ABORTED_CLIENTS": "异常断开连接数",
            "ABORTED_CONNECTS": "异常连接尝试数",
            "COM_COMMIT": "COMMIT语句数",
            "COM_ROLLBACK": "ROLLBACK语句数",
            "INNODB_ENGINE_STATUS": "InnoDB引擎状态",
            "FILE_DESCRIPTOR_LIMIT": "文件描述符限制",
            "NETWORK_CONNECTIONS": "网络连接数",
            "MYSQL_ERROR_LOG": "MySQL错误日志",
            # 表碎片详细分析
            "TABLE_FRAGMENTATION_DETAILED": "表碎片详细信息",
            # 索引分析
            "INDEX_DETAILS": "索引详细信息",
            "INDEX_STATISTICS": "索引统计",
            # MGR集群巡检项
            "MGR_GROUP_MEMBERS": "MGR集群成员",
            "MGR_MEMBER_STATS": "MGR成员统计",
            # PXC集群巡检项
            "PXC_WSREP_STATUS": "PXC集群状态",
            "PXC_WSREP_VARIABLES": "PXC集群配置"
        }
        return name_map.get(item, item)
    
    def _format_item_value(self, item: str, item_data: list) -> str:
        """格式化巡检项值 - 使用通用提取方法"""
        # 主库状态和从库状态即使无数据也走专用格式化，以支持集群/主库识别
        if item == "MASTER_STATUS":
            slave_status_data = self.inspection_data.get("SLAVE_STATUS", [])
            return self._format_master_status(item_data, slave_status_data)
        
        if item == "SLAVE_STATUS":
            return self._format_slave_status(item_data)
        
        # 通用处理：无数据时返回提示
        if not item_data or len(item_data) == 0:
            return "未检测到数据"
        
        # 特殊处理：二进制日志格式化显示
        if item == "BINARY_LOGS":
            return self._format_binary_logs(item_data)
        
        # 特殊处理：MGR集群成员格式化显示
        if item == "MGR_GROUP_MEMBERS":
            return self._format_mgr_group_members(item_data)
        
        # 特殊处理：MGR成员统计格式化显示
        if item == "MGR_MEMBER_STATS":
            return self._format_mgr_member_stats(item_data)
        
        # 特殊处理：PXC集群状态格式化显示
        if item == "PXC_WSREP_STATUS":
            return self._format_pxc_wsrep_status(item_data)
        
        # 特殊处理：PXC集群配置格式化显示
        if item == "PXC_WSREP_VARIABLES":
            return self._format_pxc_wsrep_variables(item_data)
        
        # 特殊处理：表碎片详细信息格式化显示（后续单独生成表格）
        if item == "TABLE_FRAGMENTATION_DETAILED":
            return "详见下方表格"
        
        # 特殊处理：索引详细信息格式化显示
        if item == "INDEX_DETAILS":
            return self._format_index_details(item_data)
        
        # 特殊处理：索引统计格式化显示（后续单独生成表格）
        if item == "INDEX_STATISTICS":
            return "详见下方表格"
        
        # 特殊处理：运行时间显示格式
        if item == "UPTIME":
            return self._format_uptime_value(item_data)
        
        # 特殊处理：InnoDB缓冲池大小显示格式（转换为GB）
        if item == "INNODB_BUFFER_POOL_SIZE":
            return self._format_size_to_gb(item_data, "未检测到缓冲池大小")
        
        # 新增：InnoDB日志文件大小显示格式（转换为MB/GB）
        if item == "INNODB_LOG_FILE_SIZE":
            return self._format_size_to_mb_gb(item_data, "未检测到日志文件大小")
        
        # 新增：最大允许包大小显示格式（转换为MB）
        if item == "MAX_ALLOWED_PACKET":
            return self._format_size_to_mb(item_data, "未检测到包大小")
        
        # 新增：InnoDB引擎状态显示格式
        if item == "INNODB_ENGINE_STATUS":
            if item_data:
                has_deadlock = any("LATEST DETECTED DEADLOCK" in line for line in item_data)
                if has_deadlock:
                    return "检测到死锁信息"
                else:
                    return "InnoDB引擎状态正常"
            return "未检测到InnoDB引擎状态"
        
        # 需要完整显示的项目（不包括有专门格式化函数的复制状态）
        full_display_items = ["SLAVE_STATUS_DETAILED", 
                            "MYSQL_PROCESS", 
                            "NETWORK_CONNECTIONS", "MYSQL_ERROR_LOG"]
        
        if item in full_display_items:
            if len(item_data) > 10:
                return f"共 {len(item_data)} 条记录\n" + "\n".join([self._clean_xml_text(line) for line in item_data[:10]])
            else:
                return "\n".join([self._clean_xml_text(line) for line in item_data])
        
        # 特殊处理：空密码用户显示格式
        if item == "EMPTY_PASSWORD_USERS":
            return self._format_empty_password_users(item_data)
        
        # 特殊处理：远程root用户显示格式
        if item == "REMOTE_ROOT_USERS":
            return self._format_remote_root_users(item_data)
        
        # 特殊处理：远程用户显示格式
        if item == "REMOTE_USERS":
            return self._format_remote_users(item_data)
        
        # 特殊处理：无主键表显示格式
        if item == "TABLES_WITHOUT_PRIMARY_KEY":
            return self._format_tables_without_pk(item_data)
        
        # 特殊处理：数据库大小显示格式
        if item == "DATABASE_SIZES":
            return self._format_database_sizes(item_data)
        
        # 特殊处理：文件描述符限制显示格式
        if item == "FILE_DESCRIPTOR_LIMIT":
            return self._format_file_descriptor_limit_display(item_data)
        
        # 特殊处理：InnoDB配置项显示格式
        if item == "INNODB_FLUSH_METHOD":
            val = self._extract_item_value(item_data[0]) if item_data else ""
            return val or "未检测到"
        if item == "INNODB_FILE_PER_TABLE":
            val = self._extract_item_value(item_data[0]) if item_data else ""
            return val or "未检测到"
        
        # 对于多行数据，显示摘要
        if len(item_data) > 3:
            return f"共 {len(item_data)} 条记录"
        
        # 使用通用方法提取值
        cleaned_values = []
        for line in item_data[:3]:
            val = self._extract_item_value(line)
            if val:
                cleaned_values.append(val)
        
        if len(cleaned_values) == 1:
            return cleaned_values[0]
        
        return "\n".join(cleaned_values) if cleaned_values else "未检测到数据"
    
    def _format_size_to_gb(self, item_data: list, default_msg: str) -> str:
        """格式化大小为GB"""
        if item_data and len(item_data) > 0:
            size_bytes = None
            for line in item_data:
                # 使用通用方法提取值，然后尝试解析
                val_str = self._extract_item_value(line)
                try:
                    size_bytes = int(val_str)
                    break
                except (ValueError, TypeError):
                    continue
            
            if size_bytes:
                size_gb = size_bytes / (1024 * 1024 * 1024)
                return f"{size_gb:.2f}GB"
        
        return default_msg
    
    def _format_size_to_mb_gb(self, item_data: list, default_msg: str) -> str:
        """格式化大小为MB或GB"""
        if item_data and len(item_data) > 0:
            size_bytes = None
            for line in item_data:
                val_str = self._extract_item_value(line)
                try:
                    size_bytes = int(val_str)
                    break
                except (ValueError, TypeError):
                    continue
            
            if size_bytes:
                if size_bytes >= 1024 * 1024 * 1024:
                    size_gb = size_bytes / (1024 * 1024 * 1024)
                    return f"{size_gb:.2f}GB"
                else:
                    size_mb = size_bytes / (1024 * 1024)
                    return f"{size_mb:.2f}MB"
        
        return default_msg
    
    def _format_size_to_mb(self, item_data: list, default_msg: str) -> str:
        """格式化大小为MB"""
        if item_data and len(item_data) > 0:
            size_bytes = None
            for line in item_data:
                val_str = self._extract_item_value(line)
                try:
                    size_bytes = int(val_str)
                    break
                except (ValueError, TypeError):
                    continue
            
            if size_bytes:
                size_mb = size_bytes / (1024 * 1024)
                return f"{size_mb:.2f}MB"
        
        return default_msg
    
    def _format_tables_without_pk(self, item_data: list) -> str:
        """格式化无主键表 - 已去重"""
        if len(item_data) > 0:
            table_info = []
            seen = set()
            for line in item_data:
                if line.strip() and not line.startswith('database_name'):
                    if ':' in line and not line.startswith('*'):
                        parts = line.split(':', 1)
                        if len(parts) >= 2:
                            line_content = parts[1].strip()
                            if '\t' in line_content:
                                line_parts = line_content.split('\t')
                                if len(line_parts) >= 2:
                                    db_name = line_parts[0].strip()
                                    table_name = line_parts[1].strip()
                                    key = f"{db_name}.{table_name}"
                                    if key not in seen:
                                        seen.add(key)
                                        table_info.append(key)
                    elif '\t' in line:
                        parts = line.split('\t')
                        if len(parts) >= 2:
                            db_name = parts[0].strip()
                            table_name = parts[1].strip()
                            key = f"{db_name}.{table_name}"
                            if key not in seen:
                                seen.add(key)
                                table_info.append(key)
            
            if len(table_info) > 0:
                if len(table_info) > 5:
                    return f"共 {len(table_info)} 个无主键表\n" + "\n".join(table_info[:5])
                else:
                    return "\n".join(table_info)
        return "未检测到无主键表"
    
    def _add_issue_deduplicated(self, issues_list, new_issue, key_field="issue"):
        """向问题列表中添加不重复的问题"""
        for existing in issues_list:
            if existing.get(key_field) == new_issue.get(key_field):
                return
        issues_list.append(new_issue)
    
    def _format_file_descriptor_limit_display(self, item_data: list) -> str:
        """格式化文件描述符限制显示值"""
        for line in item_data:
            line = line.strip()
            if "open files" in line.lower() or "max open files" in line.lower():
                parts = line.split()
                for p in parts:
                    try:
                        val = int(p)
                        if val >= 1024:
                            return str(val)
                    except ValueError:
                        continue
                return line
        return "未检测到"

    def _format_database_sizes(self, item_data: list) -> str:
        """格式化数据库大小"""
        if len(item_data) >= 1:
            db_info = []
            for line in item_data:
                if line.strip() and not line.startswith('database_name'):
                    parts = line.split('\t')
                    if len(parts) >= 2:
                        db_name = parts[0].strip()
                        size_mb = parts[1].strip()
                        if not size_mb.endswith('MB'):
                            size_mb = f"{size_mb}MB"
                        db_info.append(f"{db_name} {size_mb}")
            
            if len(db_info) > 5:
                return f"共 {len(db_info)} 个数据库\n" + "\n".join(db_info[:5])
            else:
                return "\n".join(db_info)
        return "未检测到数据库"
    
    def _format_empty_password_users(self, item_data: list) -> str:
        """格式化空密码用户"""
        if len(item_data) > 0:
            user_info = []
            for line in item_data:
                if line.strip():
                    parts = line.split('\t')
                    if len(parts) >= 2:
                        user = parts[0].strip()
                        host = parts[1].strip()
                        user_info.append(f"{user}@{host}")
            
            if len(user_info) > 0:
                if len(user_info) > 5:
                    return f"共 {len(user_info)} 个空密码用户\n" + "\n".join(user_info[:5])
                else:
                    return "\n".join(user_info)
        return "未检测到空密码用户"
    
    def _format_remote_root_users(self, item_data: list) -> str:
        """格式化远程root用户"""
        if len(item_data) > 0:
            user_info = []
            for line in item_data:
                if line.strip():
                    parts = line.split('\t')
                    if len(parts) >= 2:
                        user = parts[0].strip()
                        host = parts[1].strip()
                        user_info.append(f"{user}@{host}")
            
            if len(user_info) > 0:
                if len(user_info) > 5:
                    return f"共 {len(user_info)} 个远程root用户\n" + "\n".join(user_info[:5])
                else:
                    return "\n".join(user_info)
        return "未检测到远程root用户"
    
    def _format_remote_users(self, item_data: list) -> str:
        """格式化远程用户"""
        if len(item_data) > 0:
            user_info = []
            for line in item_data:
                if line.strip():
                    parts = line.split('\t')
                    if len(parts) >= 3:
                        user = parts[0].strip()
                        host = parts[1].strip()
                        user_info.append(f"{user}@{host}")
            
            if len(user_info) > 0:
                if len(user_info) > 5:
                    return f"共 {len(user_info)} 个远程用户\n" + "\n".join(user_info[:5])
                else:
                    return "\n".join(user_info)
        return "未检测到远程用户"
    
    def _analyze_inspection_item(self, item: str, item_data: list) -> Dict[str, str]:
        """分析单个巡检项的合理性"""
        # 如果没有数据，返回注意状态

        
        # 默认返回
        default_result = {
            "status": "正常",
            "suggestion": "配置合理，无需调整"
        }
        
        # 根据巡检项类型进行分析
        if item == "CURRENT_CONNECTIONS":
            return self._analyze_connections(item_data)
        elif item == "SLOW_QUERIES":
            return self._analyze_slow_queries(item_data)
        elif item == "SYSTEM_CPU_USAGE":
            return self._analyze_cpu_usage(item_data)
        elif item == "SYSTEM_MEMORY_USAGE":
            return self._analyze_memory_usage(item_data)
        elif item == "SYSTEM_DISK_USAGE":
            return self._analyze_disk_usage(item_data)
        elif item == "EMPTY_PASSWORD_USERS":
            return self._analyze_empty_password_users(item_data)
        elif item == "REMOTE_ROOT_USERS":
            return self._analyze_remote_root_users(item_data)
        elif item == "SSL_STATUS":
            return self._analyze_ssl_status(item_data)
        elif item == "REMOTE_USERS":
            return self._analyze_remote_users(item_data)
        elif item == "DATABASE_SIZES":
            return self._analyze_database_sizes(item_data)
        elif item == "FRAGMENTED_TABLES" or item == "TABLE_FRAGMENTATION_DETAILED":
            return self._analyze_fragmented_tables(item_data)
        elif item == "SYSTEM_SWAP_USAGE":
            return self._analyze_swap_usage(item_data)
        elif item == "TABLES_WITHOUT_PRIMARY_KEY":
            return self._analyze_tables_without_primary_key(item_data)
        elif item == "PARTITIONED_TABLES":
            return self._analyze_partitioned_tables(item_data)
        elif item == "GTID_MODE":
            return self._analyze_gtid_mode(item_data)
        elif item == "GTID_CONSISTENCY":
            return self._analyze_gtid_consistency(item_data)
        elif item == "GTID_EXECUTED":
            return {"status": "正常", "suggestion": "GTID执行状态正常"}
        elif item == "BINARY_LOGS":
            return self._analyze_binary_logs(item_data)
        elif item == "MGR_GROUP_MEMBERS" or item == "MGR_MEMBER_STATS":
            return self._analyze_mgr_cluster(item_data)
        elif item == "PXC_WSREP_STATUS" or item == "PXC_WSREP_VARIABLES":
            return self._analyze_pxc_cluster(item_data)
        elif item == "SLAVE_STATUS":
            return self._analyze_slave_status(item_data)
        elif item == "SLAVE_STATUS_DETAILED":
            return self._analyze_slave_status(item_data)
        elif item == "UPTIME":
            return self._analyze_uptime(item_data)
        # 新增巡检项分析
        elif item == "INNODB_FLUSH_LOG_AT_TRX_COMMIT":
            return self._analyze_innodb_flush_log(item_data)
        elif item == "SYNC_BINLOG":
            return self._analyze_sync_binlog(item_data)
        elif item == "BINLOG_FORMAT":
            return self._analyze_binlog_format(item_data)
        elif item == "TRANSACTION_ISOLATION":
            return self._analyze_transaction_isolation(item_data)
        elif item == "INNODB_ROW_LOCK_WAITS":
            return self._analyze_innodb_lock_waits(item_data)
        elif item == "INNODB_ROW_LOCK_TIME_AVG":
            return self._analyze_innodb_lock_time_avg(item_data)
        elif item == "CREATED_TMP_DISK_TABLES":
            return self._analyze_tmp_disk_tables(item_data)
        elif item == "ABORTED_CLIENTS":
            return self._analyze_aborted_clients(item_data)
        elif item == "ABORTED_CONNECTS":
            return self._analyze_aborted_connects(item_data)
        elif item == "OPEN_FILES_LIMIT":
            return self._analyze_open_files_limit(item_data)
        elif item == "INNODB_ENGINE_STATUS":
            return self._analyze_innodb_engine_status(item_data)
        elif item == "MYSQL_ERROR_LOG":
            return self._analyze_error_log(item_data)
        # 新增：其他配置项分析
        elif item == "INNODB_LOG_FILE_SIZE":
            return self._analyze_innodb_log_file_size(item_data)
        elif item == "INNODB_LOG_FILES_IN_GROUP":
            return self._analyze_config_value(item_data, "InnoDB日志文件数量")
        # 新增：性能统计项分析
        elif item == "QUESTIONS":
            return self._analyze_numeric_value(item_data, "查询总数", unit="次")
        elif item == "QUERIES":
            return self._analyze_numeric_value(item_data, "执行语句总数", unit="次")
        elif item == "COM_SELECT":
            return self._analyze_numeric_value(item_data, "SELECT语句数", unit="次")
        elif item == "COM_INSERT":
            return self._analyze_numeric_value(item_data, "INSERT语句数", unit="次")
        elif item == "COM_UPDATE":
            return self._analyze_numeric_value(item_data, "UPDATE语句数", unit="次")
        elif item == "COM_DELETE":
            return self._analyze_numeric_value(item_data, "DELETE语句数", unit="次")
        elif item == "INNODB_ROW_LOCK_TIME_MAX":
            return self._analyze_numeric_value(item_data, "InnoDB行锁最大等待时间", warning_threshold=5000, critical_threshold=1000000, unit="ms")
        elif item == "CREATED_TMP_TABLES":
            return self._analyze_numeric_value(item_data, "创建临时表数量", unit="个")
        elif item == "COM_COMMIT":
            return self._analyze_numeric_value(item_data, "COMMIT语句数", unit="次")
        elif item == "COM_ROLLBACK":
            return self._analyze_numeric_value(item_data, "ROLLBACK语句数", unit="次")
        # 新增：系统资源项分析
        elif item == "FILE_DESCRIPTOR_LIMIT":
            return self._analyze_file_descriptor_limit(item_data)
        elif item == "NETWORK_CONNECTIONS":
            return self._analyze_numeric_value(item_data, "网络连接数", unit="个")
        # 新增：表碎片和索引分析
        elif item == "TABLE_FRAGMENTATION_DETAILED":
            return self._analyze_table_fragmentation_item(item_data)
        elif item == "INDEX_DETAILS":
            return self._analyze_index_details_item(item_data)
        elif item == "INDEX_STATISTICS":
            return self._analyze_index_statistics_item(item_data)
        # 新增：数据库信息补充分析
        elif item == "DATABASE_VERSION":
            return self._analyze_database_version(item_data)
        elif item == "OS_VERSION":
            return self._analyze_os_version(item_data)
        elif item == "MASTER_STATUS":
            return self._analyze_master_status(item_data)
        elif item == "INNODB_FILE_PER_TABLE":
            return self._analyze_innodb_file_per_table(item_data)
        elif item == "INNODB_FLUSH_METHOD":
            return self._analyze_innodb_flush_method(item_data)
        elif item == "INNODB_BUFFER_POOL_SIZE":
            return self._analyze_innodb_buffer_pool(item_data)
        elif item == "MAX_CONNECTIONS":
            return self._analyze_max_connections(item_data)
        elif item == "TABLE_OPEN_CACHE":
            return self._analyze_table_open_cache(item_data)
        elif item == "MAX_ALLOWED_PACKET":
            return self._analyze_max_allowed_packet(item_data)
        elif item == "WAIT_TIMEOUT":
            return self._analyze_wait_timeout(item_data)
        elif item == "INTERACTIVE_TIMEOUT":
            return self._analyze_wait_timeout(item_data, True)
        
        return default_result
    
    def _format_uptime_value(self, item_data: list) -> str:
        """格式化运行时间值"""
        if item_data:
            try:
                # 提取运行时间（秒）
                uptime_seconds = int(item_data[0].split()[1] if len(item_data[0].split()) > 1 else item_data[0])
                
                # 转换为天数、小时、分钟、秒数
                days = uptime_seconds // (24 * 3600)
                hours = (uptime_seconds % (24 * 3600)) // 3600
                minutes = (uptime_seconds % 3600) // 60
                seconds = uptime_seconds % 60
                
                return f"{days}天{hours}小时{minutes}分钟{seconds}秒"
                    
            except (ValueError, IndexError):
                pass
        
        return "未检测到数据"
    
    def _get_instance_type(self) -> str:
        """判断实例类型：主库或从库"""
        # 检查是否有主库状态信息
        master_status_data = self.inspection_data.get("MASTER_STATUS", [])
        slave_status_data = self.inspection_data.get("SLAVE_STATUS", [])
        
        # 优先检查是否是从库：有 SLAVE_STATUS 数据且包含 Master_Host
        if slave_status_data and len(slave_status_data) > 1:
            for line in slave_status_data:
                # 处理垂直格式：Field: Value
                if ':' in line and not line.startswith('*'):
                    parts = line.split(':', 1)
                    if len(parts) >= 2 and "Master_Host" in parts[0]:
                        field_value = parts[1].strip()
                        if field_value and field_value.upper() != "NULL":
                            return "slave"
                # 处理制表符格式
                elif "Master_Host" in line and '\t' in line:
                    parts = line.split('\t')
                    if len(parts) >= 2 and parts[1].strip() and parts[1].strip().upper() != "NULL":
                        return "slave"
        
        # 再检查是否是主库：有 MASTER_STATUS 数据且有实际内容
        if master_status_data and len(master_status_data) > 1:
            for line in master_status_data:
                # 跳过表头行
                if "File" in line and "Position" in line:
                    continue
                # 检查是否有实际数据（制表符分隔的多列数据）
                if '\t' in line:
                    parts = [p.strip() for p in line.split('\t') if p.strip()]
                    if len(parts) >= 2 and parts[0]:  # 有实际的 binlog 文件名
                        return "master"
        
        return "unknown"
    
    def _is_mgr_or_pxc_cluster(self) -> bool:
        """判断当前实例是否为MGR或PXC集群"""
        has_mgr = False
        mgr_members = self.inspection_data.get("MGR_GROUP_MEMBERS", [])
        mgr_stats = self.inspection_data.get("MGR_MEMBER_STATS", [])
        if (mgr_members and len(mgr_members) > 1) or (mgr_stats and len(mgr_stats) > 1):
            has_mgr = True
        
        has_pxc = False
        pxc_status = self.inspection_data.get("PXC_WSREP_STATUS", [])
        pxc_vars = self.inspection_data.get("PXC_WSREP_VARIABLES", [])
        if (pxc_status and len(pxc_status) > 1) or (pxc_vars and len(pxc_vars) > 1):
            has_pxc = True
        
        return has_mgr or has_pxc
    
    def _format_master_status(self, item_data: list, slave_status_data: list = None) -> str:
        """格式化主库状态信息
        
        逻辑：
        1. 如果有 SHOW SLAVE STATUS 输出，说明是从库，显示主库连接信息
        2. 如果没有 SHOW SLAVE STATUS 输出，说明是主库，显示自身状态
        """
        # 获取实例类型
        instance_type = self._get_instance_type()
        is_cluster = self._is_mgr_or_pxc_cluster()
        
        # 先检查是否有从库状态信息
        slave_status_data = self.inspection_data.get("SLAVE_STATUS", [])
        if slave_status_data and len(slave_status_data) > 1:
            slave_status = {}
            for line in slave_status_data:
                # 处理垂直格式: Field: Value
                if ':' in line and not line.startswith('*'):
                    parts = line.split(':', 1)
                    if len(parts) >= 2:
                        field_name = parts[0].strip()
                        field_value = parts[1].strip()
                        slave_status[field_name] = field_value
                # 处理制表符格式
                elif '\t' in line:
                    parts = line.split('\t')
                    if len(parts) >= 2:
                        slave_status[parts[0].strip()] = parts[1].strip()
            
            # 提取主库连接信息
            if "Master_Host" in slave_status and "Master_Port" in slave_status:
                master_info = []
                master_info.append(f"主库IP: {slave_status['Master_Host']}")
                master_info.append(f"主库端口: {slave_status['Master_Port']}")
                
                # 显示更多主库连接信息
                if "Master_User" in slave_status:
                    master_info.append(f"复制用户: {slave_status['Master_User']}")
                if "Master_Log_File" in slave_status:
                    master_info.append(f"同步Binlog: {slave_status['Master_Log_File']}")
                if "Read_Master_Log_Pos" in slave_status:
                    master_info.append(f"同步位置: {slave_status['Read_Master_Log_Pos']}")
                
                return "\n".join(master_info)
        
        # 如果没有从库状态信息，说明是主库，显示自身状态
        if instance_type == "master" or instance_type == "slave_and_master":
            if item_data and len(item_data) > 0:
                # 尝试解析主库状态数据
                master_info = []
                for line in item_data:
                    # 跳过可能的表头（包含 File 或 Position 的行）
                    line_stripped = line.strip()
                    if not line_stripped:
                        continue
                    if "File" in line_stripped and "Position" in line_stripped:
                        continue
                    
                    # 处理有表头或无表头的数据
                    if "\t" in line:
                        parts = [p.strip() for p in line.split("\t") if p.strip()]
                        if len(parts) >= 2:
                            master_info.append(f"Binlog文件: {parts[0]}")
                            master_info.append(f"Binlog位置: {parts[1]}")
                            break
                    elif line_stripped:
                        # 直接值的情况
                        master_info.append(line_stripped)
                
                if master_info:
                    return "\n".join(master_info)
            
            if is_cluster:
                return "当前实例为MGR/PXC集群节点"
            return "当前实例为主库"
        elif instance_type == "slave":
            # 从从库状态信息中获取主库连接信息
            slave_status_data = self.inspection_data.get("SLAVE_STATUS", [])
            if slave_status_data and len(slave_status_data) > 1:
                slave_status = {}
                for line in slave_status_data:
                    # 处理垂直格式: Field: Value
                    if ':' in line and not line.startswith('*'):
                        parts = line.split(':', 1)
                        if len(parts) >= 2:
                            field_name = parts[0].strip()
                            field_value = parts[1].strip()
                            slave_status[field_name] = field_value
                    # 处理制表符格式
                    elif '\t' in line:
                        parts = line.split('\t')
                        if len(parts) >= 2:
                            slave_status[parts[0].strip()] = parts[1].strip()
                
                if "Master_Host" in slave_status and "Master_Port" in slave_status:
                    return f"主库IP: {slave_status['Master_Host']}\n主库端口: {slave_status['Master_Port']}"
            
            return "无法获取主库连接信息"
        else:
            return "实例类型未知"
    
    def _analyze_slave_status(self, item_data: list) -> Dict[str, str]:
        """分析从库状态，返回健康状态和优化建议"""
        # MGR/PXC集群无需传统主从同步检查
        if self._is_mgr_or_pxc_cluster():
            return {
                "status": "正常",
                "suggestion": "当前实例为MGR/PXC集群，无需传统主从同步检查"
            }
        
        if not item_data or len(item_data) <= 1:
            instance_type = self._get_instance_type()
            if instance_type == "master":
                return {
                    "status": "正常",
                    "suggestion": "当前实例为主库，无需从库状态检查"
                }
            return {
                "status": "异常",
                "suggestion": "未配置从库或未检测到状态"
            }
        
        # 解析从库状态信息
        slave_status = {}
        for line in item_data:
            # 处理垂直格式: Field: Value
            if ':' in line and not line.startswith('*'):
                parts = line.split(':', 1)
                if len(parts) >= 2:
                    field_name = parts[0].strip()
                    field_value = parts[1].strip()
                    slave_status[field_name] = field_value
            # 处理制表符格式
            elif '\t' in line:
                parts = line.split('\t')
                if len(parts) >= 2:
                    slave_status[parts[0].strip()] = parts[1].strip()
        
        # 分析健康状态
        io_running = slave_status.get("Slave_IO_Running", "").upper()
        sql_running = slave_status.get("Slave_SQL_Running", "").upper()
        
        is_healthy = io_running == "YES" and sql_running == "YES"
        status = "正常" if is_healthy else "异常"
        
        # 生成优化建议
        suggestion = []
        
        if io_running != "YES":
            last_io_error = slave_status.get("Last_IO_Error", "")
            if last_io_error:
                suggestion.append(f"IO线程异常: {last_io_error[:100]}...")
            suggestion.append("请检查主库连接或网络状况")
        
        if sql_running != "YES":
            last_sql_error = slave_status.get("Last_SQL_Error", "")
            if last_sql_error:
                suggestion.append(f"SQL线程异常: {last_sql_error[:100]}...")
            suggestion.append("请检查数据一致性或SQL语法")
        
        delay = slave_status.get("Seconds_Behind_Master", "NULL")
        if delay == "NULL":
            suggestion.append("复制延迟未知，可能存在复制异常")
        elif delay != "0" and int(delay) > 30:
            suggestion.append(f"复制延迟超过{delay}秒，建议检查网络状况或优化主库性能")
        
        auto_position = slave_status.get("Master_Auto_Position", "0")
        if auto_position == "0":
            suggestion.append("建议启用GTID自动定位以提高复制可靠性")
        
        if not suggestion:
            suggestion.append("配置合理，无需调整")
        
        return {
            "status": status,
            "suggestion": "\n".join(suggestion)
        }
    
    def _format_slave_status(self, item_data: list) -> str:
        """格式化从库状态信息，包含关键复制指标"""
        # MGR/PXC集群无需传统主从同步
        if self._is_mgr_or_pxc_cluster():
            return "当前实例为MGR/PXC集群，无需传统主从同步检查"
        
        instance_type = self._get_instance_type()
        
        if instance_type == "master":
            return "当前实例为主库，无从库状态信息"
        
        if not item_data or len(item_data) <= 1:
            return "未配置从库或未检测到状态"
        
        # 确保是从库类型才继续分析
        if instance_type == "slave" or instance_type == "slave_and_master":
            # 解析从库状态信息
            slave_status = {}
            for line in item_data:
                # 处理垂直格式: Field: Value
                if ':' in line and not line.startswith('*'):
                    parts = line.split(':', 1)
                    if len(parts) >= 2:
                        field_name = parts[0].strip()
                        field_value = parts[1].strip()
                        slave_status[field_name] = field_value
                # 处理制表符格式
                elif '\t' in line:
                    parts = line.split('\t')
                    if len(parts) >= 2:
                        slave_status[parts[0].strip()] = parts[1].strip()
            
            # 格式化显示关键信息
            formatted_info = []
            
            # 复制线程状态
            io_running = slave_status.get("Slave_IO_Running", "").upper()
            sql_running = slave_status.get("Slave_SQL_Running", "").upper()
            
            io_status = "运行中" if io_running == "YES" else "停止"
            sql_status = "运行中" if sql_running == "YES" else "停止"
            
            formatted_info.append(f"IO线程: {io_status}")
            formatted_info.append(f"SQL线程: {sql_status}")
            
            # 显示IO线程错误
            last_io_error = slave_status.get("Last_IO_Error", "")
            if io_running != "YES" and last_io_error:
                formatted_info.append(f"IO线程错误: {last_io_error[:100]}...")
            
            # 显示SQL线程错误
            last_sql_error = slave_status.get("Last_SQL_Error", "")
            if sql_running != "YES" and last_sql_error:
                formatted_info.append(f"SQL线程错误: {last_sql_error[:100]}...")
            
            # 复制延迟
            delay = slave_status.get("Seconds_Behind_Master", "NULL")
            if delay == "NULL":
                formatted_info.append("复制延迟: 未知")
            elif delay == "0":
                formatted_info.append("复制延迟: 无延迟")
            else:
                formatted_info.append(f"复制延迟: {delay}秒")
            
            # GTID自动定位状态
            auto_position = slave_status.get("Master_Auto_Position", "0")
            auto_position_status = "启用" if auto_position == "1" else "禁用"
            formatted_info.append(f"GTID自动定位: {auto_position_status}")
            
            # 已执行的GTID集
            executed_gtid = slave_status.get("Executed_Gtid_Set", "")
            if executed_gtid:
                formatted_info.append(f"已执行GTID集: {executed_gtid[:100]}...")
            
            # 最后错误信息（兼容旧版本）
            last_error = slave_status.get("Last_Error", "")
            if last_error and not (last_io_error or last_sql_error):
                formatted_info.append(f"最后错误: {last_error[:100]}...")
            
            return "\n".join(formatted_info) if formatted_info else "从库状态信息不完整"
        else:
            return "实例类型未知，无法显示从库状态"
    
    def _analyze_uptime(self, item_data: list) -> Dict[str, str]:
        """分析运行时间 - 按文档标准三态"""
        if item_data:
            try:
                uptime_seconds = int(item_data[0].split()[1] if len(item_data[0].split()) > 1 else item_data[0])
                if uptime_seconds > 864000:
                    return {"status": "正常", "suggestion": f"运行时间: {uptime_seconds//86400}天，系统稳定运行中"}
                elif uptime_seconds > 3600:
                    return {"status": "注意", "suggestion": f"运行时间: {uptime_seconds//3600}小时，建议持续监控"}
                else:
                    return {"status": "异常", "suggestion": f"运行时间不足1小时({uptime_seconds}秒)，数据库刚重启，请检查原因"}
            except (ValueError, IndexError):
                pass
        return {"status": "正常", "suggestion": "运行时间正常"}

    def _analyze_connections(self, item_data: list) -> Dict[str, str]:
        """分析连接数使用率 - 计算当前/最大百分比"""
        max_conn = self._get_value("MAX_CONNECTIONS", 1)
        if max_conn and len(item_data) > 1:
            try:
                current_conn = int(item_data[1].split()[1])
                usage_rate = current_conn / max_conn
                if usage_rate > 0.9:
                    return {"status": "异常", "suggestion": f"连接数使用率过高: {usage_rate:.1%}（{current_conn}/{max_conn}），建议增加max_connections或优化连接池"}
                elif usage_rate > 0.8:
                    return {"status": "注意", "suggestion": f"连接数使用率较高: {usage_rate:.1%}（{current_conn}/{max_conn}），建议监控并适时调整"}
            except (ValueError, IndexError):
                pass
        return {"status": "正常", "suggestion": "连接数使用率在合理范围内"}

    def _analyze_slow_queries(self, item_data: list) -> Dict[str, str]:
        """分析慢查询 - 按文档三态"""
        if len(item_data) > 1:
            try:
                slow_count = int(item_data[1].split()[1])
                if slow_count > 100:
                    return {"status": "异常", "suggestion": f"慢查询数量过多: {slow_count}，建议立即检查慢查询日志并优化SQL"}
                elif slow_count > 10:
                    return {"status": "注意", "suggestion": f"慢查询数量较多: {slow_count}，建议定期检查慢查询日志"}
                elif slow_count > 0:
                    return {"status": "正常", "suggestion": f"存在少量慢查询({slow_count})，建议持续监控"}
            except (ValueError, IndexError):
                pass
        return {"status": "正常", "suggestion": "无慢查询，性能良好"}
    
    def _analyze_cpu_usage(self, item_data: list) -> Dict[str, str]:
        """分析CPU使用率 - 按文档三态"""
        if item_data:
            cpu_usage = self._extract_percentage(item_data[0])
            if cpu_usage > 80:
                return {"status": "异常", "suggestion": f"CPU使用率过高: {cpu_usage}%，建议检查高负载进程并优化"}
            elif cpu_usage > 50:
                return {"status": "注意", "suggestion": f"CPU使用率偏高: {cpu_usage}%，建议持续监控"}
        return {"status": "正常", "suggestion": "CPU使用率正常"}
    
    def _analyze_memory_usage(self, item_data: list) -> Dict[str, str]:
        """分析内存使用率 - 按文档三态"""
        if item_data:
            memory_usage = self._extract_percentage(item_data[0])
            if memory_usage > 80:
                return {"status": "异常", "suggestion": f"内存使用率过高: {memory_usage}%，可能影响性能"}
            elif memory_usage > 60:
                return {"status": "注意", "suggestion": f"内存使用率较高: {memory_usage}%，建议监控"}
        return {"status": "正常", "suggestion": "内存使用率正常"}
    
    def _analyze_disk_usage(self, item_data: list) -> Dict[str, str]:
        """分析磁盘使用率 - 按文档三态"""
        if item_data:
            disk_usage = self._extract_percentage(item_data[0])
            if disk_usage > 85:
                return {"status": "异常", "suggestion": f"磁盘使用率过高: {disk_usage}%，需要立即清理或扩容"}
            elif disk_usage > 70:
                return {"status": "注意", "suggestion": f"磁盘使用率偏高: {disk_usage}%，建议清理或扩容"}
        return {"status": "正常", "suggestion": "磁盘空间充足"}
    
    def _analyze_empty_password_users(self, item_data: list) -> Dict[str, str]:
        """分析空密码用户"""
        if len(item_data) > 0:
            # 统计实际的空密码用户数量
            actual_count = 0
            for line in item_data:
                if line.strip():
                    actual_count += 1
            
            if actual_count > 0:
                return {
                    "status": "异常",
                    "suggestion": f"发现 {actual_count} 个空密码用户，这是严重的安全隐患！建议立即为这些用户设置强密码，并启用密码复杂度策略"
                }
        return {"status": "正常", "suggestion": "无空密码用户，安全配置良好"}
    
    def _analyze_remote_root_users(self, item_data: list) -> Dict[str, str]:
        """分析远程root用户"""
        if len(item_data) > 0:
            # 统计实际的远程root用户数量
            actual_count = 0
            for line in item_data:
                if line.strip():
                    actual_count += 1
            
            if actual_count > 0:
                return {
                    "status": "异常",
                    "suggestion": f"发现 {actual_count} 个远程root用户！root用户不应允许从远程登录，建议修改root用户的host为localhost或127.0.0.1"
                }
        return {"status": "正常", "suggestion": "无远程root用户，安全配置良好"}
    
    def _analyze_remote_users(self, item_data: list) -> Dict[str, str]:
        """分析远程用户"""
        if len(item_data) > 0:
            # 统计实际的远程用户数量
            actual_count = 0
            for line in item_data:
                if line.strip():
                    actual_count += 1
            
            if actual_count > 0:
                return {
                    "status": "注意",
                    "suggestion": f"发现 {actual_count} 个允许任意主机访问的用户（host为%）。建议限制这些用户的访问来源，使用具体的IP地址或网段"
                }
        return {"status": "正常", "suggestion": "无任意主机访问的用户，安全配置良好"}
    
    def _analyze_ssl_status(self, item_data: list) -> Dict[str, str]:
        """分析SSL状态"""
        if item_data and "YES" not in item_data[0].upper():
            return {
                "status": "注意",
                "suggestion": "SSL未启用，建议启用加密传输"
            }
        return {"status": "正常", "suggestion": "SSL已启用，数据传输安全"}
    
    def _analyze_database_sizes(self, item_data: list) -> Dict[str, str]:
        """分析数据库大小 - 按文档三态阈值"""
        if item_data and len(item_data) > 0:
            total_size = 0
            db_count = 0
            
            for line in item_data:
                line_stripped = line.strip()
                if not line_stripped or line_stripped.lower().startswith('database_name'):
                    continue
                parts = line_stripped.split('\t')
                if len(parts) >= 2:
                    try:
                        total_size += float(parts[1])
                        db_count += 1
                    except (ValueError, IndexError):
                        continue
            
            total_size_gb = total_size / 1024
            suggestion = f"共 {db_count} 个数据库，总大小: {total_size_gb:.2f}GB"
            
            if total_size_gb > 1000:
                return {"status": "异常", "suggestion": suggestion + "，数据库总容量超过1TB，建议立即清理历史数据并考虑分库分表"}
            elif total_size_gb > 500:
                return {"status": "注意", "suggestion": suggestion + "，数据库总容量超过500GB，建议定期清理历史数据"}
            return {"status": "正常", "suggestion": suggestion}
        return {"status": "正常", "suggestion": "数据库大小正常"}
    
    def _analyze_fragmented_tables(self, item_data: list) -> Dict[str, str]:
        """分析碎片化表（使用详细数据进行筛选）"""
        # 使用详细数据 TABLE_FRAGMENTATION_DETAILED 来筛选碎片大小>2GB且碎片率>20%的表
        frag_data = self.inspection_data.get("TABLE_FRAGMENTATION_DETAILED", [])
        if len(frag_data) > 0:
            # 统计碎片大小>2GB且碎片率>20%的表
            tables = []
            for line in frag_data:
                if line.strip():
                    parts = line.split('\t')
                    if len(parts) >= 7:
                        try:
                            frag_size_mb = float(parts[5])  # 碎片大小
                            frag_percent = float(parts[6])  # 碎片率
                            # 过滤条件：碎片大小 > 2048MB (2GB) 且 碎片率 > 20%
                            if frag_size_mb > 2048.0 and frag_percent > 20.0:
                                tables.append(f"{parts[0]}.{parts[1]}")
                        except (ValueError, IndexError):
                            continue
            
            if tables:
                if len(tables) > 10:
                    return {
                        "status": "异常",
                        "suggestion": f"发现 {len(tables)} 个碎片大小>2GB且碎片率>20%的表，建议立即对这些表进行OPTIMIZE TABLE操作"
                    }
                else:
                    return {
                        "status": "注意",
                        "suggestion": f"发现 {len(tables)} 个碎片大小>2GB且碎片率>20%的表，建议对这些表进行OPTIMIZE TABLE操作"
                    }
        return {"status": "正常", "suggestion": "无严重碎片化表，存储状态良好"}
    
    def _analyze_swap_usage(self, item_data: list) -> Dict[str, str]:
        """分析swap使用率"""
        if item_data:
            swap_usage = self._extract_percentage(item_data[0])
            if swap_usage > 50:
                return {
                    "status": "注意",
                    "suggestion": "swap使用率较高，可能影响性能"
                }
            elif swap_usage > 20:
                return {
                    "status": "正常",
                    "suggestion": "swap使用率适中，建议关注"
                }
        return {"status": "正常", "suggestion": "swap使用率正常"}
    
    def _analyze_tables_without_primary_key(self, item_data: list) -> Dict[str, str]:
        """分析无主键表 - 按文档三态"""
        table_count = 0
        for line in item_data:
            if line.strip() and not line.startswith('database_name'):
                if '\t' in line or ':' in line:
                    table_count += 1
        
        if table_count > 10:
            return {"status": "异常", "suggestion": f"发现 {table_count} 个无主键表，严重影响查询性能和数据完整性，建议立即添加主键"}
        elif table_count > 0:
            return {"status": "注意", "suggestion": f"发现 {table_count} 个无主键表，建议为这些表添加主键，提升查询性能和数据完整性"}
        return {"status": "正常", "suggestion": "所有表都有主键，存储结构良好"}
    
    def _analyze_partitioned_tables(self, item_data: list) -> Dict[str, str]:
        """分析分区表"""
        from datetime import datetime
        
        if len(item_data) > 1:  # 第一行是标题
            # 检查时间分区表是否需要添加新分区
            time_partitioned_tables = []
            for line in item_data[1:]:
                if "TO_DAYS" in line or "UNIX_TIMESTAMP" in line:
                    # 这里可以添加更复杂的时间分区检查逻辑
                    time_partitioned_tables.append(line.split('\t')[1] if '\t' in line else line)
            
            if time_partitioned_tables:
                return {
                    "status": "正常",
                    "suggestion": f"发现 {len(time_partitioned_tables)} 个时间分区表，建议定期检查分区是否需要扩展"
                }
            
            return {
                "status": "正常",
                "suggestion": f"发现 {len(item_data)-1} 个分区表，分区策略合理"
            }
        return {"status": "正常", "suggestion": "无分区表，存储结构简单"}
    
    def _analyze_gtid_mode(self, item_data: list) -> Dict[str, str]:
        """分析GTID模式"""
        if item_data and "ON" in item_data[0].upper():
            return {
                "status": "正常",
                "suggestion": "GTID模式已启用，复制配置良好"
            }
        return {
            "status": "注意",
            "suggestion": "GTID模式未启用，建议启用以提高复制可靠性"
        }
    
    def _analyze_gtid_consistency(self, item_data: list) -> Dict[str, str]:
        """分析GTID一致性"""
        if not item_data:
            return {
                "status": "注意",
                "suggestion": "GTID一致性状态未知，建议检查配置"
            }
        
        # 解析GTID一致性状态（支持垂直格式和制表符格式）
        gtid_consistency = None
        for line in item_data:
            # 处理垂直格式: Field: Value
            if ':' in line and not line.startswith('*'):
                parts = line.split(':', 1)
                if len(parts) >= 2 and "enforce_gtid_consistency" in parts[0].lower():
                    gtid_consistency = parts[1].strip()
                    break
            # 处理制表符格式
            elif '\t' in line:
                parts = line.split('\t')
                if len(parts) >= 2 and "enforce_gtid_consistency" in parts[0].lower():
                    gtid_consistency = parts[1].strip()
                    break
        
        if gtid_consistency and gtid_consistency.upper() == "ON":
            return {
                "status": "正常",
                "suggestion": "GTID一致性已启用，数据一致性良好"
            }
        else:
            return {
                "status": "注意",
                "suggestion": "GTID一致性未启用，建议启用以确保数据一致性"
            }

    def _format_binary_logs(self, item_data: list) -> str:
        """格式化二进制日志信息"""
        if not item_data or len(item_data) == 0:
            return "未检测到二进制日志"
        
        result = []
        total_size_bytes = 0
        
        for line in item_data:
            line = line.strip()
            if not line:
                continue
            
            parts = line.split('\t')
            if len(parts) >= 2:
                filename = parts[0].strip()
                size_str = parts[1].strip()
                
                try:
                    size_bytes = int(size_str)
                    total_size_bytes += size_bytes
                    
                    # 转换为GB
                    size_gb = size_bytes / (1024 * 1024 * 1024)
                    if size_gb >= 1:
                        size_display = f"{size_gb:.2f}GB"
                    else:
                        size_mb = size_bytes / (1024 * 1024)
                        size_display = f"{size_mb:.2f}MB"
                    
                    result.append(f"{filename}: {size_display}")
                except ValueError:
                    continue
        
        if total_size_bytes > 0:
            total_gb = total_size_bytes / (1024 * 1024 * 1024)
            result.insert(0, f"总大小: {total_gb:.2f}GB")
            result.insert(1, f"日志数量: {len(result) - 1}个")
            result.insert(2, "-" * 30)
        
        return "\n".join(result) if result else "未检测到有效的二进制日志"
    
    def _analyze_binary_logs(self, item_data: list) -> Dict[str, str]:
        """分析二进制日志 - 检查是否开启及保留天数"""
        if not item_data or len(item_data) == 0:
            return {"status": "异常", "suggestion": "二进制日志未开启，主从复制和高可用场景必须开启binlog"}
        
        log_count = len(item_data)
        total_size_bytes = 0
        for line in item_data:
            line = line.strip()
            if not line:
                continue
            parts = line.split('\t')
            if len(parts) >= 2:
                try:
                    size_bytes = int(parts[1].strip())
                    total_size_bytes += size_bytes
                except ValueError:
                    continue
        
        total_gb = total_size_bytes / (1024 * 1024 * 1024)
        
        # 检查expire_logs_days
        expire_days = None
        expire_data = self.inspection_data.get("EXPIRE_LOGS_DAYS", [])
        if expire_data:
            expire_days = self._extract_value(expire_data)
        
        status = "正常"
        suggestions = [f"二进制日志已开启，共 {log_count} 个文件，总大小: {total_gb:.2f}GB"]
        
        if expire_days is None or expire_days == "0":
            status = "注意"
            suggestions.append("二进制日志保留天数未设置或为0，建议设置 expire_logs_days=7 以自动清理")
        elif expire_days and int(expire_days) < 7:
            status = "注意"
            suggestions.append(f"二进制日志保留天数过短: {expire_days}天，建议不少于7天")
        else:
            suggestions.append(f"保留天数: {expire_days}天，配置合理")
        
        if total_gb > 50:
            status = "异常"
            suggestions.append(f"二进制日志占用空间过大: {total_gb:.2f}GB，建议立即使用 PURGE BINARY LOGS 清理")
        
        return {"status": status, "suggestion": "\n".join(suggestions)}
    
    def _format_mgr_group_members(self, item_data: list) -> str:
        """格式化MGR集群成员信息"""
        if not item_data or len(item_data) == 0:
            return "未检测到MGR集群成员信息"
        
        result = []
        for line in item_data:
            line = line.strip()
            if not line:
                continue
            result.append(line)
        
        return "\n".join(result) if result else "未检测到有效的MGR集群成员信息"
    
    def _format_mgr_member_stats(self, item_data: list) -> str:
        """格式化MGR成员统计信息"""
        if not item_data or len(item_data) == 0:
            return "未检测到MGR成员统计信息"
        
        result = []
        for line in item_data:
            line = line.strip()
            if not line:
                continue
            result.append(line)
        
        return "\n".join(result) if result else "未检测到有效的MGR成员统计信息"
    
    def _format_pxc_wsrep_status(self, item_data: list) -> str:
        """格式化PXC集群状态信息"""
        if not item_data or len(item_data) == 0:
            return "未检测到PXC集群状态信息"
        
        result = []
        for line in item_data:
            line = line.strip()
            if not line:
                continue
            result.append(line)
        
        return "\n".join(result) if result else "未检测到有效的PXC集群状态信息"
    
    def _format_pxc_wsrep_variables(self, item_data: list) -> str:
        """格式化PXC集群配置信息"""
        if not item_data or len(item_data) == 0:
            return "未检测到PXC集群配置信息"
        
        result = []
        for line in item_data:
            line = line.strip()
            if not line:
                continue
            result.append(line)
        
        return "\n".join(result) if result else "未检测到有效的PXC集群配置信息"
    
    def _analyze_mgr_cluster(self, item_data: list) -> Dict[str, str]:
        """分析MGR集群状态
        会同时分析MGR_GROUP_MEMBERS和MGR_MEMBER_STATS
        """
        # 同时获取两个MGR相关数据
        mgr_members_data = self.inspection_data.get("MGR_GROUP_MEMBERS", [])
        mgr_stats_data = self.inspection_data.get("MGR_MEMBER_STATS", [])
        
        if (not mgr_members_data or len(mgr_members_data) <= 1) and \
           (not mgr_stats_data or len(mgr_stats_data) <= 1):
            return {"status": "注意", "suggestion": "未检测到MGR集群信息，请确认是否使用--mgr-check参数"}
        
        issues = []
        warnings = []
        
        # 分析MGR集群成员分析
        if mgr_members_data and len(mgr_members_data) > 1:
            for i in range(1, len(mgr_members_data)):
                line = mgr_members_data[i].strip()
                if not line:
                    continue
                
                parts = line.split('\t')
                if len(parts) < 8:
                    continue
                
                member_id = parts[1] if len(parts) > 1 else ""
                host = parts[2] if len(parts) > 2 else ""
                state = parts[4] if len(parts) > 4 else ""
                role = parts[5] if len(parts) > 5 else ""
                
                # 检查成员状态
                if state.upper() != "ONLINE":
                    issues.append(f"成员 {host}({member_id}) 状态异常: {state}")
                
                # 检查是否有主节点
                if role.upper() == "PRIMARY":
                    has_primary = True
        
        # 分析MGR成员统计分析
        if mgr_stats_data and len(mgr_stats_data) > 1:
            for i in range(1, len(mgr_stats_data)):
                line = mgr_stats_data[i].strip()
                if not line:
                    continue
                
                parts = line.split('\t')
                if len(parts) < 14:
                    continue
                
                member_id = parts[2] if len(parts) > 2 else ""
                count_conflict = int(parts[5]) if len(parts) > 5 else 0
                queue_transactions = int(parts[10]) if len(parts) > 10 else 0
                
                # 检查冲突次数
                if count_conflict > 10:
                    warnings.append(f"成员 {member_id} 冲突次数过高: {count_conflict}")
                
                # 检查队列事务数
                if queue_transactions > 1000:
                    warnings.append(f"成员 {member_id} 队列事务数过高: {queue_transactions}")
        
        # 综合分析结果
        if issues:
            return {
                "status": "异常",
                "suggestion": "; ".join(issues) + ("; " + "; ".join(warnings) if warnings else "")
            }
        elif warnings:
            return {
                "status": "注意",
                "suggestion": "; ".join(warnings)
            }
        else:
            return {
                "status": "正常",
                "suggestion": "MGR集群状态良好"
            }
    
    def _analyze_pxc_cluster(self, item_data: list) -> Dict[str, str]:
        """分析PXC集群状态 - 按文档标准"""
        if not item_data or len(item_data) <= 1:
            return {"status": "注意", "suggestion": "未检测到PXC集群信息，请确认是否使用--pxc-check参数"}
        
        all_primary = False
        all_synced = True
        node_count = 0
        for line in item_data[1:]:
            line = line.strip()
            if not line:
                continue
            parts = line.split('\t')
            if len(parts) >= 3:
                node_count += 1
                status_val = parts[2].strip().upper() if len(parts) > 2 else ""
                if "PRIMARY" in status_val:
                    all_primary = True
                if "SYNCED" not in status_val and "DONOR" not in status_val:
                    all_synced = False
        
        if node_count == 0:
            return {"status": "异常", "suggestion": "未检测到PXC节点，集群可能已关闭"}
        
        if all_primary and all_synced:
            return {"status": "正常", "suggestion": f"PXC集群状态良好: {node_count} 个节点，均为Primary且同步正常"}
        elif all_primary:
            return {"status": "注意", "suggestion": f"PXC集群: {node_count} 个节点，Primary但部分节点未同步，建议检查网络状况"}
        else:
            return {"status": "异常", "suggestion": f"PXC集群异常: {node_count} 个节点，非Primary状态，请检查集群健康状况"}
    
    def _format_table_fragmentation_detailed(self, item_data: list) -> list:
        """格式化表碎片详细信息，过滤只显示碎片率>20%的表，返回列表用于表格"""
        if not item_data or len(item_data) <= 1:
            return []
        
        result = []
        
        for i in range(1, len(item_data)):
            line = item_data[i].strip()
            if not line:
                continue
            
            parts = line.split('\t')
            if len(parts) < 6:
                continue
            
            try:
                frag_percent = float(parts[5])
                if frag_percent > 20:
                    result.append(parts)
            except (ValueError, IndexError):
                continue
        
        return result
    
    def _format_index_details(self, item_data: list) -> str:
        """格式化索引详细信息（美化展示）"""
        if not item_data or len(item_data) <= 1:
            return "未检测到索引信息"
        
        result = []
        show_count = min(50, len(item_data) - 1)
        
        current_db = ""
        current_table = ""
        
        for i in range(1, show_count + 1):
            line = item_data[i].strip()
            if not line:
                continue
            
            parts = line.split('\t')
            if len(parts) < 4:
                result.append(line)
                continue
            
            table_schema = parts[0]
            table_name = parts[1]
            index_name = parts[2]
            columns = parts[3]
            
            # 按数据库和表分组展示
            if table_schema != current_db:
                current_db = table_schema
                current_table = table_name
                result.append(f"\n【数据库】{table_schema}")
                result.append(f"  【表】{table_name}")
                result.append(f"    • 索引: {index_name}")
                result.append(f"      列: {columns}")
            elif table_name != current_table:
                current_table = table_name
                result.append(f"  【表】{table_name}")
                result.append(f"    • 索引: {index_name}")
                result.append(f"      列: {columns}")
            else:
                result.append(f"    • 索引: {index_name}")
                result.append(f"      列: {columns}")
        
        if len(item_data) - 1 > show_count:
            result.append(f"\n... 还有 {len(item_data) - 1 - show_count} 个索引未显示")
        
        return "\n".join(result)
    
    def _format_index_statistics(self, item_data: list) -> list:
        """格式化索引统计信息，返回列表用于表格展示"""
        if not item_data or len(item_data) <= 1:
            return []
        
        result = []
        
        for i in range(1, len(item_data)):
            line = item_data[i].strip()
            if not line:
                continue
            
            parts = line.split('\t')
            if len(parts) < 3:
                parts = line.split()
                if len(parts) < 3:
                    continue
            
            result.append(parts)
        
        return result
    
    def _analyze_innodb_flush_log(self, item_data: list) -> Dict[str, str]:
        """分析InnoDB日志刷新策略"""
        if not item_data or len(item_data) == 0:
            return {"status": "注意", "suggestion": "未检测到数据，请检查巡检脚本配置"}
        value = self._extract_value(item_data)
        if value == "1":
            return {"status": "正常", "suggestion": "innodb_flush_log_at_trx_commit=1，最高数据安全级，每次事务刷新日志"}
        elif value == "2":
            return {"status": "注意", "suggestion": "innodb_flush_log_at_trx_commit=2，性能优先，每秒刷新一次，崩溃可能丢失1秒数据"}
        elif value == "0":
            return {"status": "异常", "suggestion": "innodb_flush_log_at_trx_commit=0，性能最高，但崩溃可能丢失数据，生产环境不建议"}
        return {"status": "正常", "suggestion": "配置合理"}
    
    def _analyze_sync_binlog(self, item_data: list) -> Dict[str, str]:
        """分析Binlog同步策略"""
        if not item_data or len(item_data) == 0:
            return {"status": "注意", "suggestion": "未检测到数据，请检查巡检脚本配置"}
        value = self._extract_value(item_data)
        if value == "1":
            return {"status": "正常", "suggestion": "sync_binlog=1，最高数据安全级，每次事务同步binlog"}
        elif value == "0":
            return {"status": "注意", "suggestion": "sync_binlog=0，由操作系统控制同步，性能较高但数据安全性降低"}
        else:
            return {"status": "正常", "suggestion": f"sync_binlog={value}，每{value}次事务同步一次"}
    
    def _analyze_binlog_format(self, item_data: list) -> Dict[str, str]:
        """分析Binlog格式"""
        if not item_data or len(item_data) == 0:
            return {"status": "注意", "suggestion": "未检测到数据，请检查巡检脚本配置"}
        value = self._extract_value(item_data)
        if value.upper() == "ROW":
            return {"status": "正常", "suggestion": "binlog_format=ROW，行级复制，数据一致性最好"}
        elif value.upper() == "MIXED":
            return {"status": "正常", "suggestion": "binlog_format=MIXED，混合模式，自动选择合适格式"}
        elif value.upper() == "STATEMENT":
            return {"status": "注意", "suggestion": "binlog_format=STATEMENT，语句级复制，可能存在数据一致性问题"}
        return {"status": "正常", "suggestion": "配置合理"}
    
    def _analyze_transaction_isolation(self, item_data: list) -> Dict[str, str]:
        """分析事务隔离级别 - 按文档标准"""
        if not item_data or len(item_data) == 0:
            return {"status": "注意", "suggestion": "未检测到数据，请检查巡检脚本配置"}
        value = self._extract_value(item_data)
        if "READ-COMMITTED" in value.upper():
            return {"status": "正常", "suggestion": "事务隔离级别为READ-COMMITTED，推荐用于大多数OLTP场景"}
        elif "REPEATABLE-READ" in value.upper():
            return {"status": "注意", "suggestion": "事务隔离级别为REPEATABLE-READ（MySQL默认），建议改用READ-COMMITTED以降低锁竞争"}
        elif "SERIALIZABLE" in value.upper():
            return {"status": "异常", "suggestion": "事务隔离级别为SERIALIZABLE，最高级别但性能影响大，不建议使用"}
        elif "READ-UNCOMMITTED" in value.upper():
            return {"status": "异常", "suggestion": "事务隔离级别为READ-UNCOMMITTED，存在脏读风险，不建议使用"}
        return {"status": "正常", "suggestion": "配置合理"}
    
    def _analyze_innodb_lock_waits(self, item_data: list) -> Dict[str, str]:
        """分析InnoDB行锁等待次数 - 按文档三态"""
        if not item_data or len(item_data) == 0:
            return {"status": "注意", "suggestion": "未检测到数据，请检查巡检脚本配置"}
        value = self._extract_value(item_data)
        try:
            wait_count = float(value)
            if wait_count > 100:
                return {"status": "异常", "suggestion": f"InnoDB行锁等待次数过高: {wait_count}，建议检查锁等待情况并优化事务"}
            elif wait_count > 0:
                return {"status": "注意", "suggestion": f"InnoDB行锁等待次数: {wait_count}，建议持续监控"}
        except ValueError:
            pass
        return {"status": "正常", "suggestion": "InnoDB行锁等待在正常范围内"}
    
    def _analyze_innodb_lock_time_avg(self, item_data: list) -> Dict[str, str]:
        """分析InnoDB行锁平均等待时间 - 按文档三态"""
        if not item_data or len(item_data) == 0:
            return {"status": "注意", "suggestion": "未检测到数据，请检查巡检脚本配置"}
        value = self._extract_value(item_data)
        try:
            avg_time = float(value)
            if avg_time > 1000:
                return {"status": "异常", "suggestion": f"InnoDB行锁平均等待时间过长: {avg_time}ms，建议检查锁等待情况和索引优化"}
            elif avg_time > 100:
                return {"status": "注意", "suggestion": f"InnoDB行锁平均等待时间较长: {avg_time}ms，建议检查锁等待情况"}
        except ValueError:
            pass
        return {"status": "正常", "suggestion": "InnoDB行锁平均等待时间正常"}
    
    def _analyze_tmp_disk_tables(self, item_data: list) -> Dict[str, str]:
        """分析磁盘临时表使用情况"""
        if not item_data or len(item_data) == 0:
            return {"status": "注意", "suggestion": "未检测到数据，请检查巡检脚本配置"}
        value = self._extract_value(item_data)
        try:
            disk_tables = int(value)
            if disk_tables > 1000:
                return {"status": "异常", "suggestion": f"磁盘临时表创建过多: {disk_tables}，建议增加tmp_table_size和max_heap_table_size"}
            elif disk_tables > 100:
                return {"status": "注意", "suggestion": f"磁盘临时表创建较多: {disk_tables}，建议优化SQL或增加内存临时表大小"}
        except ValueError:
            pass
        return {"status": "正常", "suggestion": "临时表使用情况正常"}
    
    def _analyze_aborted_clients(self, item_data: list) -> Dict[str, str]:
        """分析异常断开连接数"""
        if not item_data or len(item_data) == 0:
            return {"status": "注意", "suggestion": "未检测到数据，请检查巡检脚本配置"}
        value = self._extract_value(item_data)
        try:
            aborted = int(value)
            if aborted > 100:
                return {"status": "异常", "suggestion": f"异常断开连接数过多: {aborted}，可能存在网络问题或应用连接未正确关闭"}
            elif aborted > 10:
                return {"status": "注意", "suggestion": f"异常断开连接数较多: {aborted}，建议检查网络和应用连接池配置"}
        except ValueError:
            pass
        return {"status": "正常", "suggestion": "异常断开连接数在正常范围内"}
    
    def _analyze_aborted_connects(self, item_data: list) -> Dict[str, str]:
        """分析异常连接尝试数"""
        if not item_data or len(item_data) == 0:
            return {"status": "注意", "suggestion": "未检测到数据，请检查巡检脚本配置"}
        value = self._extract_value(item_data)
        try:
            aborted = int(value)
            if aborted > 50:
                return {"status": "异常", "suggestion": f"异常连接尝试数过多: {aborted}，可能存在暴力破解或配置错误"}
            elif aborted > 10:
                return {"status": "注意", "suggestion": f"异常连接尝试数较多: {aborted}，建议检查连接配置和安全设置"}
        except ValueError:
            pass
        return {"status": "正常", "suggestion": "异常连接尝试数在正常范围内"}
    
    def _analyze_open_files_limit(self, item_data: list) -> Dict[str, str]:
        """分析打开文件限制 - 按文档三态"""
        if not item_data or len(item_data) == 0:
            return {"status": "注意", "suggestion": "未检测到数据，请检查巡检脚本配置"}
        value = self._extract_value(item_data)
        try:
            limit = int(value)
            if limit >= 65535:
                return {"status": "正常", "suggestion": f"打开文件限制: {limit}，配置合理"}
            elif limit >= 20000:
                return {"status": "注意", "suggestion": f"打开文件限制: {limit}，建议增加至65535"}
            else:
                return {"status": "异常", "suggestion": f"打开文件限制过低: {limit}，生产环境建议至少65535"}
        except ValueError:
            pass
        return {"status": "正常", "suggestion": "打开文件限制配置合理"}
    
    def _analyze_innodb_engine_status(self, item_data: list) -> Dict[str, str]:
        """分析InnoDB引擎状态（检查死锁）"""
        if not item_data or len(item_data) == 0:
            return {"status": "注意", "suggestion": "未检测到数据，请检查巡检脚本配置"}
        status_text = "\n".join(item_data)
        if "LATEST DETECTED DEADLOCK" in status_text:
            return {"status": "异常", "suggestion": "检测到死锁！建议检查最近的死锁信息，优化事务和索引设计"}
        return {"status": "正常", "suggestion": "未检测到死锁，InnoDB引擎状态正常"}
    
    def _analyze_error_log(self, item_data: list) -> Dict[str, str]:
        """分析MySQL错误日志"""
        if not item_data or len(item_data) == 0:
            return {"status": "注意", "suggestion": "未检测到数据，请检查巡检脚本配置"}
        log_text = "\n".join(item_data)
        error_keywords = ["ERROR", "Warning", "crash", "failed", "aborted"]
        for keyword in error_keywords:
            if keyword.lower() in log_text.lower():
                return {"status": "注意", "suggestion": "错误日志中存在异常信息，建议检查详细日志内容"}
        return {"status": "正常", "suggestion": "错误日志检查正常"}
    
    def _extract_value(self, item_data: list) -> str:
        """从数据中提取值"""
        for line in item_data:
            if "\t" in line:
                parts = line.split("\t")
                if len(parts) > 1:
                    return parts[1].strip()
            elif ":" in line:
                parts = line.split(":", 1)
                if len(parts) > 1:
                    return parts[1].strip()
        return ""
    
    def _analyze_numeric_value(self, item_data: list, item_name: str, warning_threshold: int = None, 
                               critical_threshold: int = None, unit: str = "", 
                               is_greater: bool = True) -> Dict[str, str]:
        """通用数值分析函数"""
        if not item_data or len(item_data) == 0:
            return {"status": "注意", "suggestion": "未检测到数据，请检查巡检脚本配置"}
        
        value_str = self._extract_value(item_data)
        try:
            value = int(value_str)
            status = "正常"
            suggestion = f"{item_name}: {value}{unit}"
            
            if is_greater:
                if critical_threshold and value > critical_threshold:
                    status = "异常"
                    suggestion = f"{item_name}过高: {value}{unit}，建议检查和优化"
                elif warning_threshold and value > warning_threshold:
                    status = "注意"
                    suggestion = f"{item_name}较高: {value}{unit}，建议关注"
            else:
                if critical_threshold and value < critical_threshold:
                    status = "异常"
                    suggestion = f"{item_name}过低: {value}{unit}，建议检查和优化"
                elif warning_threshold and value < warning_threshold:
                    status = "注意"
                    suggestion = f"{item_name}偏低: {value}{unit}，建议关注"
            
            return {"status": status, "suggestion": suggestion}
        except ValueError:
            return {"status": "正常", "suggestion": f"{item_name}: {value_str}"}
    
    def _analyze_config_value(self, item_data: list, item_name: str, 
                             expected_value: str = None) -> Dict[str, str]:
        """通用配置值分析函数"""
        if not item_data or len(item_data) == 0:
            return {"status": "注意", "suggestion": "未检测到数据，请检查巡检脚本配置"}
        
        value = self._extract_value(item_data)
        if expected_value and value.upper() != expected_value.upper():
            return {"status": "注意", "suggestion": f"{item_name}: {value}（预期: {expected_value}）"}
        return {"status": "正常", "suggestion": f"{item_name}: {value}"}
    
    def _analyze_database_version(self, item_data: list) -> Dict[str, str]:
        """分析数据库版本 - 按文档标准"""
        if item_data:
            version = item_data[0].strip()
            if any(v in version for v in ["8.0", "8.", "5.7"]):
                return {"status": "正常", "suggestion": f"数据库版本: {version}，在支持范围内"}
            elif any(v in version for v in ["5.6"]):
                return {"status": "注意", "suggestion": f"数据库版本: {version}（MySQL 5.6），建议升级至5.7或8.0"}
            else:
                return {"status": "异常", "suggestion": f"数据库版本: {version}（MySQL 5.5及以下），请立即升级"}
        return {"status": "注意", "suggestion": "未检测到数据库版本信息"}
    
    def _analyze_os_version(self, item_data: list) -> Dict[str, str]:
        """分析操作系统版本 - 按文档标准"""
        if item_data:
            os_info = item_data[0].strip().lower()
            if any(v in os_info for v in ["centos 7", "centos 8", "centos 9", "ubuntu 20", "ubuntu 22", "ubuntu 24", "rhel 7", "rhel 8", "rhel 9"]):
                return {"status": "正常", "suggestion": f"操作系统: {item_data[0].strip()}，在支持范围内"}
            elif any(v in os_info for v in ["centos 6", "ubuntu 18", "rhel 6"]):
                return {"status": "注意", "suggestion": f"操作系统: {item_data[0].strip()}，建议升级到受支持版本"}
            else:
                return {"status": "异常", "suggestion": f"操作系统: {item_data[0].strip()}，版本过旧，请尽快升级"}
        return {"status": "注意", "suggestion": "未检测到操作系统版本信息"}
    
    def _analyze_master_status(self, item_data: list) -> Dict[str, str]:
        """分析主库状态 - 按文档标准"""
        if item_data and len(item_data) > 0:
            return {"status": "正常", "suggestion": "主库运行正常"}
        return {"status": "异常", "suggestion": "主库状态异常，请检查数据库是否正常运行"}
    
    def _analyze_innodb_file_per_table(self, item_data: list) -> Dict[str, str]:
        """分析独立表空间 - 按文档标准"""
        value = self._extract_value(item_data) if item_data else ""
        if value.upper() == "ON":
            return {"status": "正常", "suggestion": "独立表空间已开启(ON)，建议保持开启"}
        return {"status": "异常", "suggestion": "独立表空间未开启(OFF)，建议开启以提高管理灵活性"}
    
    def _analyze_innodb_flush_method(self, item_data: list) -> Dict[str, str]:
        """分析InnoDB刷新方法 - 按文档标准"""
        value = self._extract_value(item_data) if item_data else ""
        if value.upper() == "O_DIRECT":
            return {"status": "正常", "suggestion": f"InnoDB刷新方法: {value}，配置推荐"}
        return {"status": "正常", "suggestion": f"InnoDB刷新方法: {value}"}
    
    def _analyze_innodb_buffer_pool(self, item_data: list) -> Dict[str, str]:
        """分析InnoDB缓冲池大小 - 按文档三态"""
        value_str = self._extract_value(item_data) if item_data else ""
        try:
            pool_bytes = int(value_str)
            pool_gb = pool_bytes / (1024**3)
            return {"status": "正常", "suggestion": f"InnoDB缓冲池大小: {pool_gb:.2f}GB，请根据内存总量判断是否达到50%以上"}
        except (ValueError, TypeError):
            return {"status": "正常", "suggestion": f"InnoDB缓冲池大小: {value_str}"}
    
    def _analyze_max_connections(self, item_data: list) -> Dict[str, str]:
        """分析最大连接数 - 按文档三态"""
        value = self._extract_value(item_data) if item_data else ""
        try:
            conn = int(value)
            if conn >= 200:
                return {"status": "正常", "suggestion": f"最大连接数: {conn}，配置合理"}
            elif conn >= 100:
                return {"status": "注意", "suggestion": f"最大连接数: {conn}，建议增加至200以上"}
            else:
                return {"status": "异常", "suggestion": f"最大连接数: {conn}，过低，请立即增加"}
        except ValueError:
            return {"status": "正常", "suggestion": f"最大连接数: {value}"}
    
    def _analyze_table_open_cache(self, item_data: list) -> Dict[str, str]:
        """分析表打开缓存 - 按文档三态"""
        value = self._extract_value(item_data) if item_data else ""
        try:
            cache = int(value)
            if cache >= 4000:
                return {"status": "正常", "suggestion": f"表打开缓存: {cache}，配置合理"}
            elif cache >= 2000:
                return {"status": "注意", "suggestion": f"表打开缓存: {cache}，建议增加至4000以上"}
            else:
                return {"status": "异常", "suggestion": f"表打开缓存: {cache}，过低，建议增加"}
        except ValueError:
            return {"status": "正常", "suggestion": f"表打开缓存: {value}"}
    
    def _analyze_max_allowed_packet(self, item_data: list) -> Dict[str, str]:
        """分析最大允许包大小 - 按文档三态"""
        value = self._extract_value(item_data) if item_data else ""
        try:
            bytes_val = int(value)
            mb_val = bytes_val / (1024*1024)
            if mb_val >= 64:
                return {"status": "正常", "suggestion": f"最大包大小: {mb_val:.0f}MB，配置合理"}
            elif mb_val >= 16:
                return {"status": "注意", "suggestion": f"最大包大小: {mb_val:.0f}MB，建议增加至64MB以上"}
            else:
                return {"status": "异常", "suggestion": f"最大包大小: {mb_val:.0f}MB，过低，建议增加"}
        except ValueError:
            return {"status": "正常", "suggestion": f"最大包大小: {value}"}
    
    def _analyze_wait_timeout(self, item_data: list, is_interactive: bool = False) -> Dict[str, str]:
        """分析超时设置 - 按文档三态"""
        name = "交互超时" if is_interactive else "等待超时"
        value = self._extract_value(item_data) if item_data else ""
        try:
            timeout = int(value)
            if timeout >= 28800:
                return {"status": "正常", "suggestion": f"{name}: {timeout}秒(8小时)，配置合理"}
            elif timeout >= 600:
                return {"status": "注意", "suggestion": f"{name}: {timeout}秒，建议设置为28800(8小时)"}
            else:
                return {"status": "异常", "suggestion": f"{name}: {timeout}秒，过低，建议设置为28800(8小时)"}
        except ValueError:
            return {"status": "正常", "suggestion": f"{name}: {value}"}
    
    def _analyze_innodb_log_file_size(self, item_data: list) -> Dict[str, str]:
        """分析InnoDB日志文件大小 - 按文档三态（byte→GB）"""
        if not item_data or len(item_data) == 0:
            return {"status": "注意", "suggestion": "未检测到数据，请检查巡检脚本配置"}
        value_str = self._extract_value(item_data)
        try:
            size_bytes = int(value_str)
            size_gb = size_bytes / (1024 * 1024 * 1024)
            if size_gb >= 2:
                return {"status": "正常", "suggestion": f"InnoDB日志文件大小: {size_gb:.2f}GB，配置合理"}
            elif size_gb >= 1:
                return {"status": "注意", "suggestion": f"InnoDB日志文件大小: {size_gb:.2f}GB，建议增大至2GB以上以减少日志切换频率"}
            else:
                return {"status": "异常", "suggestion": f"InnoDB日志文件大小: {size_gb:.2f}GB，过小，建议增大至2GB以上"}
        except ValueError:
            return {"status": "正常", "suggestion": f"InnoDB日志文件大小: {value_str}"}
    
    def _analyze_file_descriptor_limit(self, item_data: list) -> Dict[str, str]:
        """分析文件描述符限制 - 按文档三态（处理/proc格式）"""
        if not item_data or len(item_data) == 0:
            return {"status": "注意", "suggestion": "未检测到数据，请检查巡检脚本配置"}
        limit = None
        for line in item_data:
            line = line.strip()
            if "open files" in line.lower() or "max open files" in line.lower():
                parts = line.split()
                if len(parts) >= 4:
                    try:
                        limit = int(parts[3])
                    except ValueError:
                        try:
                            limit = int(parts[2])
                        except ValueError:
                            pass
                    break
        if limit is None:
            return {"status": "注意", "suggestion": "无法解析文件描述符限制值"}
        if limit >= 65535:
            return {"status": "正常", "suggestion": f"文件描述符限制: {limit}，配置合理"}
        elif limit >= 20000:
            return {"status": "注意", "suggestion": f"文件描述符限制: {limit}，建议增加至65535"}
        else:
            return {"status": "异常", "suggestion": f"文件描述符限制: {limit}，过低，建议增加至65535"}
    
    def _generate_recommendations(self, doc: Document, issues_summary: Dict[str, Any]):
        """生成处理建议 - 表格样式"""
        doc.add_page_break()
        self._add_heading_with_style(doc, self._clean_xml_text("五、处理建议"), level=1)
        
        critical_issues = issues_summary.get('critical_issues', [])
        warning_issues = issues_summary.get('warning_issues', [])
        info_issues = issues_summary.get('info_issues', [])
        
        all_with_suggestions = [
            ("P0 严重问题 — 需立即处理", critical_issues, 'FFF0F0', 'FFD7D7', 'CC0000'),
            ("P1 警告问题 — 建议近期处理", warning_issues, 'FFF8E6', 'FFEDB8', 'CC7700'),
        ]
        
        has_issues = False
        for title, issues, header_bg, row_bg, color_code in all_with_suggestions:
            if not issues:
                continue
            has_issues = True
            
            heading = doc.add_paragraph(self._clean_xml_text(title))
            heading_run = heading.add_run()
            heading_run.font.bold = True
            heading_run.font.size = Pt(13)
            heading_run.font.color.rgb = RGBColor(int(color_code[:2], 16), int(color_code[2:4], 16), int(color_code[4:], 16))
            
            rec_table = doc.add_table(rows=1, cols=3)
            rec_table.style = 'Light Grid Accent 1'
            rec_table.autofit = False
            rec_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            rec_table.columns[0].width = Inches(1.0)
            rec_table.columns[1].width = Inches(1.6)
            rec_table.columns[2].width = Inches(3.6)
            
            rec_header = rec_table.rows[0].cells
            for j, (cell, text) in enumerate(zip(rec_header, ['问题分类', '问题描述', '处理建议'])):
                cell.text = self._clean_xml_text(text)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._setup_cell_style(cell, bold=True, font_size=10, color=RGBColor(255, 255, 255))
                self._set_cell_shading(cell, "003366")
            
            for i, issue in enumerate(issues):
                row_cells = rec_table.add_row().cells
                row_cells[0].text = self._clean_xml_text(issue.get('category', ''))
                self._setup_cell_style(row_cells[0], bold=True, color=RGBColor(0, 0, 0))
                
                row_cells[1].text = self._clean_xml_text(issue.get('issue', ''))
                self._setup_cell_style(row_cells[1], color=RGBColor(51, 51, 51))
                
                row_cells[2].text = self._clean_xml_text(issue.get('suggestion', ''))
                self._setup_cell_style(row_cells[2], color=RGBColor(102, 102, 102))
                
                shading_color = row_bg if i % 2 == 0 else header_bg
                for cell in row_cells:
                    self._set_cell_shading(cell, shading_color)
                    cell.paragraphs[0].space_before = Pt(2)
                    cell.paragraphs[0].space_after = Pt(2)
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
        
        # 常规维护建议
        maintenance_title = doc.add_paragraph(self._clean_xml_text("常规维护建议"))
        maintenance_run = maintenance_title.add_run()
        maintenance_run.font.bold = True
        maintenance_run.font.size = Pt(12)
        maintenance_run.font.color.rgb = RGBColor(0, 102, 0)
        
        maintenance_table = doc.add_table(rows=1, cols=2)
        maintenance_table.style = 'Light Grid Accent 1'
        maintenance_table.autofit = False
        maintenance_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        maintenance_table.columns[0].width = Inches(0.4)
        maintenance_table.columns[1].width = Inches(5.8)
        
        maint_header = maintenance_table.rows[0].cells
        for j, (cell, text) in enumerate(zip(maint_header, ['序号', '建议内容'])):
            cell.text = self._clean_xml_text(text)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._setup_cell_style(cell, bold=True, font_size=10, color=RGBColor(255, 255, 255))
            self._set_cell_shading(cell, "003366")
        
        maintenance_items = [
            "定期备份数据库，确保数据安全",
            "监控数据库性能指标，及时发现异常",
            "及时更新数据库补丁和版本",
            "定期优化表和索引，减少碎片",
            "建立完善的监控告警体系"
        ]
        
        for i, item in enumerate(maintenance_items, 1):
            row_cells = maintenance_table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._setup_cell_style(row_cells[0], color=RGBColor(0, 0, 0))
            
            row_cells[1].text = self._clean_xml_text(item)
            self._setup_cell_style(row_cells[1], color=RGBColor(51, 51, 51))
            
            shading_color = "F0FFF0" if i % 2 == 0 else "E0F0E0"
            for cell in row_cells:
                self._set_cell_shading(cell, shading_color)
                cell.paragraphs[0].space_before = Pt(3)
                cell.paragraphs[0].space_after = Pt(3)
        
        if not has_issues:
            para = doc.add_paragraph("数据库状态良好，建议继续保持定期巡检。")
    
    def _save_docx_report(self, doc: Document, inspection_data: Dict[str, Any]) -> str:
        """保存Word报告文件"""
        output_dir = "reports"
        os.makedirs(output_dir, exist_ok=True)
        
        # 从日志文件名中提取IP和端口
        log_filename = os.path.basename(self.log_file)
        instance_ip = ""
        instance_port = ""
        
        # 尝试从日志文件名中提取IP和端口
        # 日志文件名格式: mysql_inspection_<IP>_<PORT>_<date>.log
        import re
        # 匹配 IP 地址格式
        ip_pattern = r'(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})'
        # 匹配 端口号格式
        port_pattern = r'_(\d{1,5})_'
        
        ip_match = re.search(ip_pattern, log_filename)
        if ip_match:
            instance_ip = ip_match.group(1)
        
        port_match = re.search(port_pattern, log_filename)
        if port_match:
            instance_port = port_match.group(1)
        
        # 如果从日志文件名中提取失败，再尝试从巡检数据中提取
        if not instance_ip or not instance_port:
            summary_data = inspection_data.get('INSPECTION_SUMMARY', [])
            for line in summary_data:
                if not instance_ip and "实例IP:" in line:
                    instance_ip = line.replace("实例IP:", "").strip()
                elif not instance_port and "实例端口:" in line:
                    instance_port = line.replace("实例端口:", "").strip()
        
        # 生成文件名
        current_date = datetime.now().strftime('%Y%m%d')
        
        if instance_ip and instance_port:
            filename = f"MySQL_inspection_{instance_ip}_{instance_port}_{current_date}.docx"
        else:
            # 如果没有IP和端口，使用默认格式
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"MySQL_inspection_{timestamp}.docx"
        
        filepath = os.path.join(output_dir, filename)
        
        doc.save(filepath)
        return filepath

def main():
    """主函数"""
    import sys
    
    # 解析 --font 参数
    font_name = "微软雅黑"
    clean_argv = []
    skip_next = False
    for i, arg in enumerate(sys.argv):
        if skip_next:
            skip_next = False
            continue
        if arg == '--font' and i + 1 < len(sys.argv):
            font_name = sys.argv[i + 1]
            skip_next = True
        elif arg.startswith('--font='):
            font_name = arg.split('=', 1)[1]
        else:
            clean_argv.append(arg)
    global FONT_NAME
    FONT_NAME = font_name
    sys.argv = clean_argv
    
    if len(sys.argv) < 2:
        print("用法: python log_analyzer.py <巡检日志文件> [输出格式]")
        print("输出格式: docx (默认) 或 txt")
        print("")
        print("示例:")
        print("  python log_analyzer.py inspection_logs/mysql_inspection_20241201_120000.log")
        print("  python log_analyzer.py inspection_logs/mysql_inspection_20241201_120000.log txt")
        sys.exit(1)
    
    log_file = sys.argv[1]
    output_format = sys.argv[2] if len(sys.argv) > 2 else "docx"
    
    try:
        analyzer = LogAnalyzer(log_file)
        report_file = analyzer.generate_report(output_format)
        
        print(f"巡检报告生成完成: {report_file}")
        
    except Exception as e:
        print(f"生成报告失败: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()