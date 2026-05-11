#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
批量日志分析脚本 - 并行处理优化版
用于批量处理多个巡检日志文件，支持并行分析和进度显示
"""

import os
import sys
import re
import glob
from datetime import datetime
from log_analyzer import LogAnalyzer
from concurrent.futures import ProcessPoolExecutor, as_completed
from multiprocessing import cpu_count

try:
    from docx import Document
    from docx.shared import Inches, Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# 全局字体设置
FONT_NAME = "微软雅黑"


try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False


def analyze_single_log(log_file: str, batch_dir: str, output_format: str):
    """分析单个日志文件的函数（用于并行处理）
    
    Args:
        log_file: 日志文件路径
        batch_dir: 批量报告目录
        output_format: 输出格式
    
    Returns:
        dict: 分析结果，包含成功状态、错误信息、巡检数据等
    """
    try:
        analyzer = LogAnalyzer(log_file)
        inspection_data = analyzer.parse_log_file()
        issues_summary = analyzer.analyze_issues()
        
        # 生成报告
        report_file = analyzer.generate_report(output_format)
        
        # 移动报告到批量目录
        report_filename = os.path.basename(report_file)
        new_report_path = os.path.join(batch_dir, report_filename)
        os.rename(report_file, new_report_path)
        
        return {
            'success': True,
            'log_file': log_file,
            'inspection_data': inspection_data,
            'issues_summary': issues_summary,
            'report_path': new_report_path
        }
    except Exception as e:
        return {
            'success': False,
            'log_file': log_file,
            'error': str(e)
        }


def update_progress(current: int, total: int, prefix: str = "进度"):
    """显示进度百分比
    
    Args:
        current: 当前完成数
        total: 总数
        prefix: 进度条前缀文本
    """
    percent = (current / total) * 100
    bar_length = 50
    filled_length = int(bar_length * current // total)
    bar = '█' * filled_length + '-' * (bar_length - filled_length)
    
    # 使用回车符回到行首
    sys.stdout.write(f'\r{prefix}: |{bar}| {percent:.1f}% ({current}/{total})')
    sys.stdout.flush()


def batch_analyze_logs(logs_directory: str, output_format: str = "docx", max_workers: int = None, generate_excel: bool = False):
    """批量分析日志文件 - 并行处理版
    
    Args:
        logs_directory: 日志目录
        output_format: 输出格式
        max_workers: 最大并行工作进程数，默认为 CPU 核心数
        generate_excel: 是否额外生成Excel统计表格
    """
    
    # 查找所有日志文件
    log_files = glob.glob(os.path.join(logs_directory, "*.log"))
    
    if not log_files:
        print(f"在目录 {logs_directory} 中未找到日志文件")
        return
    
    total_files = len(log_files)
    print(f"找到 {total_files} 个日志文件")
    
    # 设置最大工作进程数
    if max_workers is None:
        max_workers = max(1, cpu_count() - 1)  # 留一个核心给系统
    max_workers = min(max_workers, total_files)  # 不超过文件数
    print(f"使用 {max_workers} 个并行进程进行分析")
    
    # 创建批量报告目录
    batch_dir = f"batch_reports_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    os.makedirs(batch_dir, exist_ok=True)
    
    success_count = 0
    failed_files = []
    all_inspection_data_list = []  # 存储所有实例的巡检数据用于汇总
    completed_count = 0
    
    print("\n开始并行分析...")
    update_progress(0, total_files)
    
    # 并行处理
    with ProcessPoolExecutor(max_workers=max_workers) as executor:
        # 提交所有任务
        future_to_log = {
            executor.submit(analyze_single_log, log_file, batch_dir, output_format): log_file
            for log_file in log_files
        }
        
        # 处理完成的任务
        for future in as_completed(future_to_log):
            result = future.result()
            completed_count += 1
            
            if result['success']:
                print(f"\n✓ 分析成功: {os.path.basename(result['log_file'])}")
                success_count += 1
                all_inspection_data_list.append({
                    'log_file': result['log_file'],
                    'inspection_data': result['inspection_data'],
                    'issues_summary': result['issues_summary']
                })
            else:
                print(f"\n✗ 分析失败: {os.path.basename(result['log_file'])} - {result['error']}")
                failed_files.append((os.path.basename(result['log_file']), result['error']))
            
            # 更新进度
            update_progress(completed_count, total_files)
    
    # 确保进度条完整显示
    update_progress(total_files, total_files)
    print("\n")  # 换行
    
    # 生成批量分析摘要
    generate_batch_summary(batch_dir, log_files, success_count, failed_files)
    
    # 生成批量汇总报告
    if DOCX_AVAILABLE and success_count > 0:
        generate_batch_summary_report(batch_dir, all_inspection_data_list)
    
    # 生成Excel统计表格
    if generate_excel and success_count > 0:
        if EXCEL_AVAILABLE:
            generate_batch_statistics_excel(batch_dir, all_inspection_data_list)
        else:
            print("\n⚠ openpyxl 未安装，无法生成Excel统计表格")
            print("  请执行: pip install openpyxl")
    
    print(f"\n{'='*50}")
    print(f"批量分析完成")
    print(f"成功: {success_count} 个文件")
    print(f"失败: {len(failed_files)} 个文件")
    print(f"报告目录: {batch_dir}")
    print(f"{'='*50}")
    
    if failed_files:
        print("\n失败文件列表:")
        for filename, error in failed_files:
            print(f"  {filename}: {error}")


def _enable_auto_update_fields(doc: Document):
    """设置文档打开时自动更新所有域（包括目录）"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    settings = doc.settings._element
    update_fields = settings.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        settings.append(update_fields)
    update_fields.set(qn('w:val'), 'true')


def generate_batch_summary_report(batch_dir: str, inspection_data_list: list):
    """生成批量汇总报告"""
    try:
        doc = Document()
        
        # 设置文档打开时自动更新域（目录自动生成）
        _enable_auto_update_fields(doc)
        
        # 设置文档样式
        setup_document_styles(doc)
        
        # 生成封面页
        generate_cover_page(doc)
        
        # 生成目录
        generate_table_of_contents(doc)
        
        # 生成汇总内容
        generate_summary_content(doc, inspection_data_list)
        
        # 生成详细内容（按IP+端口组织）
        generate_detailed_content_by_instance(doc, inspection_data_list)
        
        # 保存报告
        summary_report_filename = os.path.join(batch_dir, "批量巡检汇总报告.docx")
        doc.save(summary_report_filename)
        
        print(f"\n✓ 批量汇总报告生成成功: {summary_report_filename}")
        
    except Exception as e:
        print(f"✗ 生成批量汇总报告失败: {str(e)}")
        import traceback
        print(f"详细错误: {traceback.format_exc()}")


def generate_batch_statistics_excel(batch_dir: str, inspection_data_list: list):
    """生成Excel统计表格
    
    包含IP、端口、P0/P1/P2问题数量统计及问题描述
    """
    try:
        wb = openpyxl.Workbook()
        
        # ============================================================
        # Sheet 1: 问题统计汇总
        # ============================================================
        ws_summary = wb.active
        ws_summary.title = "问题统计汇总"
        
        # 定义样式
        title_font = Font(name=FONT_NAME, size=16, bold=True, color='FFFFFF')
        header_font = Font(name=FONT_NAME, size=11, bold=True, color='FFFFFF')
        body_font = Font(name=FONT_NAME, size=10)
        p0_font = Font(name=FONT_NAME, size=10, bold=True, color='CC0000')
        p1_font = Font(name=FONT_NAME, size=10, bold=True, color='E68A00')
        p2_font = Font(name=FONT_NAME, size=10, color='006600')
        
        title_fill = PatternFill(start_color='003366', end_color='003366', fill_type='solid')
        header_fill = PatternFill(start_color='004080', end_color='004080', fill_type='solid')
        p0_fill = PatternFill(start_color='FFE6E6', end_color='FFE6E6', fill_type='solid')
        p1_fill = PatternFill(start_color='FFF2D6', end_color='FFF2D6', fill_type='solid')
        p2_fill = PatternFill(start_color='E6F5E6', end_color='E6F5E6', fill_type='solid')
        alt_fill = PatternFill(start_color='F5F8FC', end_color='F5F8FC', fill_type='solid')
        
        center_align = Alignment(horizontal='center', vertical='center', wrap_text=True)
        left_align = Alignment(horizontal='left', vertical='center', wrap_text=True)
        
        thin_border = Border(
            left=Side(style='thin', color='B0B0B0'),
            right=Side(style='thin', color='B0B0B0'),
            top=Side(style='thin', color='B0B0B0'),
            bottom=Side(style='thin', color='B0B0B0')
        )
        
        # 标题行
        ws_summary.merge_cells('A1:I1')
        title_cell = ws_summary['A1']
        title_cell.value = f"MySQL巡检批量统计报表（{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}）"
        title_cell.font = title_font
        title_cell.fill = title_fill
        title_cell.alignment = center_align
        ws_summary.row_dimensions[1].height = 40
        
        # 表头
        headers = ['序号', 'IP地址', '端口', 'P0严重问题', 'P0问题描述', 'P1警告问题', 'P1问题描述', 'P2信息提示', 'P2问题描述']
        col_widths = [6, 18, 10, 14, 45, 14, 45, 14, 45]
        
        for col_idx, (header, width) in enumerate(zip(headers, col_widths), 1):
            cell = ws_summary.cell(row=2, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = thin_border
            ws_summary.column_dimensions[get_column_letter(col_idx)].width = width
        
        ws_summary.row_dimensions[2].height = 30
        ws_summary.auto_filter.ref = f'A2:I{len(inspection_data_list) + 2}'
        
        # 填充数据
        row_idx = 3
        total_p0 = 0
        total_p1 = 0
        total_p2 = 0
        
        for idx, data in enumerate(inspection_data_list, 1):
            inspection_data = data.get('inspection_data', {})
            issues_summary = data.get('issues_summary', {})
            
            # 获取IP和端口
            summary_section = inspection_data.get("INSPECTION_SUMMARY", [])
            instance_ip = ""
            instance_port = ""
            
            # 先尝试从日志文件名提取
            log_filename = os.path.basename(data.get('log_file', ''))
            ip_match = re.search(r'(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})', log_filename)
            port_match = re.search(r'_(\d{1,5})_', log_filename)
            if ip_match:
                instance_ip = ip_match.group(1)
            if port_match:
                instance_port = port_match.group(1)
            
            # 如果文件名提取失败，从巡检数据提取
            if not instance_ip or not instance_port:
                for line in summary_section:
                    if not instance_ip and "实例IP:" in line:
                        instance_ip = line.replace("实例IP:", "").strip()
                    elif not instance_port and "实例端口:" in line:
                        instance_port = line.replace("实例端口:", "").strip()
            
            critical_issues = issues_summary.get('critical_issues', [])
            warning_issues = issues_summary.get('warning_issues', [])
            info_issues = issues_summary.get('info_issues', [])
            
            p0_count = len(critical_issues)
            p1_count = len(warning_issues)
            p2_count = len(info_issues)
            
            total_p0 += p0_count
            total_p1 += p1_count
            total_p2 += p2_count
            
            # 构建问题描述（全部展示，无缩略）
            p0_desc = '\n'.join([f"【{i['category']}】{i['issue']}" for i in critical_issues]) if critical_issues else '无'
            p1_desc = '\n'.join([f"【{i['category']}】{i['issue']}" for i in warning_issues]) if warning_issues else '无'
            p2_desc = '\n'.join([f"【{i['category']}】{i['issue']}" for i in info_issues]) if info_issues else '无'
            
            row_data = [
                idx,
                instance_ip or '未知',
                instance_port or '未知',
                p0_count,
                p0_desc or '无',
                p1_count,
                p1_desc or '无',
                p2_count,
                p2_desc or '无'
            ]
            
            is_alt_row = (idx % 2 == 0)
            
            for col_idx, value in enumerate(row_data, 1):
                cell = ws_summary.cell(row=row_idx, column=col_idx, value=value)
                cell.font = body_font
                cell.border = thin_border
                
                if col_idx <= 3 or col_idx in (4, 6, 8):
                    cell.alignment = center_align
                else:
                    cell.alignment = left_align
                
                if is_alt_row:
                    cell.fill = alt_fill
            
            # 对P0/P1/P2数量列着色
            p0_cell = ws_summary.cell(row=row_idx, column=4)
            if p0_count > 0:
                p0_cell.font = p0_font
                p0_cell.fill = p0_fill
            
            p1_cell = ws_summary.cell(row=row_idx, column=6)
            if p1_count > 0:
                p1_cell.font = p1_font
                p1_cell.fill = p1_fill
            
            p2_cell = ws_summary.cell(row=row_idx, column=8)
            if p2_count > 0:
                p2_cell.font = p2_font
                p2_cell.fill = p2_fill
            
            p2_desc_cell = ws_summary.cell(row=row_idx, column=9)
            if p2_count > 0:
                p2_desc_cell.font = p2_font
                p2_desc_cell.fill = p2_fill
            
            ws_summary.row_dimensions[row_idx].height = max(45, 17 * max(
                1,
                row_data[4].count('\n') + 1 if row_data[4] != '无' else 1,
                row_data[6].count('\n') + 1 if row_data[6] != '无' else 1,
                row_data[8].count('\n') + 1 if row_data[8] != '无' else 1
            ))
            for desc_col in (5, 7, 9):
                desc_cell = ws_summary.cell(row=row_idx, column=desc_col)
                desc_cell.alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
            row_idx += 1
        
        # 汇总行
        summary_row = row_idx
        ws_summary.merge_cells(f'A{summary_row}:C{summary_row}')
        total_label = ws_summary.cell(row=summary_row, column=1, value='合计')
        total_label.font = Font(name=FONT_NAME, size=11, bold=True)
        total_label.fill = PatternFill(start_color='E0E8F0', end_color='E0E8F0', fill_type='solid')
        total_label.alignment = center_align
        total_label.border = thin_border
        
        for col in range(2, 4):
            cell = ws_summary.cell(row=summary_row, column=col)
            cell.fill = PatternFill(start_color='E0E8F0', end_color='E0E8F0', fill_type='solid')
            cell.border = thin_border
        
        total_values = [total_p0, '', total_p1, '', total_p2, '']
        total_cols = [4, 5, 6, 7, 8, 9]
        for col_idx, val in zip(total_cols, total_values):
            cell = ws_summary.cell(row=summary_row, column=col_idx, value=val)
            cell.font = Font(name=FONT_NAME, size=11, bold=True)
            cell.fill = PatternFill(start_color='E0E8F0', end_color='E0E8F0', fill_type='solid')
            cell.alignment = center_align if col_idx in (4, 6, 8) else left_align
            cell.border = thin_border
        
        ws_summary.row_dimensions[summary_row].height = 28
        
        # ============================================================
        # Sheet 2: 问题详情
        # ============================================================
        ws_detail = wb.create_sheet(title="问题详情")
        
        detail_headers = ['序号', 'IP地址', '端口', '问题级别', '分类', '问题描述', '优化建议']
        detail_col_widths = [6, 18, 10, 12, 12, 40, 50]
        
        # 标题行
        ws_detail.merge_cells('A1:G1')
        detail_title = ws_detail['A1']
        detail_title.value = f"MySQL巡检问题详情（{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}）"
        detail_title.font = title_font
        detail_title.fill = title_fill
        detail_title.alignment = center_align
        ws_detail.row_dimensions[1].height = 40
        
        for col_idx, (header, width) in enumerate(zip(detail_headers, detail_col_widths), 1):
            cell = ws_detail.cell(row=2, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = thin_border
            ws_detail.column_dimensions[get_column_letter(col_idx)].width = width
        
        ws_detail.row_dimensions[2].height = 30
        ws_detail.auto_filter.ref = f'A2:G{len(inspection_data_list) * 20 + 2}'
        
        # 填充详细数据
        detail_row = 3
        for idx, data in enumerate(inspection_data_list, 1):
            inspection_data = data.get('inspection_data', {})
            issues_summary = data.get('issues_summary', {})
            
            # 获取IP和端口（复用上面逻辑）
            summary_section = inspection_data.get("INSPECTION_SUMMARY", [])
            instance_ip = ""
            instance_port = ""
            log_filename = os.path.basename(data.get('log_file', ''))
            ip_match = re.search(r'(\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3})', log_filename)
            port_match = re.search(r'_(\d{1,5})_', log_filename)
            if ip_match:
                instance_ip = ip_match.group(1)
            if port_match:
                instance_port = port_match.group(1)
            if not instance_ip or not instance_port:
                for line in summary_section:
                    if not instance_ip and "实例IP:" in line:
                        instance_ip = line.replace("实例IP:", "").strip()
                    elif not instance_port and "实例端口:" in line:
                        instance_port = line.replace("实例端口:", "").strip()
            
            ip_val = instance_ip or '未知'
            port_val = instance_port or '未知'
            
            # P0级别
            for issue in issues_summary.get('critical_issues', []):
                row_data = [idx, ip_val, port_val, 'P0严重', issue.get('category', ''), issue.get('issue', ''), issue.get('suggestion', '')]
                for col_idx, value in enumerate(row_data, 1):
                    cell = ws_detail.cell(row=detail_row, column=col_idx, value=value)
                    cell.font = Font(name=FONT_NAME, size=10)
                    cell.border = thin_border
                    cell.alignment = left_align if col_idx >= 5 else center_align
                    if col_idx == 4:
                        cell.font = Font(name=FONT_NAME, size=10, bold=True, color='CC0000')
                ws_detail.row_dimensions[detail_row].height = 28
                detail_row += 1
            
            # P1级别
            for issue in issues_summary.get('warning_issues', []):
                row_data = [idx, ip_val, port_val, 'P1警告', issue.get('category', ''), issue.get('issue', ''), issue.get('suggestion', '')]
                for col_idx, value in enumerate(row_data, 1):
                    cell = ws_detail.cell(row=detail_row, column=col_idx, value=value)
                    cell.font = Font(name=FONT_NAME, size=10)
                    cell.border = thin_border
                    cell.alignment = left_align if col_idx >= 5 else center_align
                    if col_idx == 4:
                        cell.font = Font(name=FONT_NAME, size=10, bold=True, color='E68A00')
                ws_detail.row_dimensions[detail_row].height = 28
                detail_row += 1
            
            # P2级别
            for issue in issues_summary.get('info_issues', []):
                row_data = [idx, ip_val, port_val, 'P2提示', issue.get('category', ''), issue.get('issue', ''), issue.get('suggestion', '')]
                for col_idx, value in enumerate(row_data, 1):
                    cell = ws_detail.cell(row=detail_row, column=col_idx, value=value)
                    cell.font = Font(name=FONT_NAME, size=10)
                    cell.border = thin_border
                    cell.alignment = left_align if col_idx >= 5 else center_align
                    if col_idx == 4:
                        cell.font = Font(name=FONT_NAME, size=10, color='006600')
                ws_detail.row_dimensions[detail_row].height = 28
                detail_row += 1
        
        # 保存
        excel_filename = os.path.join(batch_dir, f"巡检统计报表_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
        wb.save(excel_filename)
        print(f"\n✓ Excel统计表格生成成功: {excel_filename}")
        
    except Exception as e:
        print(f"✗ 生成Excel统计表格失败: {str(e)}")
        import traceback
        print(f"详细错误: {traceback.format_exc()}")


def setup_document_styles(doc: Document):
    """设置文档样式"""
    from datetime import datetime
    
    # 设置页面边距和页眉页脚
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        
        # 添加页眉
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_para.add_run("MySQL数据库巡检汇总报告")
        header_run.font.size = Pt(10)
        header_run.font.name = FONT_NAME
        header_run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        header_run.font.color.rgb = RGBColor(102, 102, 102)
        
        # 添加页脚
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        footer_run1 = footer_para.add_run("专业数据库服务提供商")
        footer_run1.font.size = Pt(9)
        footer_run1.font.name = FONT_NAME
        footer_run1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        footer_run1.font.color.rgb = RGBColor(102, 102, 102)
        
        footer_para.add_run(" | ")
        footer_run2 = footer_para.add_run("第 ")
        footer_run2.font.size = Pt(9)
        footer_run2.font.color.rgb = RGBColor(102, 102, 102)
        
        footer_run3 = footer_para.add_run()
        footer_run3._element.append(parse_xml('<w:fldSimple {} w:instr="PAGE \\* MERGEFORMAT"/>'.format(nsdecls('w'))))
        footer_run3.font.size = Pt(9)
        footer_run3.font.color.rgb = RGBColor(102, 102, 102)
        
        footer_para.add_run(" 页")
        
        footer_para.add_run(" | ")
        footer_run5 = footer_para.add_run(datetime.now().strftime('%Y-%m-%d'))
        footer_run5.font.size = Pt(9)
        footer_run5.font.color.rgb = RGBColor(102, 102, 102)
    
    # 设置专业字体
    doc.styles['Normal'].font.name = FONT_NAME
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    
    # 创建专业标题样式
    if 'ProfessionalTitle' not in doc.styles:
        title_style = doc.styles.add_style('ProfessionalTitle', 1)
        title_style.font.name = FONT_NAME
        title_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        title_style.font.size = Pt(28)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 0, 0)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(24)
    
    # 创建章节标题样式
    if 'SectionHeading' not in doc.styles:
        heading_style = doc.styles.add_style('SectionHeading', 1)
        heading_style.font.name = FONT_NAME
        heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        heading_style.font.size = Pt(18)
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(12)
    
    # 创建子标题样式
    if 'SubHeading' not in doc.styles:
        subheading_style = doc.styles.add_style('SubHeading', 1)
        subheading_style.font.name = FONT_NAME
        subheading_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        subheading_style.font.size = Pt(14)
        subheading_style.font.bold = True
        subheading_style.font.color.rgb = RGBColor(0, 0, 0)
        subheading_style.paragraph_format.space_before = Pt(12)
        subheading_style.paragraph_format.space_after = Pt(6)
    
    # 设置 Word 内置标题样式 - Heading 1 (一级标题)
    heading1_style = doc.styles['Heading 1']
    heading1_style.font.name = FONT_NAME
    heading1_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True
    heading1_style.font.color.rgb = RGBColor(0, 0, 0)
    heading1_style.paragraph_format.space_before = Pt(18)
    heading1_style.paragraph_format.space_after = Pt(12)
    
    # 设置 Word 内置标题样式 - Heading 2 (二级标题)
    heading2_style = doc.styles['Heading 2']
    heading2_style.font.name = FONT_NAME
    heading2_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    heading2_style.font.size = Pt(16)
    heading2_style.font.bold = True
    heading2_style.font.color.rgb = RGBColor(0, 0, 0)
    heading2_style.paragraph_format.space_before = Pt(12)
    heading2_style.paragraph_format.space_after = Pt(8)
    
    # 设置 Word 内置标题样式 - Heading 3 (三级标题)
    heading3_style = doc.styles['Heading 3']
    heading3_style.font.name = FONT_NAME
    heading3_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    heading3_style.font.size = Pt(14)
    heading3_style.font.bold = True
    heading3_style.font.color.rgb = RGBColor(0, 0, 0)
    heading3_style.paragraph_format.space_before = Pt(8)
    heading3_style.paragraph_format.space_after = Pt(4)


def add_heading_with_style(doc: Document, text: str, level: int):
    """添加带有正确样式的标题 - 优化版"""
    heading = doc.add_heading(text, level=level)
    
    # 快速路径：如果已经有运行对象，只设置第一个run的样式即可
    if heading.runs:
        run = heading.runs[0]
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
        
        # 根据标题级别设置字体大小
        if level == 1:
            run.font.size = Pt(18)
        elif level == 2:
            run.font.size = Pt(16)
        elif level == 3:
            run.font.size = Pt(14)
        
        run.font.bold = True
    
    return heading


def generate_table_of_contents(doc: Document):
    """在封面后插入Word原生目录域，打开后在Word中更新域即可生成带页码、可点击的目录"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # 添加"目录"标题（使用普通段落+样式，避免作为Heading出现在目录自身）
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("目　　录")
    title_run.font.name = FONT_NAME
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    title_para.paragraph_format.space_after = Pt(12)

    # 插入TOC域代码：TOC \o "1-3" \h \z \u
    # \o "1-3" — 捕获大纲级别1-3的标题
    # \h       — 目录条目自动变为超链接（可点击跳转）
    # \z       — Web版式视图中隐藏页码和制表符前导符
    # \u       — 使用段落的大纲级别（Heading样式自动带大纲级别）
    toc_para = doc.add_paragraph()
    toc_run = toc_para.add_run()

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    toc_run._element.append(fld_begin)

    toc_run2 = toc_para.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    toc_run2._element.append(instr)

    toc_run3 = toc_para.add_run()
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    toc_run3._element.append(fld_sep)

    toc_run4 = toc_para.add_run()
    toc_run4.text = '（请在Word中右键点击此处，选择"更新域" → "更新整个目录" 以生成完整目录）'
    toc_run4.font.color.rgb = RGBColor(160, 160, 160)
    toc_run4.font.size = Pt(11)
    toc_run4.font.name = FONT_NAME
    toc_run4._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

    toc_run5 = toc_para.add_run()
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    toc_run5._element.append(fld_end)

    doc.add_page_break()


def generate_cover_page(doc: Document):
    """生成封面页"""
    from datetime import datetime
    
    # 添加公司logo
    logo_path = os.path.join(os.path.dirname(__file__), "LOGO.jpg")
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    if logo_path and os.path.exists(logo_path):
        try:
            logo_run = logo_para.add_run()
            logo_run.add_picture(logo_path, width=Inches(1.5))
        except Exception as e:
            print(f"添加logo失败: {e}")
            logo_run = logo_para.add_run("[LOGO]")
            logo_run.font.size = Pt(24)
            logo_run.font.color.rgb = RGBColor(0, 51, 102)
    else:
        logo_run = logo_para.add_run("[LOGO]")
        logo_run.font.size = Pt(24)
        logo_run.font.color.rgb = RGBColor(0, 51, 102)
    
    # 主标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.style = 'ProfessionalTitle'
    title_run = title_para.add_run("MySQL数据库巡检汇总报告")
    
    # 副标题
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("多实例批量健康检查与分析汇总")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.name = FONT_NAME
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
    subtitle_run.italic = True
    
    doc.add_paragraph().add_run("\n" * 6)
    
    # 专业信息表格
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(4.2)
        for cell in row.cells:
            cell.paragraphs[0].style = doc.styles['Normal']
    
    # 表头行
    header_cells = table.rows[0].cells
    header_cells[0].merge(header_cells[1])
    header_cells[0].text = "巡检汇总基本信息"
    header_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header_cells[0].paragraphs[0].runs:
        header_cells[0].paragraphs[0].runs[0].font.bold = True
        header_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # 设置表头背景色
    shading_elm = parse_xml(r'<w:shd {} w:fill="003366"/>'.format(nsdecls('w')))
    header_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
    
    # 填充表格数据
    data = [
        ("报告编号", f"MYSQL-BATCH-{datetime.now().strftime('%Y%m%d')}"),
        ("汇总时间", datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
        ("巡检范围", "批量多实例数据库巡检"),
        ("巡检人员", "自动化巡检系统"),
        ("报告版本", "V2.0 并行处理版"),
        ("生成时间", datetime.now().strftime('%Y年%m月%d日 %H:%M:%S'))
    ]
    
    for i, (label, value) in enumerate(data, 1):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = str(value)
        
        if table.cell(i, 0).paragraphs[0].runs:
            table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
            table.cell(i, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    doc.add_paragraph().add_run("\n" * 4)
    
    # 添加保密声明
    confidential_para = doc.add_paragraph()
    confidential_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    confidential_run = confidential_para.add_run("本报告包含敏感信息，请妥善保管")
    confidential_run.font.size = Pt(10)
    confidential_run.font.color.rgb = RGBColor(153, 0, 0)
    confidential_run.italic = True
    
  


def generate_summary_content(doc: Document, inspection_data_list: list):
    """生成汇总内容"""
    add_heading_with_style(doc, "一、汇总执行摘要", level=1)
    
    # 统计总体问题
    total_instances = len(inspection_data_list)
    total_issues = 0
    total_critical = 0
    total_warning = 0
    total_info = 0
    
    instances_with_issues = 0
    
    for data in inspection_data_list:
        issues_summary = data.get('issues_summary', {})
        summary = issues_summary.get('summary', {})
        total_issues += summary.get('total_issues', 0)
        total_critical += summary.get('critical_count', 0)
        total_warning += summary.get('warning_count', 0)
        total_info += summary.get('info_count', 0)
        
        if summary.get('total_issues', 0) > 0:
            instances_with_issues += 1
    
    # 生成健康状态
    health_status = "良好"
    if total_critical > 0:
        health_status = "严重"
    elif total_warning > 10:
        health_status = "警告"
    elif total_warning > 0:
        health_status = "注意"
    
    summary_text = f"""
本次批量巡检共检查 {total_instances} 个数据库实例，其中 {instances_with_issues} 个实例存在问题。
总计发现 {total_issues} 个问题，按严重程度分类如下：

P0 严重问题: {total_critical} 个（最高优先级，需立即处理）
P1 警告问题: {total_warning} 个（中等优先级，需近期处理）
P2 信息提示: {total_info} 个（低优先级，可后续处理）

总体健康状态: {health_status}
"""
    
    doc.add_paragraph(summary_text)
    
    # 统计每个实例列表
    doc.add_paragraph()
    add_heading_with_style(doc, "1.1 巡检实例列表", level=2)
    
    for idx, data in enumerate(inspection_data_list, 1):
        inspection_data = data.get('inspection_data', {})
        summary_section = inspection_data.get("INSPECTION_SUMMARY", [])
        
        instance_ip = ""
        instance_port = ""
        inspection_time = ""
        
        for line in summary_section:
            if "实例IP:" in line:
                instance_ip = line.replace("实例IP:", "").strip()
            elif "实例端口:" in line:
                instance_port = line.replace("实例端口:", "").strip()
            elif "巡检时间:" in line:
                inspection_time = line.replace("巡检时间:", "").strip()
        
        issues_summary = data.get('issues_summary', {})
        summary = issues_summary.get('summary', {})
        critical_count = summary.get('critical_count', 0)
        warning_count = summary.get('warning_count', 0)
        
        instance_identifier = f"{instance_ip}:{instance_port}" if instance_ip and instance_port else "未知实例"
        
        instance_para = doc.add_paragraph()
        instance_para.add_run(f"{idx}. {instance_identifier}")
        
        if inspection_time:
            instance_para.add_run(f" - 巡检时间: {inspection_time}")
        
        if critical_count > 0 or warning_count > 0:
            issues_text = f" - 严重问题: {critical_count}个, 警告问题: {warning_count}个"
            run = instance_para.add_run(issues_text)
            run.font.color.rgb = RGBColor(255, 0, 0) if critical_count > 0 else RGBColor(255, 165, 0)
    
    doc.add_page_break()


def generate_detailed_content_by_instance(doc: Document, inspection_data_list: list):
    """生成详细内容 - 按IP+端口组织"""
    add_heading_with_style(doc, "二、详细问题汇总（按实例）", level=1)
    
    for idx, data in enumerate(inspection_data_list, 1):
        inspection_data = data.get('inspection_data', {})
        issues_summary = data.get('issues_summary', {})
        
        # 获取实例信息
        summary_section = inspection_data.get("INSPECTION_SUMMARY", [])
        instance_ip = ""
        instance_port = ""
        
        for line in summary_section:
            if "实例IP:" in line:
                instance_ip = line.replace("实例IP:", "").strip()
            elif "实例端口:" in line:
                instance_port = line.replace("实例端口:", "").strip()
        
        instance_identifier = f"{instance_ip}:{instance_port}" if instance_ip and instance_port else f"实例{idx}"
        
        # 添加实例标题
        add_heading_with_style(doc, f"实例{idx}: {instance_identifier}", level=2)
        
        # 显示该实例的问题
        critical_issues = issues_summary.get('critical_issues', [])
        warning_issues = issues_summary.get('warning_issues', [])
        info_issues = issues_summary.get('info_issues', [])
        
        issue_levels = [
            ("严重问题", critical_issues, "P0", RGBColor(255, 0, 0)),
            ("警告问题", warning_issues, "P1", RGBColor(255, 165, 0)),
            ("信息提示", info_issues, "P2", RGBColor(0, 128, 0))
        ]
        
        has_issues = False
        for level_name, issues, icon, color in issue_levels:
            if issues:
                has_issues = True
                level_para = doc.add_paragraph()
                level_run = level_para.add_run(f"{icon} {level_name}（{len(issues)}个）")
                level_run.font.bold = True
                level_run.font.color.rgb = color
                
                for issue_idx, issue in enumerate(issues, 1):
                    issue_text = f"{issue_idx}. 【{issue['category']}】{issue['issue']}\n"
                    issue_text += f"   描述: {issue['description']}\n"
                    issue_text += f"   建议: {issue['suggestion']}\n"
                    doc.add_paragraph(issue_text)
        
        if not has_issues:
            doc.add_paragraph("该实例无问题，状态良好")
        
        doc.add_paragraph()
        
        # 检查是否有MGR相关问题
        has_mgr_issues = False
        for issue in critical_issues + warning_issues:
            if "MGR" in issue.get("issue", "") or "集群" in issue.get("category", ""):
                has_mgr_issues = True
                break
        
        # 添加MGR集群信息表格（只有在有问题时才显示）
        mgr_members_data = inspection_data.get("MGR_GROUP_MEMBERS", [])
        mgr_stats_data = inspection_data.get("MGR_MEMBER_STATS", [])
        
        if (mgr_members_data or mgr_stats_data) and has_mgr_issues:
            add_heading_with_style(doc, "2.1 MGR集群信息", level=3)
            
            # 显示MGR集群成员信息
            if mgr_members_data:
                doc.add_paragraph("MGR集群成员信息：")
                
                # 创建表格
                mgr_members_table = doc.add_table(rows=1, cols=4)
                mgr_members_table.style = 'Light Grid Accent 1'
                mgr_members_table.autofit = False
                mgr_members_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 设置表格列宽
                mgr_members_table.columns[0].width = Inches(1.5)
                mgr_members_table.columns[1].width = Inches(2.0)
                mgr_members_table.columns[2].width = Inches(1.5)
                mgr_members_table.columns[3].width = Inches(1.5)
                
                # 设置表头
                mgr_members_header = mgr_members_table.rows[0].cells
                mgr_members_header[0].text = "成员ID"
                mgr_members_header[1].text = "主机"
                mgr_members_header[2].text = "状态"
                mgr_members_header[3].text = "角色"
                
                # 设置表头样式
                for cell in mgr_members_header:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                
                # 设置表头背景色
                shading_elm = parse_xml(r'<w:shd {} w:fill="003366"/>'.format(nsdecls('w')))
                for cell in mgr_members_header:
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                
                # 填充数据
                for line in mgr_members_data:
                    stripped_line = line.strip()
                    if not stripped_line:
                        continue
                    parts = stripped_line.split('\t')
                    if len(parts) < 8:
                        continue
                    
                    row_cells = mgr_members_table.add_row().cells
                    row_cells[0].text = parts[1]  # 成员ID
                    row_cells[1].text = parts[2]  # 主机
                    row_cells[2].text = parts[4]  # 状态
                    row_cells[3].text = parts[5]  # 角色
                    
                    for i in range(4):
                        row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph()
            
            # 显示MGR成员统计信息
            if mgr_stats_data:
                doc.add_paragraph("MGR成员统计信息：")
                
                # 创建表格
                mgr_stats_table = doc.add_table(rows=1, cols=4)
                mgr_stats_table.style = 'Light Grid Accent 1'
                mgr_stats_table.autofit = False
                mgr_stats_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # 设置表格列宽
                mgr_stats_table.columns[0].width = Inches(1.5)
                mgr_stats_table.columns[1].width = Inches(1.8)
                mgr_stats_table.columns[2].width = Inches(1.8)
                mgr_stats_table.columns[3].width = Inches(1.8)
                
                # 设置表头
                mgr_stats_header = mgr_stats_table.rows[0].cells
                mgr_stats_header[0].text = "成员ID"
                mgr_stats_header[1].text = "计数提交"
                mgr_stats_header[2].text = "计数冲突"
                mgr_stats_header[3].text = "队列事务"
                
                # 设置表头样式
                for cell in mgr_stats_header:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.bold = True
                        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                
                # 设置表头背景色
                shading_elm = parse_xml(r'<w:shd {} w:fill="003366"/>'.format(nsdecls('w')))
                for cell in mgr_stats_header:
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                
                # 填充数据
                for line in mgr_stats_data:
                    stripped_line = line.strip()
                    if not stripped_line:
                        continue
                    parts = stripped_line.split('\t')
                    if len(parts) < 14:
                        continue
                    
                    row_cells = mgr_stats_table.add_row().cells
                    row_cells[0].text = parts[2]  # 成员ID
                    row_cells[1].text = parts[3]  # 计数提交
                    row_cells[2].text = parts[5]  # 计数冲突
                    row_cells[3].text = parts[10] # 队列事务
                    
                    for i in range(4):
                        row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph()
        
        # 每个实例之间加分隔线
        if idx < len(inspection_data_list):
            doc.add_paragraph("-" * 80)
            doc.add_paragraph()


def generate_batch_summary(batch_dir: str, log_files: list, success_count: int, failed_files: list):
    """生成批量分析摘要"""
    
    summary_content = [
        "MySQL巡检批量分析报告摘要",
        "=" * 50,
        f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"分析文件总数: {len(log_files)}",
        f"成功分析: {success_count}",
        f"分析失败: {len(failed_files)}",
        "",
        "文件列表:",
        ""
    ]
    
    for log_file in log_files:
        filename = os.path.basename(log_file)
        status = "成功" if not any(f[0] == filename for f in failed_files) else "失败"
        summary_content.append(f"  {filename} - {status}")
    
    if failed_files:
        summary_content.extend([
            "",
            "失败详情:",
            ""
        ])
        
        for filename, error in failed_files:
            summary_content.append(f"  {filename}:")
            summary_content.append(f"    错误: {error}")
    
    # 保存摘要
    summary_file = os.path.join(batch_dir, "批量分析摘要.txt")
    with open(summary_file, 'w', encoding='utf-8') as f:
        f.write('\n'.join(summary_content))


def main():
    """主函数"""
    
    import argparse
    
    parser = argparse.ArgumentParser(description='批量分析MySQL巡检日志')
    parser.add_argument('logs_directory', help='日志目录路径')
    parser.add_argument('-f', '--format', default='docx', choices=['docx', 'txt'],
                      help='输出格式 (默认: docx)')
    parser.add_argument('-w', '--workers', type=int, default=None,
                      help='最大并行工作进程数 (默认: CPU核心数-1)')
    parser.add_argument('--excel', action='store_true',
                      help='额外生成Excel统计表格（需安装openpyxl）')
    parser.add_argument('--font', default='微软雅黑',
                      help='Word/Excel报告字体 (默认: 微软雅黑)')
    
    args = parser.parse_args()
    
    logs_dir = args.logs_directory
    output_format = args.format
    max_workers = args.workers
    generate_excel = args.excel
    
    global FONT_NAME
    FONT_NAME = args.font
    
    # 提前检查依赖
    if generate_excel and not EXCEL_AVAILABLE:
        print("错误: 生成Excel统计表格需要 openpyxl 库，但当前未安装")
        print("请执行以下命令安装: pip install openpyxl")
        sys.exit(1)
    
    if not os.path.exists(logs_dir):
        print(f"错误: 目录不存在 - {logs_dir}")
        sys.exit(1)
    
    if not os.path.isdir(logs_dir):
        print(f"错误: 不是目录 - {logs_dir}")
        sys.exit(1)
    
    try:
        batch_analyze_logs(logs_dir, output_format, max_workers, generate_excel)
    except Exception as e:
        print(f"批量分析失败: {str(e)}")
        import traceback
        print(f"详细错误: {traceback.format_exc()}")
        sys.exit(1)


if __name__ == "__main__":
    main()
